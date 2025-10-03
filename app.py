from flask import Flask, render_template, jsonify, request, Response, send_file
import psycopg2
import psycopg2.extras # Adicionado para DictCursor
import json
import math
from datetime import date, datetime, timedelta
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import io
import tempfile
from docxtpl import DocxTemplate
import os
import re
import shutil
import time
import unicodedata
from pypdf import PdfWriter, PdfReader
# Importação do docx2pdf será feita condicionalmente dentro da função

# Global map for action types (Hiperdia)
TIPO_ACAO_MAP_PY = {
    1: "Iniciar MRPA",
    2: "Avaliar MRPA",
    3: "Modificar tratamento",
    4: "Solicitar Exames",
    5: "Avaliar Exames",
    6: "Avaliar RCV",
    7: "Encaminhar para nutrição",
    8: "Registrar consulta nutrição",
    9: "Agendar Hiperdia",
    10: "Encaminhar Cardiologia",
    11: "Registrar Cardiologia"
}

# Global map for diabetes action types (Hiperdia Diabetes)
TIPO_ACAO_DIABETES_MAP_PY = {
    1: "Agendar Novo Acompanhamento",
    2: "Solicitar Hemoglobina Glicada, Glicemia Média e Glicemia de Jejum",
    3: "Solicitar Mapeamento Residencial de Glicemias",
    4: "Avaliar Tratamento",
    5: "Modificar Tratamento",
    6: "Finalizar Acompanhamento"
}

# Global map for Plafam action types
TIPO_ACAO_PLAFAM_MAP_PY = {
    1: "Convite com o Agente",
    2: "Convite com o Cliente", 
    3: "Deseja iniciar após convite",
    4: "Deseja iniciar via consulta",
    5: "Já em uso - Mensal",
    6: "Já em uso - Trimestral",
    7: "Já em uso - Pílula",
    8: "Já em uso - DIU",
    9: "Já em uso - Implante",
    10: "Já em uso - Laqueadura/Histerectomia",
    11: "Já em uso - Vasectomia",
    12: "Já em uso - Outros",
    13: "Cliente não encontrado",
    14: "Reavaliar em 6 meses",
    15: "Reavaliar em 1 ano",
    16: "Fora da área - Outra área",
    17: "Fora da área - Não reside na cidade",
    18: "Fora da área - Sem informações"
}
# Função auxiliar para conversão segura para float
def safe_float_conversion(value):
    try:
        if value is None:
            return None
        return float(value)
    except (ValueError, TypeError):
        print(f"Warning: Could not convert value '{value}' to float. Returning None.")
        return None

# Função auxiliar para remover acentos
def remove_acentos(texto):
    """Remove acentos do texto"""
    if not texto:
        return texto
    return unicodedata.normalize('NFD', texto).encode('ascii', 'ignore').decode('ascii')

# Funções para formatação de medicamentos de diabetes
def formatar_medicamento_oral_diabetes(nome, dose, frequencia, posologia=None, observacoes=None):
    """
    Formatar medicamento oral conforme especificações:
    - Calcular comprimidos mensais
    - Distribuir horários automaticamente
    """
    
    # Calcular quantidade mensal
    comprimidos_mes = dose * frequencia * 30
    quantidade_texto = f"{comprimidos_mes:02d} comprimidos"
    
    # Distribuir horários baseado na frequência
    horarios_map = {
        1: ["06:00"],
        2: ["06:00", "18:00"], 
        3: ["06:00", "14:00", "22:00"],
        4: ["06:00", "12:00", "18:00", "22:00"]
    }
    
    horarios = horarios_map.get(frequencia, ["06:00"])
    
    # Montar instruções
    instrucoes = []
    
    # Instruções por horário individual
    for horario in horarios:
        instrucao = f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} às {horario} horas"
        instrucoes.append(instrucao)
    
    # Adicionar posologia se existir
    if posologia and posologia.strip():
        instrucoes.append(f"Posologia: {posologia}")
    
    # Adicionar observações se existir  
    if observacoes and observacoes.strip():
        instrucoes.append(f"Obs: {observacoes}")
    
    return {
        'nome': nome,
        'quantidade': quantidade_texto,
        'instrucoes': instrucoes
    }

def formatar_insulina_diabetes(tipo_insulina, doses_estruturadas_json, observacoes=None):
    """
    Formatar insulina conforme especificações:
    - Calcular canetas (300U por caneta)
    - Aplicar regra "ao deitar" para último horário
    """
    
    try:
        doses = json.loads(doses_estruturadas_json) if isinstance(doses_estruturadas_json, str) else doses_estruturadas_json
    except (json.JSONDecodeError, TypeError):
        doses = []
    
    if not doses:
        return None
    
    # Calcular unidades totais diárias
    total_unidades_dia = sum(int(dose.get('dose', 0)) for dose in doses)
    
    # Calcular canetas (300U por caneta, para 30 dias)
    unidades_mes = total_unidades_dia * 30
    canetas_necessarias = math.ceil(unidades_mes / 300)
    quantidade_texto = f"{canetas_necessarias:02d} caneta{'s' if canetas_necessarias > 1 else ''}"
    
    # Processar horários e aplicar regra "ao deitar"
    instrucoes = []
    doses_ordenadas = sorted(doses, key=lambda x: x.get('horario', '00:00'))
    
    for i, dose in enumerate(doses_ordenadas):
        unidades = int(dose.get('dose', 0))
        horario = dose.get('horario', '00:00')
        
        # Regra: último horário sempre "ao deitar subcutâneo"
        if i == len(doses_ordenadas) - 1 and len(doses_ordenadas) > 1:
            instrucao = f"Aplicar {unidades:02d} unidades ao deitar subcutâneo"
        else:
            instrucao = f"Aplicar {unidades:02d} unidades às {horario} subcutâneo"
        
        instrucoes.append(instrucao)
    
    # Adicionar observações se existir
    if observacoes and observacoes.strip():
        instrucoes.append(f"Obs: {observacoes}")
    
    return {
        'nome': tipo_insulina,
        'quantidade': quantidade_texto,
        'instrucoes': instrucoes
    }
 
app = Flask(__name__)

# --- Configurações do Banco de Dados PostgreSQL ---
DB_HOST = "localhost"
DB_PORT = "5433"
DB_NAME = "esus"
DB_USER = "postgres"
DB_PASS = "EUC[x*x~Mc#S+H_Ui#xZBr0O~"

def get_db_connection():
    conn = psycopg2.connect(
        host=DB_HOST,
        port=DB_PORT,
        database=DB_NAME,
        user=DB_USER,
        password=DB_PASS
    )
    return conn

# --- Função Auxiliar para construir a query ---
def build_filtered_query(args):
    """Constrói a cláusula WHERE e ORDER BY com base nos filtros e ordenação da requisição."""
    equipe = args.get('equipe', 'Todas')
    search_term = args.get('search', None)
    metodos = args.getlist('metodo')
    faixas_etarias = args.getlist('faixa_etaria')
    status_list = args.getlist('status')
    status_timeline = args.get('status_timeline', 'Todos')  # Novo filtro para as abas
    sort_by = args.get('sort_by', 'nome_asc')
    
    # Filtros de aplicação
    aplicacao_data_inicial = args.get('aplicacao_data_inicial')
    aplicacao_data_final = args.get('aplicacao_data_final')
    aplicacao_metodo = args.get('aplicacao_metodo', 'trimestral')
    
    # Filtro para pacientes específicos (usado no plano semanal)
    pacientes_ids = args.get('pacientes_ids', None)

    print(f"DEBUG: status_timeline recebido: {status_timeline}")

    query_params = {}
    where_clauses = []

    # Filtro de Equipe
    if equipe != 'Todas':
        where_clauses.append("m.nome_equipe = %(equipe)s")
        query_params['equipe'] = equipe
    
    # Filtro de Microárea
    microarea = args.get('microarea', 'Todas as áreas')
    if microarea != 'Todas as áreas':
        if ' - ' in microarea:
            parts = microarea.split(' - ', 1)
            micro_area_str = parts[0].replace('Área ', '').strip()
        else:
            micro_area_str = microarea.replace('Área ', '').strip()
        
        if micro_area_str:
            where_clauses.append("m.microarea = %(microarea)s")
            query_params['microarea'] = micro_area_str
    
    # Filtro de Busca por Nome
    if search_term:
        where_clauses.append("UPPER(m.nome_paciente) LIKE UPPER(%(search)s)")
        query_params['search'] = f"%{search_term}%"

    # Filtro para pacientes específicos (usado no plano semanal)
    if pacientes_ids:
        print(f"DEBUG: Filtro pacientes_ids recebido: {pacientes_ids}")
        # Convert comma-separated string to list of integers
        try:
            ids_list = [int(id.strip()) for id in pacientes_ids.split(',') if id.strip()]
            if ids_list:
                print(f"DEBUG: IDs convertidos para lista: {ids_list}")
                where_clauses.append("m.cod_paciente = ANY(%(pacientes_ids)s)")
                query_params['pacientes_ids'] = ids_list
                print(f"DEBUG: Filtro de pacientes_ids adicionado à query")
        except (ValueError, AttributeError) as e:
            print(f"DEBUG: Erro ao processar pacientes_ids: {e}")
            # If invalid format, ignore the filter
            pass

    # Filtro de Métodos Contraceptivos
    if metodos:
        where_clauses.append("m.metodo = ANY(%(metodos)s)")
        query_params['metodos'] = metodos
        
    # Filtro de Faixa Etária
    if faixas_etarias:
        age_conditions = []
        for faixa in faixas_etarias:
            min_age, max_age = faixa.split('-')
            age_conditions.append(f"m.idade_calculada BETWEEN {int(min_age)} AND {int(max_age)}")
        if age_conditions:
            where_clauses.append(f"({ ' OR '.join(age_conditions) })")

    # Filtro de Status das Abas (status_timeline)
    if status_timeline != 'Todos':
        print(f"DEBUG: Aplicando filtro status_timeline: {status_timeline}")
        if status_timeline == 'SemMetodo':
            where_clauses.append("(m.metodo IS NULL OR m.metodo = '')")
            print("DEBUG: Adicionado filtro SemMetodo")
        elif status_timeline == 'Gestante':
            where_clauses.append("m.status_gravidez = 'Grávida'")
            print("DEBUG: Adicionado filtro Gestante")
        elif status_timeline == 'MetodoVencido':
            where_clauses.append("""
                (
                    (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                        ( (m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days') ) OR
                        ( m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days') )
                    ))
                )
            """)
            print("DEBUG: Adicionado filtro MetodoVencido")
        elif status_timeline == 'MetodoEmDia':
            where_clauses.append("""
                (
                    (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                        ( (m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days') ) OR
                        ( m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days') )
                    )) OR
                    ( m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%' )
                ) AND (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida')
            """)
            print("DEBUG: Adicionado filtro MetodoEmDia")

    # Filtro de Status de Acompanhamento (mantido para compatibilidade)
    if status_list:
        status_conditions = []
        if 'sem_metodo' in status_list:
            status_conditions.append("(m.metodo IS NULL OR m.metodo = '')")
        if 'gestante' in status_list:
            status_conditions.append("m.status_gravidez = 'Grávida'")
        if 'atrasado' in status_list:
            status_conditions.append("""
                (
                    (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                        ( (m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days') ) OR
                        ( m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days') )
                    ))
                )
            """)
        if 'em_dia' in status_list:
            status_conditions.append("""
                (
                    (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                        ( (m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days') ) OR
                        ( m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days') )
                    )) OR
                    ( m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%' )
                ) AND (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida')
            """)
        if status_conditions:
             where_clauses.append(f"({ ' OR '.join(status_conditions) })")

    # Filtro de Controle de Aplicações
    if aplicacao_data_inicial and aplicacao_data_final and aplicacao_metodo == 'trimestral':
        # Para o trimestral, buscar pacientes cuja próxima aplicação deve ocorrer no período escolhido
        # Se o usuário escolheu o período 01/10/2025 a 10/10/2025
        # Devemos buscar pacientes cuja aplicação + 90 dias está entre 01/10/2025 e 10/10/2025
        where_clauses.append("""
            (
                m.metodo ILIKE '%%trimestral%%' AND
                m.data_aplicacao IS NOT NULL AND
                m.data_aplicacao != '' AND
                (TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') + INTERVAL '90 days') BETWEEN 
                    TO_DATE(%(aplicacao_data_inicial)s, 'YYYY-MM-DD') AND
                    TO_DATE(%(aplicacao_data_final)s, 'YYYY-MM-DD')
            )
        """)
        query_params['aplicacao_data_inicial'] = aplicacao_data_inicial
        query_params['aplicacao_data_final'] = aplicacao_data_final
        print(f"DEBUG: Adicionado filtro de aplicações - Data inicial: {aplicacao_data_inicial}, Data final: {aplicacao_data_final}")

    # Filtro de Ações (status_acompanhamento)
    acao_filters = args.getlist('acao_filter')
    if acao_filters:
        print(f"DEBUG: Filtros de ações recebidos: {acao_filters}")
        # Converter filtros para inteiros, tratando "0" como None (Sem Ações)
        acao_conditions = []
        for acao in acao_filters:
            if acao == '0':  # Sem Ações
                acao_conditions.append("(pa.status_acompanhamento IS NULL)")
            else:
                try:
                    acao_int = int(acao)
                    acao_conditions.append(f"pa.status_acompanhamento = {acao_int}")
                except ValueError:
                    print(f"WARNING: Valor de ação inválido: {acao}")
        
        if acao_conditions:
            where_clauses.append(f"({ ' OR '.join(acao_conditions) })")
            print(f"DEBUG: Adicionado filtro de ações: {' OR '.join(acao_conditions)}")

    # Lógica de Ordenação
    sort_mapping = {
        'nome_asc': 'm.nome_paciente ASC',
        'nome_desc': 'm.nome_paciente DESC',
        'idade_asc': 'm.idade_calculada ASC',
        'idade_desc': 'm.idade_calculada DESC',
        'metodo_asc': """
            CASE
                -- Status do método para determinar ordem da data
                WHEN m.status_gravidez = 'Grávida' THEN 1
                WHEN (m.metodo IS NULL OR m.metodo = '') THEN 4
                WHEN m.data_aplicacao IS NULL OR m.data_aplicacao = '' THEN 3
                -- Método em atraso - ordenar por data mais recente primeiro (DESC)
                WHEN (
                    ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                    (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days')) OR
                    (m.metodo ILIKE '%%diu%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '3650 days')) OR
                    (m.metodo ILIKE '%%implante%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '1095 days'))
                ) THEN 2
                -- Método em dia - ordenar por datas mais próximas de hoje primeiro
                ELSE 3
            END ASC,
            CASE
                -- Para métodos em atraso: datas mais próximas de hoje primeiro (mais recentes)
                WHEN (
                    ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                    (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days')) OR
                    (m.metodo ILIKE '%%diu%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '3650 days')) OR
                    (m.metodo ILIKE '%%implante%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '1095 days'))
                ) THEN TO_DATE(m.data_aplicacao, 'DD/MM/YYYY')
                ELSE NULL
            END DESC NULLS LAST,
            CASE
                -- Para métodos em dia: datas mais próximas de hoje primeiro (ASC a partir de hoje)
                WHEN NOT (
                    ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                    (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days')) OR
                    (m.metodo ILIKE '%%diu%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '3650 days')) OR
                    (m.metodo ILIKE '%%implante%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '1095 days'))
                ) AND m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' THEN TO_DATE(m.data_aplicacao, 'DD/MM/YYYY')
                ELSE NULL
            END ASC NULLS LAST,
            m.nome_paciente ASC
        """,
        'status_asc': """
            CASE
                -- Status do método para determinar ordem da data
                WHEN m.status_gravidez = 'Grávida' THEN 1
                WHEN (m.metodo IS NULL OR m.metodo = '') THEN 4
                WHEN m.data_aplicacao IS NULL OR m.data_aplicacao = '' THEN 3
                -- Método em atraso - ordenar por data mais recente primeiro (DESC)
                WHEN (
                    ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                    (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days')) OR
                    (m.metodo ILIKE '%%diu%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '3650 days')) OR
                    (m.metodo ILIKE '%%implante%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '1095 days'))
                ) THEN 2
                -- Método em dia - ordenar por datas mais próximas de hoje primeiro
                ELSE 3
            END ASC,
            CASE
                -- Para métodos em atraso: datas mais próximas de hoje primeiro (mais recentes)
                WHEN (
                    ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                    (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days')) OR
                    (m.metodo ILIKE '%%diu%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '3650 days')) OR
                    (m.metodo ILIKE '%%implante%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '1095 days'))
                ) THEN TO_DATE(m.data_aplicacao, 'DD/MM/YYYY')
                ELSE NULL
            END DESC NULLS LAST,
            CASE
                -- Para métodos em dia: datas mais próximas de hoje primeiro (ASC a partir de hoje)
                WHEN NOT (
                    ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                    (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days')) OR
                    (m.metodo ILIKE '%%diu%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '3650 days')) OR
                    (m.metodo ILIKE '%%implante%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '1095 days'))
                ) AND m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' THEN TO_DATE(m.data_aplicacao, 'DD/MM/YYYY')
                ELSE NULL
            END ASC NULLS LAST,
            m.nome_paciente ASC
        """
    }
    # Verificar se há filtros de ações aplicados para ordenação especial
    acao_filters = args.getlist('acao_filter')
    if acao_filters:
        # Quando filtros de ações são aplicados, ordenar por data_acompanhamento ascendente (mais antiga para mais recente)
        order_by_clause = " ORDER BY pa.data_acompanhamento ASC NULLS LAST, m.nome_paciente ASC"
        print(f"DEBUG: Aplicando ordenação por data_acompanhamento ASC devido aos filtros de ações")
    else:
        # Usar ordenação padrão quando não há filtros de ações
        order_by_clause = " ORDER BY " + sort_mapping.get(sort_by, 'm.nome_paciente ASC')

    where_clause_str = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""
    
    print(f"DEBUG: where_clause_str final: {where_clause_str}")
    print(f"DEBUG: query_params: {query_params}")
    
    return where_clause_str, order_by_clause, query_params

# --- Rotas do Aplicativo Flask ---
@app.route('/') # Rota para a página inicial
def home():
    return render_template('index.html')

@app.route('/painel-plafam')
def painel_plafam():
    return render_template('Painel-Plafam.html')

@app.route('/painel-adolescentes')
def painel_adolescentes():
    return render_template('painel-adolescentes.html')

@app.route('/painel-hiperdia-has')
def painel_hiperdia_has():
    return render_template('painel-hiperdia-has.html')

@app.route('/painel-plafam-analise')
def painel_plafam_analise():
    return render_template('painel-plafam-analise.html')

@app.route('/painel-plafam-analise-adolescentes')
def painel_plafam_analise_adolescentes():
    return render_template('painel-plafam-analise-adolescentes.html')


@app.route('/api/pacientes_plafam')
def api_pacientes_plafam():
    conn = None
    cur = None
    try:
        print(f"DEBUG API: Requisição recebida com args: {dict(request.args)}")
        
        conn = get_db_connection()
        cur = conn.cursor()

        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 10))  # Limite configurável, padrão 10
        offset = (page - 1) * limit

        where_clause, order_by_clause, query_params = build_filtered_query(request.args)
        
        print(f"DEBUG API: where_clause construída: {where_clause}")
        print(f"DEBUG API: query_params: {query_params}")
        
        base_query = """
        SELECT
            m.cod_paciente, m.nome_paciente, m.cartao_sus, m.idade_calculada, m.microarea,
            m.metodo, m.nome_equipe, m.data_aplicacao, m.status_gravidez, m.data_provavel_parto,
            pa.status_acompanhamento, pa.data_acompanhamento,
            ag.nome_agente
        FROM sistemaaps.mv_plafam m
        LEFT JOIN sistemaaps.tb_plafam_acompanhamento pa ON m.cod_paciente = pa.co_cidadao
        LEFT JOIN sistemaaps.tb_agentes ag ON m.nome_equipe = ag.nome_equipe AND m.microarea = ag.micro_area
        """
        
        final_query = base_query + where_clause
        count_query = """SELECT COUNT(*) FROM sistemaaps.mv_plafam m
        LEFT JOIN sistemaaps.tb_plafam_acompanhamento pa ON m.cod_paciente = pa.co_cidadao
        LEFT JOIN sistemaaps.tb_agentes ag ON m.nome_equipe = ag.nome_equipe AND m.microarea = ag.micro_area""" + where_clause

        count_params = query_params.copy()
        
        cur.execute(count_query, count_params)
        total_pacientes = cur.fetchone()[0]
        print(f"DEBUG API: Total de pacientes encontrados: {total_pacientes}")
        
        query_params['limit'] = limit
        query_params['offset'] = offset
        final_query += order_by_clause + " LIMIT %(limit)s OFFSET %(offset)s"
        
        print(f"DEBUG API: Query final: {final_query}")
        print(f"DEBUG API: Parâmetros finais: {query_params}")
        
        cur.execute(final_query, query_params)
        dados = cur.fetchall()
        print(f"DEBUG API: Dados retornados: {len(dados)} registros")

        colunas_db = [desc[0] for desc in cur.description]
        colunas_frontend = [col.replace('microarea', 'micro_area').replace('status_gravidez', 'gestante') for col in colunas_db]
        
        resultados = []
        for linha in dados:
            linha_dict = dict(zip(colunas_frontend, linha))

            # --- Tratamento específico para data_aplicacao ---
            # Objetivo: Enviar para o frontend como 'YYYY-MM-DD' ou null.
            data_app_val = linha_dict.get('data_aplicacao')
            if data_app_val:
                if isinstance(data_app_val, date):  # Se for um objeto date do Python
                    linha_dict['data_aplicacao'] = data_app_val.strftime('%Y-%m-%d')
                elif isinstance(data_app_val, str):  # Se for uma string
                    try:
                        # Tenta converter de 'dd/mm/yyyy' (formato da view, como informado)
                        dt_obj = datetime.strptime(data_app_val, '%d/%m/%Y')
                        linha_dict['data_aplicacao'] = dt_obj.strftime('%Y-%m-%d')
                    except ValueError:
                        # Se não for 'dd/mm/yyyy', verifica se já é 'yyyy-mm-dd'
                        try:
                            datetime.strptime(data_app_val, '%Y-%m-%d')
                            # Se for, já está no formato correto, não precisa alterar.
                        except ValueError:
                            # Se não for nenhum dos formatos esperados, loga e define como None.
                            print(f"Alerta: Formato de data_aplicacao ('{data_app_val}') não reconhecido na API. Enviando como null.")
                            linha_dict['data_aplicacao'] = None
                else: # Algum outro tipo inesperado
                    print(f"Alerta: Tipo de data_aplicacao ({type(data_app_val)}) inesperado na API. Enviando como null.")
                    linha_dict['data_aplicacao'] = None
            else: # Se for None, False, 0, string vazia etc. vindo do banco
                linha_dict['data_aplicacao'] = None # Garante que seja null para o JS

            if linha_dict.get('data_provavel_parto') and isinstance(linha_dict['data_provavel_parto'], date):
                linha_dict['data_provavel_parto'] = linha_dict['data_provavel_parto'].strftime('%d/%m/%Y')
            if linha_dict.get('data_acompanhamento') and isinstance(linha_dict['data_acompanhamento'], date):
                linha_dict['data_acompanhamento'] = linha_dict['data_acompanhamento'].strftime('%d/%m/%Y')
            linha_dict['gestante'] = True if linha_dict.get('gestante') == 'Grávida' else False
            resultados.append(linha_dict)

        return jsonify({
            'pacientes': resultados,
            'total': total_pacientes,
            'page': page,
            'limit': limit,
            'pages': (total_pacientes + limit - 1) // limit if total_pacientes > 0 else 0
        })

    except Exception as e:
        print(f"Erro ao conectar ou consultar o banco de dados: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

@app.route('/api/update_acompanhamento', methods=['POST'])
def update_acompanhamento():
    data = request.get_json()
    co_cidadao = data.get('co_cidadao')
    status_str = data.get('status')

    if not co_cidadao or status_str is None:
        return jsonify({'sucesso': False, 'erro': 'Dados inválidos'}), 400

    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Mapear os códigos de status para ações do Plafam - NOVO SISTEMA COM SUBMENUS
        status_mapping = {
            # Convite
            '1': 1,      # Convite com o agente
            '2': 2,      # Convite entregue ao cliente
            # Deseja iniciar  
            '3': 3,      # Deseja iniciar (após convite)
            '4': 4,      # Deseja iniciar (via consulta)
            # Já em uso - Métodos específicos
            '5': 5,      # Já em uso - Mensal
            '6': 6,      # Já em uso - Vasectomia (esposo)
            '7': 7,      # Já em uso - Trimestral
            '8': 8,      # Já em uso - DIU
            '9': 9,      # Já em uso - Implante
            '10': 10,    # Já em uso - Laqueadura
            '11': 11,    # Já em uso - Histerectomia (esposo)
            '12': 12,    # Já em uso - Outros
            # Outros status
            '13': 13,    # Cliente não encontrado
            '14': 14,    # Reavaliar em 6 meses
            '15': 15,    # Reavaliar em 1 ano
            # Fora da área
            '16': 16,    # Fora da área - Outra área
            '17': 17,    # Fora da área - Não reside na cidade
            '18': 18,    # Fora da área - Sem informação
            '19': 19,    # Fora da área - Área indígena
            # Reset
            '0': None    # Resetar ações
        }
        
        status_code = status_mapping.get(status_str)
        
        if status_code is None:
            sql_upsert = """
                INSERT INTO sistemaaps.tb_plafam_acompanhamento (co_cidadao, status_acompanhamento, data_acompanhamento)
                VALUES (%(co_cidadao)s, NULL, NULL)
                ON CONFLICT (co_cidadao) DO UPDATE
                SET status_acompanhamento = NULL,
                    data_acompanhamento = NULL;
            """
            cur.execute(sql_upsert, {'co_cidadao': co_cidadao})
        else:
            sql_upsert = """
                INSERT INTO sistemaaps.tb_plafam_acompanhamento (co_cidadao, status_acompanhamento, data_acompanhamento)
                VALUES (%(co_cidadao)s, %(status)s, CURRENT_DATE)
                ON CONFLICT (co_cidadao) DO UPDATE
                SET status_acompanhamento = EXCLUDED.status_acompanhamento,
                    data_acompanhamento = EXCLUDED.data_acompanhamento;
            """
            cur.execute(sql_upsert, {'status': status_code, 'co_cidadao': co_cidadao})
        
        conn.commit()
        
        return jsonify({'sucesso': True, 'mensagem': 'Acompanhamento atualizado com sucesso!'})

    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro ao atualizar acompanhamento: {e}")
        return jsonify({'sucesso': False, 'erro': f"Erro no servidor: {e}"}), 500
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

@app.route('/api/export_data')
def api_export_data():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        where_clause, order_by_clause, query_params = build_filtered_query(request.args)

        base_query = """
        SELECT m.cod_paciente, m.nome_paciente, m.cartao_sus, m.idade_calculada, m.microarea, m.metodo, m.nome_equipe, 
               m.data_aplicacao, m.status_gravidez, m.data_provavel_parto, pa.status_acompanhamento, pa.data_acompanhamento,
               ag.nome_agente
        FROM sistemaaps.mv_plafam m
        LEFT JOIN sistemaaps.tb_plafam_acompanhamento pa ON m.cod_paciente = pa.co_cidadao
        LEFT JOIN sistemaaps.tb_agentes ag ON m.nome_equipe = ag.nome_equipe AND m.microarea = ag.micro_area
        """
        
        final_query = base_query + where_clause + order_by_clause
        
        cur.execute(final_query, query_params)
        dados = cur.fetchall()

        colunas_db = [desc[0] for desc in cur.description]
        colunas_frontend = [col.replace('microarea', 'micro_area').replace('status_gravidez', 'gestante') for col in colunas_db]
        
        resultados = []
        for linha in dados:
            linha_dict = dict(zip(colunas_frontend, linha))
            if linha_dict.get('data_aplicacao') and isinstance(linha_dict['data_aplicacao'], date):
                linha_dict['data_aplicacao'] = linha_dict['data_aplicacao'].strftime('%d/%m/%Y')
            if linha_dict.get('data_provavel_parto') and isinstance(linha_dict['data_provavel_parto'], date):
                linha_dict['data_provavel_parto'] = linha_dict['data_provavel_parto'].strftime('%d/%m/%Y')
            if linha_dict.get('data_acompanhamento') and isinstance(linha_dict['data_acompanhamento'], date):
                linha_dict['data_acompanhamento'] = linha_dict['data_acompanhamento'].strftime('%d/%m/%Y')
            linha_dict['gestante'] = True if linha_dict.get('gestante') == 'Grávida' else False
            resultados.append(linha_dict)

        return jsonify(resultados)

    except Exception as e:
        print(f"Erro na exportação: {e}")
        return jsonify({"erro": f"Erro no servidor durante a exportação: {e}"}), 500
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()


@app.route('/api/equipes')
def api_equipes():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT nome_equipe, COUNT(*) as total_pacientes
            FROM sistemaaps.mv_plafam
            WHERE nome_equipe IS NOT NULL
            GROUP BY nome_equipe
            ORDER BY total_pacientes DESC;
        """)
        equipes = [row[0] for row in cur.fetchall()]
        return jsonify(equipes)
    except Exception as e:
        print(f"Erro ao buscar equipes: {e}")
        return jsonify({"erro": f"Não foi possível buscar equipes: {e}"}), 500
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

@app.route('/api/equipes_com_agentes_adolescentes')
def api_equipes_com_agentes_adolescentes():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # Pega todas as equipes distintas de mv_plafam ou tb_agentes
        cur.execute("""
            SELECT DISTINCT nome_equipe 
            FROM (
                SELECT nome_equipe FROM sistemaaps.mv_plafam
                UNION
                SELECT nome_equipe FROM sistemaaps.tb_agentes
            ) AS equipes_unidas
            WHERE nome_equipe IS NOT NULL 
            ORDER BY nome_equipe
        """)
        equipes_db = cur.fetchall()
        
        resultado_final = []
        for eq_row in equipes_db:
            equipe_nome = eq_row['nome_equipe']
            
            # Contar adolescentes (14-18 anos) para a equipe atual
            cur.execute("""
                SELECT COUNT(DISTINCT m.cod_paciente)
                FROM sistemaaps.mv_plafam m
                WHERE m.nome_equipe = %s AND m.idade_calculada BETWEEN 14 AND 18
            """, (equipe_nome,))
            num_adolescentes = cur.fetchone()[0] or 0
            
            cur.execute("""
                SELECT micro_area, nome_agente 
                FROM sistemaaps.tb_agentes 
                WHERE nome_equipe = %s 
                ORDER BY micro_area, nome_agente
            """, (equipe_nome,))
            agentes_db = cur.fetchall()
            agentes = []
            if agentes_db:
                agentes = [{"micro_area": ag['micro_area'], "nome_agente": ag['nome_agente']} for ag in agentes_db]
            
            resultado_final.append({"nome_equipe": equipe_nome, "agentes": agentes, "num_adolescentes": num_adolescentes})
        
        # Ordenar as equipes pelo número de adolescentes em ordem decrescente
        resultado_final_ordenado = sorted(resultado_final, key=lambda x: x.get('num_adolescentes', 0), reverse=True)
                
        return jsonify(resultado_final_ordenado)
    except Exception as e:
        print(f"Erro ao buscar equipes com agentes: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()


@app.route('/api/estatisticas_painel_adolescentes')
def api_estatisticas_painel_adolescentes():
    equipe_req = request.args.get('equipe', 'Todas')
    # Espera "Área MA - Agente NOME" ou "Todas"
    agente_selecionado_req = request.args.get('agente_selecionado', 'Todas') 

    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        micro_area_filtro = None
        # nome_agente_filtro = None # Não usado diretamente no filtro SQL se mv_plafam não tiver nome_agente
        aplicar_filtro_agente_area = False

        if agente_selecionado_req != 'Todas' and agente_selecionado_req != 'Todas as áreas':
            # Exemplo: "Área 01 - Agente X" ou apenas "Área 01"
            if ' - ' in agente_selecionado_req:
                parts = agente_selecionado_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else: # Apenas "Área 01"
                micro_area_str = agente_selecionado_req.replace('Área ', '').strip()
            
            if micro_area_str: # micro_area_str pode ser apenas o número da área
                micro_area_filtro = micro_area_str
                aplicar_filtro_agente_area = True

        # --- Total de Adolescentes (14-18 anos) ---
        # Este é para o card "Adolescentes - 14 a 18 anos" (respeita filtro de equipe e agente/área)
        query_adolescentes_base = "SELECT COUNT(DISTINCT m.cod_paciente) FROM sistemaaps.mv_plafam m "
        
        # Cláusulas e parâmetros base para filtros de idade, equipe E área/agente
        base_where_clauses = ["m.idade_calculada BETWEEN 14 AND 18"]
        base_params = []

        if equipe_req != 'Todas':
            base_where_clauses.append("m.nome_equipe = %s")
            base_params.append(equipe_req)

        if aplicar_filtro_agente_area and micro_area_filtro:
            base_where_clauses.append("m.microarea = %s")
            base_params.append(micro_area_filtro)

        query_total_adolescentes_final = query_adolescentes_base + " WHERE " + " AND ".join(base_where_clauses)
        cur.execute(query_total_adolescentes_final, tuple(base_params))
        total_adolescentes = cur.fetchone()[0] or 0

        # --- Adolescentes (14-18) SEM MÉTODO (respeita filtro de equipe e agente/área) ---
        where_clauses_sem_metodo = list(base_where_clauses)
        condicao_sem_metodo = "(m.metodo IS NULL OR m.metodo = '')"
        where_clauses_sem_metodo.append(condicao_sem_metodo)
        query_adolescentes_sem_metodo_final = query_adolescentes_base + " WHERE " + " AND ".join(where_clauses_sem_metodo)
        cur.execute(query_adolescentes_sem_metodo_final, tuple(base_params)) # Reutiliza base_params pois a condição extra não tem placeholder
        adolescentes_sem_metodo_count = cur.fetchone()[0] or 0
        
        # --- Adolescentes (14-18) GESTANTES (respeita filtro de equipe e agente/área) ---
        where_clauses_gestantes = list(base_where_clauses)
        condicao_gestantes = "m.status_gravidez = 'Grávida'"
        where_clauses_gestantes.append(condicao_gestantes)
        query_adolescentes_gestantes_final = query_adolescentes_base + " WHERE " + " AND ".join(where_clauses_gestantes)
        cur.execute(query_adolescentes_gestantes_final, tuple(base_params))
        adolescentes_gestantes_count = cur.fetchone()[0] or 0

        # --- Adolescentes (14-18) COM MÉTODO EM DIA (respeita filtro de equipe e agente/área) ---
        where_clauses_metodo_em_dia = list(base_where_clauses)
        # A condição para 'MetodoEmDia' já está definida em build_timeline_query_filters, podemos reutilizar ou redefinir
        # Para clareza, vamos redefinir aqui, garantindo que não seja gestante.
        condicao_metodo_em_dia = """
        (
            (m.metodo IS NOT NULL AND m.metodo != '') AND 
            (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND 
            (
                (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                    ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                    (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                )) OR
                (m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%') 
            )
        )
        """
        where_clauses_metodo_em_dia.append(condicao_metodo_em_dia)
        query_adolescentes_metodo_em_dia_final = query_adolescentes_base + " WHERE " + " AND ".join(where_clauses_metodo_em_dia)
        cur.execute(query_adolescentes_metodo_em_dia_final, tuple(base_params))
        adolescentes_com_metodo_em_dia_count = cur.fetchone()[0] or 0

        # --- Adolescentes (14-18) COM MÉTODO ATRASADO (respeita filtro de equipe e agente/área) ---
        where_clauses_metodo_atrasado = list(base_where_clauses)
        condicao_metodo_atrasado = """
        (
            (m.metodo IS NOT NULL AND m.metodo != '') AND -- Garante que tem um método para estar atrasado
            (
                (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                    ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days'))
                )) OR
                (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                    (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days'))
                ))
            )
        )
        """
        where_clauses_metodo_atrasado.append(condicao_metodo_atrasado)
        query_adolescentes_metodo_atrasado_final = query_adolescentes_base + " WHERE " + " AND ".join(where_clauses_metodo_atrasado)
        cur.execute(query_adolescentes_metodo_atrasado_final, tuple(base_params))
        adolescentes_com_metodo_atrasado_count = cur.fetchone()[0] or 0

        return jsonify({
            "total_adolescentes": total_adolescentes, 
            "adolescentes_sem_metodo": adolescentes_sem_metodo_count,
            "adolescentes_com_metodo_atrasado": adolescentes_com_metodo_atrasado_count,
            "adolescentes_gestantes": adolescentes_gestantes_count,
            "adolescentes_metodo_em_dia": adolescentes_com_metodo_em_dia_count
        })
    except Exception as e:
        print(f"Erro ao buscar estatísticas de adolescentes: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/estatisticas_painel_plafam')
def api_estatisticas_painel_plafam():
    equipe_req = request.args.get('equipe', 'Todas')
    microarea_req = request.args.get('microarea', 'Todas as áreas')

    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # --- Total de Mulheres (14-45 anos) ---
        query_mulheres_base = "SELECT COUNT(DISTINCT m.cod_paciente) FROM sistemaaps.mv_plafam m "
        
        # Cláusulas e parâmetros base para filtros de idade, equipe E área/agente
        base_where_clauses = ["m.idade_calculada BETWEEN 14 AND 45"]
        base_params = []

        if equipe_req != 'Todas':
            base_where_clauses.append("m.nome_equipe = %s")
            base_params.append(equipe_req)

        if microarea_req != 'Todas as áreas':
            if ' - ' in microarea_req:
                parts = microarea_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else:
                micro_area_str = microarea_req.replace('Área ', '').strip()
            
            if micro_area_str:
                base_where_clauses.append("m.microarea = %s")
                base_params.append(micro_area_str)

        query_total_mulheres_final = query_mulheres_base + " WHERE " + " AND ".join(base_where_clauses)
        cur.execute(query_total_mulheres_final, tuple(base_params))
        total_mulheres = cur.fetchone()[0] or 0

        # --- Mulheres (14-45) SEM MÉTODO ---
        where_clauses_sem_metodo = list(base_where_clauses)
        condicao_sem_metodo = "(m.metodo IS NULL OR m.metodo = '')"
        where_clauses_sem_metodo.append(condicao_sem_metodo)
        query_mulheres_sem_metodo_final = query_mulheres_base + " WHERE " + " AND ".join(where_clauses_sem_metodo)
        cur.execute(query_mulheres_sem_metodo_final, tuple(base_params))
        mulheres_sem_metodo_count = cur.fetchone()[0] or 0
        
        # --- Mulheres (14-45) GESTANTES ---
        where_clauses_gestantes = list(base_where_clauses)
        condicao_gestantes = "m.status_gravidez = 'Grávida'"
        where_clauses_gestantes.append(condicao_gestantes)
        query_mulheres_gestantes_final = query_mulheres_base + " WHERE " + " AND ".join(where_clauses_gestantes)
        cur.execute(query_mulheres_gestantes_final, tuple(base_params))
        mulheres_gestantes_count = cur.fetchone()[0] or 0

        # --- Mulheres (14-45) COM MÉTODO EM DIA ---
        where_clauses_metodo_em_dia = list(base_where_clauses)
        condicao_metodo_em_dia = """
        (
            (m.metodo IS NOT NULL AND m.metodo != '') AND 
            (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND 
            (
                (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                    ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                    (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                )) OR
                (m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%') 
            )
        )
        """
        where_clauses_metodo_em_dia.append(condicao_metodo_em_dia)
        query_mulheres_metodo_em_dia_final = query_mulheres_base + " WHERE " + " AND ".join(where_clauses_metodo_em_dia)
        cur.execute(query_mulheres_metodo_em_dia_final, tuple(base_params))
        mulheres_com_metodo_em_dia_count = cur.fetchone()[0] or 0

        # --- Mulheres (14-45) COM MÉTODO ATRASADO ---
        where_clauses_metodo_atrasado = list(base_where_clauses)
        condicao_metodo_atrasado = """
        (
            (m.metodo IS NOT NULL AND m.metodo != '') AND 
            (
                (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                    ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days'))
                )) OR
                (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                    (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days'))
                ))
            )
        )
        """
        where_clauses_metodo_atrasado.append(condicao_metodo_atrasado)
        query_mulheres_metodo_atrasado_final = query_mulheres_base + " WHERE " + " AND ".join(where_clauses_metodo_atrasado)
        cur.execute(query_mulheres_metodo_atrasado_final, tuple(base_params))
        mulheres_com_metodo_atrasado_count = cur.fetchone()[0] or 0

        return jsonify({
            "total_adolescentes": total_mulheres,  # Mantém o nome para compatibilidade com o frontend
            "adolescentes_sem_metodo": mulheres_sem_metodo_count,
            "adolescentes_metodo_atrasado": mulheres_com_metodo_atrasado_count,  # Corrigido nome do campo
            "adolescentes_gestantes": mulheres_gestantes_count,
            "adolescentes_metodo_em_dia": mulheres_com_metodo_em_dia_count
        })
    except Exception as e:
        print(f"Erro ao buscar estatísticas do painel Plafam: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

def build_timeline_query_filters(args):
    equipe = args.get('equipe', 'Todas')
    agente_selecionado = args.get('agente_selecionado', 'Todas as áreas')
    search_term = args.get('search_timeline', None)
    status_timeline = args.get('status_timeline', 'Todos')
    sort_by = args.get('sort_by_timeline', 'proxima_acao_asc') # Novo padrão de ordenação
    proxima_acao = args.get('proxima_acao', 'all')
    metodo_filter = args.get('metodo_filter', 'all')

    query_params = {}
    where_clauses = ["m.idade_calculada BETWEEN 14 AND 18"] # Foco em adolescentes

    if equipe != 'Todas':
        where_clauses.append("m.nome_equipe = %(equipe)s")
        query_params['equipe'] = equipe

    if agente_selecionado != 'Todas as áreas':
        if ' - ' in agente_selecionado:
            parts = agente_selecionado.split(' - ', 1)
            micro_area_str = parts[0].replace('Área ', '').strip()
            # nome_agente_str = parts[1].strip() # Poderia ser usado se o filtro fosse mais granular
            if micro_area_str:
                where_clauses.append("m.microarea = %(microarea_timeline)s")
                query_params['microarea_timeline'] = micro_area_str
        # Se não tiver ' - ', pode ser apenas a microárea, mas a lógica atual do JS envia com ' - ' ou 'Todas as áreas'

    if search_term:
        where_clauses.append("(UPPER(m.nome_paciente) LIKE UPPER(%(search_timeline)s) OR UPPER(m.nome_responsavel) LIKE UPPER(%(search_timeline)s))")
        query_params['search_timeline'] = f"%{search_term}%"

    if status_timeline == 'SemMetodo':
        where_clauses.append("(m.metodo IS NULL OR m.metodo = '')")
    elif status_timeline == 'MetodoVencido':
        where_clauses.append("""
            (
                (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                    ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                    (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days'))
                ))
            )
        """)
    elif status_timeline == 'MetodoEmDia':
        where_clauses.append("""
            (
                (m.metodo IS NOT NULL AND m.metodo != '') AND -- Paciente tem um método
                (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND -- Paciente não está grávida
                (
                    (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                        ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                        (m.metodo ILIKE '%%trimestral%%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                    )) OR
                    (m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%') 
                )
            )
        """)
    elif status_timeline == 'Gestante':
        where_clauses.append("m.status_gravidez = 'Grávida'")
    # 'Todos' não adiciona filtro de status específico do método

    # Filtro por próxima ação
    if proxima_acao != 'all':
        try:
            tipo_acao = int(proxima_acao)
            where_clauses.append("pa_futura.tipo_abordagem = %(proxima_acao_tipo)s")
            query_params['proxima_acao_tipo'] = tipo_acao
        except (ValueError, TypeError):
            pass  # Ignorar valores inválidos

    # Filtro por método contraceptivo
    if metodo_filter != 'all':
        if metodo_filter == 'mensal':
            where_clauses.append("(m.metodo ILIKE %(metodo_pattern)s)")
            query_params['metodo_pattern'] = '%mensal%'
        elif metodo_filter == 'trimestral':
            where_clauses.append("(m.metodo ILIKE %(metodo_pattern)s)")
            query_params['metodo_pattern'] = '%trimestral%'
        elif metodo_filter == 'pilula':
            where_clauses.append("(m.metodo ILIKE %(metodo_pattern)s)")
            query_params['metodo_pattern'] = '%pílula%'
        elif metodo_filter == 'diu':
            where_clauses.append("(m.metodo ILIKE %(metodo_pattern)s)")
            query_params['metodo_pattern'] = '%diu%'
        elif metodo_filter == 'laqueadura':
            where_clauses.append("(m.metodo ILIKE %(metodo_pattern)s)")
            query_params['metodo_pattern'] = '%laqueadura%'
        elif metodo_filter == 'sem_metodo':
            where_clauses.append("(m.metodo IS NULL OR m.metodo = '')")
        elif metodo_filter == 'gestante':
            where_clauses.append("(m.status_gravidez = %(status_gravidez)s)")
            query_params['status_gravidez'] = 'Grávida'

    # Definir direção da ordenação baseada no filtro de status
    proxima_aplicacao_direction = "DESC"  # Padrão: mais recente primeiro
    if status_timeline == 'MetodoEmDia':
        proxima_aplicacao_direction = "ASC"  # Método em dia: mais antigo primeiro
    elif status_timeline == 'MetodoVencido':
        proxima_aplicacao_direction = "DESC"  # Método vencido: mais recente primeiro
    
    sort_mapping_timeline = {
        'nome_asc': 'm.nome_paciente ASC',
        'nome_desc': 'm.nome_paciente DESC',
        'idade_asc': 'm.idade_calculada ASC',
        'idade_desc': 'm.idade_calculada DESC',
        'metodo_asc': 'm.metodo ASC NULLS LAST, m.nome_paciente ASC',
        'proxima_aplicacao_desc': f"""
            CASE 
                WHEN m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' THEN
                    CASE 
                        WHEN m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%' THEN 
                            TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') + INTERVAL '30 days'
                        WHEN m.metodo ILIKE '%%trimestral%%' THEN 
                            TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') + INTERVAL '90 days'
                        ELSE TO_DATE(m.data_aplicacao, 'DD/MM/YYYY')
                    END
                ELSE NULL
            END {proxima_aplicacao_direction} NULLS LAST, m.nome_paciente ASC
        """,
        'proxima_acao_asc': 'data_proxima_acao_ordenacao ASC NULLS LAST, m.nome_paciente ASC'
    }
    # Default seguro para ordenação, caso sort_by seja inválido
    order_by_clause = " ORDER BY " + sort_mapping_timeline.get(sort_by, 'data_proxima_acao_ordenacao ASC NULLS LAST, m.nome_paciente ASC')
    
    where_clause_str = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""
    return where_clause_str, order_by_clause, query_params, proxima_acao

@app.route('/api/timeline_adolescentes')
def api_timeline_adolescentes(): # Rota para a timeline de adolescentes
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        page = int(request.args.get('page_timeline', 1))
        limit = int(request.args.get('limit', 10)) # Limite de 10 por página (padrão)
        offset = (page - 1) * limit

        where_clause, order_by_clause, query_params, proxima_acao = build_timeline_query_filters(request.args)

        base_query_fields = """
            m.cod_paciente, m.nome_paciente, m.cartao_sus, m.idade_calculada, m.microarea,
            m.metodo, m.nome_equipe, m.data_aplicacao, m.status_gravidez, m.data_provavel_parto,
            ag.nome_agente, m.nome_responsavel,
            pa_futura.data_acao AS data_proxima_acao_ordenacao,
            pa_futura.tipo_abordagem AS tipo_proxima_acao_ordenacao,
            pa_ultima.tipo_abordagem AS ultimo_tipo_abordagem,
            pa_ultima.resultado_abordagem AS ultimo_resultado_abordagem,
            pa_ultima.data_acao AS ultima_data_acao,
            pa_ultima.responsavel_pela_acao AS responsavel_pela_acao,
            pa_penultima.tipo_abordagem AS penultima_tipo_abordagem,
            pa_penultima.resultado_abordagem AS penultima_resultado_abordagem,
            pa_penultima.data_acao AS penultima_data_acao
        """
        
        from_join_clause = """
        FROM sistemaaps.mv_plafam m
        LEFT JOIN sistemaaps.tb_agentes ag ON m.nome_equipe = ag.nome_equipe AND m.microarea = ag.micro_area
        LEFT JOIN LATERAL (
            SELECT pa.data_acao, pa.tipo_abordagem
            FROM sistemaaps.tb_plafam_adolescentes pa
            WHERE pa.co_cidadao = m.cod_paciente
              AND pa.data_acao >= CURRENT_DATE
              AND pa.resultado_abordagem IS NULL
            ORDER BY pa.data_acao ASC
            LIMIT 1
        ) pa_futura ON TRUE
        LEFT JOIN LATERAL (
            SELECT pa.tipo_abordagem, pa.resultado_abordagem, pa.data_acao, pa.responsavel_pela_acao
            FROM sistemaaps.tb_plafam_adolescentes pa
            WHERE pa.co_cidadao = m.cod_paciente
              AND pa.resultado_abordagem IS NOT NULL
            ORDER BY pa.data_acao DESC, pa.co_abordagem DESC
            LIMIT 1
        ) pa_ultima ON TRUE
        LEFT JOIN LATERAL (
            SELECT pa.tipo_abordagem, pa.resultado_abordagem, pa.data_acao
            FROM sistemaaps.tb_plafam_adolescentes pa
            WHERE pa.co_cidadao = m.cod_paciente
              AND pa.resultado_abordagem IS NOT NULL
            ORDER BY pa.data_acao DESC, pa.co_abordagem DESC
            LIMIT 1 OFFSET 1
        ) pa_penultima ON TRUE
        """

        # Para contagem, usar FROM clause completo se filtro de próxima ação for aplicado
        if proxima_acao != 'all':
            from_clause_for_count = from_join_clause
        else:
            # FROM clause simplificado para a contagem, para evitar problemas de performance
            from_clause_for_count = """
        FROM sistemaaps.mv_plafam m
        LEFT JOIN sistemaaps.tb_agentes ag ON m.nome_equipe = ag.nome_equipe AND m.microarea = ag.micro_area
        LEFT JOIN LATERAL (
            SELECT pa.tipo_abordagem, pa.resultado_abordagem, pa.data_acao, pa.responsavel_pela_acao
            FROM sistemaaps.tb_plafam_adolescentes pa
            WHERE pa.co_cidadao = m.cod_paciente
              AND pa.resultado_abordagem IS NOT NULL
            ORDER BY pa.data_acao DESC, pa.co_abordagem DESC
            LIMIT 1
        ) pa_ultima ON TRUE
        """
        count_query = "SELECT COUNT(DISTINCT m.cod_paciente) " + from_clause_for_count + where_clause
        cur.execute(count_query, query_params)
        total_adolescentes = cur.fetchone()[0] or 0

        query_params_paginated = query_params.copy()
        query_params_paginated['limit'] = limit
        query_params_paginated['offset'] = offset
        final_query = "SELECT " + base_query_fields + from_join_clause + where_clause + order_by_clause + " LIMIT %(limit)s OFFSET %(offset)s"
        cur.execute(final_query, query_params_paginated)
        
        dados_adolescentes = []
        for row in cur.fetchall():
            row_dict = dict(row)
            # Inicializa os campos de próxima ação
            data_prox_acao_ord = row_dict.pop('data_proxima_acao_ordenacao', None)
            tipo_prox_acao_ord = row_dict.pop('tipo_proxima_acao_ordenacao', None)

            if data_prox_acao_ord and isinstance(data_prox_acao_ord, date):
                row_dict['proxima_acao_data_formatada'] = data_prox_acao_ord.strftime('%d/%m/%Y')
                tipo_abordagem_map_py = {
                    1: "Abordagem com pais",
                    2: "Abordagem direta com adolescente",
                    3: "Consulta na UBS",
                    4: "Entrega de convite",
                    5: "Fora de área",
                    6: "Iniciar método em domicílio",
                    7: "Remover do acompanhamento",
                    8: "Atualizar no PEC"
                }
                row_dict['proxima_acao_descricao'] = tipo_abordagem_map_py.get(tipo_prox_acao_ord, 'Ação futura')
                row_dict['proxima_acao_tipo'] = tipo_prox_acao_ord
            else:
                row_dict['proxima_acao_data_formatada'] = None
                row_dict['proxima_acao_descricao'] = None
                row_dict['proxima_acao_tipo'] = None

            # Processar dados da última ação realizada
            ultima_data_acao = row_dict.get('ultima_data_acao')
            if ultima_data_acao and isinstance(ultima_data_acao, date):
                row_dict['ultima_data_acao_formatada'] = ultima_data_acao.strftime('%d/%m/%Y')
            else:
                row_dict['ultima_data_acao_formatada'] = None

            # Processar dados da penúltima ação realizada
            penultima_data_acao = row_dict.get('penultima_data_acao')
            if penultima_data_acao and isinstance(penultima_data_acao, date):
                row_dict['penultima_data_acao_formatada'] = penultima_data_acao.strftime('%d/%m/%Y')
            else:
                row_dict['penultima_data_acao_formatada'] = None


            # Tratamento de datas para cada linha
            data_app_val = row_dict.get('data_aplicacao')
            if data_app_val:
                if isinstance(data_app_val, date):
                    row_dict['data_aplicacao'] = data_app_val.strftime('%Y-%m-%d')
                elif isinstance(data_app_val, str):
                    try:
                        dt_obj = datetime.strptime(data_app_val, '%d/%m/%Y')
                        row_dict['data_aplicacao'] = dt_obj.strftime('%Y-%m-%d')
                    except ValueError:
                        try:
                            datetime.strptime(data_app_val, '%Y-%m-%d')
                            # Se já estiver no formato YYYY-MM-DD, não faz nada
                        except ValueError:
                            row_dict['data_aplicacao'] = None # Formato não reconhecido
                else:
                    row_dict['data_aplicacao'] = None # Tipo inesperado
            else:
                row_dict['data_aplicacao'] = None # Valor original era nulo ou vazio

            if row_dict.get('data_provavel_parto') and isinstance(row_dict.get('data_provavel_parto'), date):
                row_dict['data_provavel_parto'] = row_dict.get('data_provavel_parto').strftime('%d/%m/%Y')
            
            dados_adolescentes.append(row_dict)

        return jsonify({
            'adolescentes': dados_adolescentes,
            'total': total_adolescentes,
            'page': page,
            'limit': limit,
            'pages': (total_adolescentes + limit - 1) // limit if total_adolescentes > 0 else 0
        })
    except Exception as e:
        print(f"Erro na API timeline_adolescentes: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/adolescente_detalhes_timeline/<co_cidadao>')
def api_adolescente_detalhes_timeline(co_cidadao):
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        # Buscar detalhes básicos da adolescente em mv_plafam
        cur.execute("""
            SELECT m.nome_paciente, m.idade_calculada, m.microarea, m.nome_equipe, m.cartao_sus, m.nome_responsavel,
                   ag.nome_agente,
                   m.metodo, m.data_aplicacao, m.status_gravidez, m.data_provavel_parto -- Para status no modal 
            FROM sistemaaps.mv_plafam m
            LEFT JOIN sistemaaps.tb_agentes ag ON m.nome_equipe = ag.nome_equipe AND m.microarea = ag.micro_area
            WHERE m.cod_paciente = %s AND m.idade_calculada BETWEEN 14 AND 18
        """, (co_cidadao,))
        detalhes_mv = cur.fetchone()

        if not detalhes_mv:
            return jsonify({"erro": "Adolescente não encontrada ou fora da faixa etária."}), 404
        
        # O nome_responsavel agora vem diretamente de mv_plafam

        cur.execute("""
            SELECT co_abordagem, tipo_abordagem, resultado_abordagem, observacoes, data_acao, responsavel_pela_acao
            FROM sistemaaps.tb_plafam_adolescentes
            WHERE co_cidadao = %(co_cidadao)s
            ORDER BY data_acao DESC, co_abordagem DESC
        """, {'co_cidadao': co_cidadao})
        eventos_timeline_db = cur.fetchall()

        eventos_timeline = []
        for evento in eventos_timeline_db:
            evento_dict = dict(evento)
            if evento_dict['data_acao'] and isinstance(evento_dict['data_acao'], date):
                evento_dict['data_acao'] = evento_dict['data_acao'].strftime('%Y-%m-%d')
            eventos_timeline.append(evento_dict)

        detalhes_finais = dict(detalhes_mv)
        detalhes_finais['proxima_acao_data_formatada'] = None
        detalhes_finais['proxima_acao_descricao'] = None

        # Buscar a próxima ação futura específica para esta adolescente para o modal
        cur.execute("""
            SELECT tipo_abordagem, data_acao
            FROM sistemaaps.tb_plafam_adolescentes
            WHERE co_cidadao = %(co_cidadao)s 
              AND data_acao >= CURRENT_DATE  -- Mudança para >=
              AND resultado_abordagem IS NULL
            ORDER BY data_acao ASC
            LIMIT 1;
        """, {'co_cidadao': co_cidadao})
        proxima_acao_especifica = cur.fetchone()

        if proxima_acao_especifica:
            tipo_abordagem_map_py = {
                1: "Abordagem com pais", 2: "Abordagem direta com adolescente",
                3: "Consulta na UBS", 4: "Entrega de convite"
            }
            detalhes_finais['proxima_acao_data_formatada'] = proxima_acao_especifica['data_acao'].strftime('%d/%m/%Y')
            detalhes_finais['proxima_acao_descricao'] = tipo_abordagem_map_py.get(proxima_acao_especifica['tipo_abordagem'], 'Ação futura')       

        return jsonify({
            "detalhes": detalhes_finais,
            "eventos_timeline": eventos_timeline
        })

    except Exception as e:
        print(f"Erro em api_adolescente_detalhes_timeline: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/adolescente/registrar_acao', methods=['POST'])
def api_registrar_acao_adolescente():
    data = request.get_json()
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        co_cidadao = data['co_cidadao']
        nome_adolescente = data.get('nome_adolescente')
        nome_responsavel_atual = data.get('nome_responsavel_atual')
        acao_atual_payload = data['acao_atual']
        proxima_acao_payload = data.get('proxima_acao')

        # Mapping function to convert method names to numeric codes
        def convert_method_to_code(method_name):
            if not method_name:
                return None
            method_mapping = {
                'Pílula anticoncepcional': 1,
                'Injetável mensal': 2,
                'Injetável trimestral': 3,
                'DIU': 4,
                'Implante subdérmico': 5
            }
            return method_mapping.get(method_name, None)

        # Convert method names to codes in payloads
        if acao_atual_payload and 'metodo_desejado' in acao_atual_payload:
            acao_atual_payload['metodo_desejado'] = convert_method_to_code(acao_atual_payload['metodo_desejado'])
        
        if proxima_acao_payload and 'metodo_desejado' in proxima_acao_payload:
            proxima_acao_payload['metodo_desejado'] = convert_method_to_code(proxima_acao_payload['metodo_desejado'])

        # Obter o próximo valor para co_abordagem de forma segura para múltiplas inserções
        cur.execute("SELECT MAX(co_abordagem) FROM sistemaaps.tb_plafam_adolescentes")
        max_co_abordagem = cur.fetchone()[0]
        next_co_abordagem_counter = (max_co_abordagem + 1) if max_co_abordagem is not None else 1

        # --- Etapa 1: Lidar com a Ação Atual ---
        # Verificar se existe uma ação futura pendente para esta adolescente para ser atualizada
        cur.execute("""
            SELECT co_abordagem
            FROM sistemaaps.tb_plafam_adolescentes
            WHERE co_cidadao = %(co_cidadao)s 
              AND data_acao >= CURRENT_DATE 
              AND resultado_abordagem IS NULL  -- Adicionado: apenas ações futuras PENDENTES
            ORDER BY data_acao ASC, co_abordagem ASC -- Garante seleção consistente
            LIMIT 1;
        """, {'co_cidadao': co_cidadao})
        pending_future_action = cur.fetchone()

        if pending_future_action:
            co_abordagem_to_update = pending_future_action[0]
            # Atualizar a ação futura existente com os detalhes da "acao_atual"
            cur.execute("""
                UPDATE sistemaaps.tb_plafam_adolescentes
                SET tipo_abordagem = %(tipo_abordagem)s,
                    resultado_abordagem = %(resultado_abordagem)s,
                    observacoes = %(observacoes)s,
                    data_acao = %(data_acao)s, -- Data em que a ação foi efetivamente realizada
                    responsavel_pela_acao = %(responsavel_pela_acao)s,
                    metodo_desejado = %(metodo_desejado)s,
                    nome_responsavel = %(nome_responsavel)s, -- Atualiza nome do responsável
                    nome_adolescente = %(nome_adolescente)s -- Atualiza nome da adolescente
                WHERE co_abordagem = %(co_abordagem)s;
            """, {
                'tipo_abordagem': acao_atual_payload.get('tipo_abordagem'),
                'resultado_abordagem': acao_atual_payload.get('resultado_abordagem'),
                'observacoes': acao_atual_payload.get('observacoes'),
                'data_acao': acao_atual_payload.get('data_acao'),
                'responsavel_pela_acao': acao_atual_payload.get('responsavel_pela_acao'),
                'metodo_desejado': acao_atual_payload.get('metodo_desejado'),
                'nome_responsavel': nome_responsavel_atual,
                'nome_adolescente': nome_adolescente,
                'co_abordagem': co_abordagem_to_update
            })
        else:
            # Nenhuma ação futura pendente, então insere "acao_atual" como um novo registro
            cur.execute("""
                INSERT INTO sistemaaps.tb_plafam_adolescentes
                (co_abordagem, co_cidadao, nome_adolescente, nome_responsavel,
                 tipo_abordagem, resultado_abordagem, observacoes,
                 data_acao, responsavel_pela_acao, metodo_desejado)
                VALUES (%(co_abordagem)s, %(co_cidadao)s, %(nome_adolescente)s, %(nome_responsavel)s,
                        %(tipo_abordagem)s, %(resultado_abordagem)s, %(observacoes)s,
                        %(data_acao)s, %(responsavel_pela_acao)s, %(metodo_desejado)s);
            """, {
                'co_abordagem': next_co_abordagem_counter,
                'co_cidadao': co_cidadao,
                'nome_adolescente': nome_adolescente,
                'nome_responsavel': nome_responsavel_atual,
                'tipo_abordagem': acao_atual_payload.get('tipo_abordagem'),
                'resultado_abordagem': acao_atual_payload.get('resultado_abordagem'),
                'observacoes': acao_atual_payload.get('observacoes'),
                'data_acao': acao_atual_payload.get('data_acao'),
                'responsavel_pela_acao': acao_atual_payload.get('responsavel_pela_acao'),
                'metodo_desejado': acao_atual_payload.get('metodo_desejado')
            })
            next_co_abordagem_counter += 1 # Incrementar para a próxima possível inserção

        # --- Etapa 2: Lidar com a Próxima Ação (se fornecida) ---
        # Esta é sempre uma inserção de uma nova ação futura
        if proxima_acao_payload:
            cur.execute("""
                INSERT INTO sistemaaps.tb_plafam_adolescentes
                (co_abordagem, co_cidadao, nome_adolescente, nome_responsavel,
                 tipo_abordagem, resultado_abordagem, observacoes,
                 data_acao, responsavel_pela_acao, metodo_desejado)
                VALUES (%(co_abordagem)s, %(co_cidadao)s, %(nome_adolescente)s, %(nome_responsavel)s,
                        %(tipo_abordagem)s, %(resultado_abordagem)s, %(observacoes)s,
                        %(data_acao)s, %(responsavel_pela_acao)s, %(metodo_desejado)s);
            """, {
                'co_abordagem': next_co_abordagem_counter, # Usa o contador, possivelmente incrementado
                'co_cidadao': co_cidadao,
                'nome_adolescente': nome_adolescente,
                'nome_responsavel': nome_responsavel_atual,
                'tipo_abordagem': proxima_acao_payload.get('tipo_abordagem'),
                'resultado_abordagem': None, # Próxima ação não tem resultado ainda
                'observacoes': proxima_acao_payload.get('observacoes'), # Usar observação enviada do frontend
                'data_acao': proxima_acao_payload.get('data_acao'), # Data futura
                'responsavel_pela_acao': proxima_acao_payload.get('responsavel_pela_acao'),
                'metodo_desejado': proxima_acao_payload.get('metodo_desejado')
            })

        conn.commit()
        return jsonify({"sucesso": True, "mensagem": "Ação(ões) registrada(s) com sucesso!"})
    except Exception as e:
        if conn: conn.rollback()
        print(f"Erro ao registrar ação do adolescente: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {str(e)}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/graficos_painel_adolescentes')
def api_graficos_painel_adolescentes():
    equipe_req = request.args.get('equipe', 'Todas')
    # agente_selecionado_req = request.args.get('agente_selecionado', 'Todas') # Não usado para estes gráficos

    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        # --- Dados para Gráfico de Pizza (Distribuição da Equipe) ---
        # Contagem total por status para a equipe selecionada (ou todas)
        
        base_where_pizza = ["m.idade_calculada BETWEEN 14 AND 18"]
        params_pizza = {}
        if equipe_req != 'Todas':
            base_where_pizza.append("m.nome_equipe = %(equipe)s")
            params_pizza['equipe'] = equipe_req
        
        where_clause_pizza_str = " WHERE " + " AND ".join(base_where_pizza) if base_where_pizza else ""

        query_pizza_status = f"""
            SELECT 
                SUM(CASE WHEN m.status_gravidez = 'Grávida' THEN 1 ELSE 0 END) as gestantes,
                SUM(CASE WHEN (m.metodo IS NULL OR m.metodo = '') AND (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') THEN 1 ELSE 0 END) as sem_metodo,
                SUM(CASE WHEN 
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days'))
                            ))
                        )
                    THEN 1 ELSE 0 END) as metodo_atraso,
                SUM(CASE WHEN
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                            )) OR
                            (m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%')
                        )
                    THEN 1 ELSE 0 END) as metodo_em_dia
            FROM sistemaaps.mv_plafam m
            {where_clause_pizza_str};
        """
        cur.execute(query_pizza_status, params_pizza)
        dados_pizza = cur.fetchone()

        # --- Dados para Gráfico de Barras (Distribuição por Micro-área) ---
        # Contagem por status AGRUPADO POR MICROÁREA para a equipe selecionada (ou todas)
        
        dados_barras_db = []

        if equipe_req != 'Todas':
            # Query para agrupar por AGENTE (via microarea) DENTRO DA EQUIPE SELECIONADA
            base_where_barras_agente = ["m.idade_calculada BETWEEN 14 AND 18", "m.nome_equipe = %(equipe)s"]
            params_barras_agente = {'equipe': equipe_req}
            where_clause_barras_agente_str = " WHERE " + " AND ".join(base_where_barras_agente)

            query_barras_status_por_agente = f"""
            SELECT 
                m.microarea, ag.nome_agente,
                SUM(CASE WHEN m.status_gravidez = 'Grávida' THEN 1 ELSE 0 END) as gestantes,
                SUM(CASE WHEN (m.metodo IS NULL OR m.metodo = '') AND (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') THEN 1 ELSE 0 END) as sem_metodo,
                SUM(CASE WHEN 
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days'))
                            ))
                        )
                    THEN 1 ELSE 0 END) as metodo_atraso,
                SUM(CASE WHEN
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                            )) OR
                            (m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%')
                        )
                    THEN 1 ELSE 0 END) as metodo_em_dia
            FROM sistemaaps.mv_plafam m
            LEFT JOIN sistemaaps.tb_agentes ag ON m.nome_equipe = ag.nome_equipe AND m.microarea = ag.micro_area
            {where_clause_barras_agente_str}
            GROUP BY m.microarea, ag.nome_agente
            ORDER BY m.microarea, ag.nome_agente;
            """
            cur.execute(query_barras_status_por_agente, params_barras_agente)
            dados_barras_db = cur.fetchall()
        else: # "Todas as Equipes" selecionado
            # Query para agrupar por EQUIPE
            query_barras_status_por_equipe = f"""
            SELECT 
                m.nome_equipe,
                SUM(CASE WHEN m.status_gravidez = 'Grávida' THEN 1 ELSE 0 END) as gestantes,
                SUM(CASE WHEN (m.metodo IS NULL OR m.metodo = '') AND (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') THEN 1 ELSE 0 END) as sem_metodo,
                SUM(CASE WHEN 
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days'))
                            ))
                        )
                    THEN 1 ELSE 0 END) as metodo_atraso,
                SUM(CASE WHEN
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                            )) OR
                            (m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%')
                        )
                    THEN 1 ELSE 0 END) as metodo_em_dia
            FROM sistemaaps.mv_plafam m
            WHERE m.idade_calculada BETWEEN 14 AND 18 AND m.nome_equipe IS NOT NULL
            GROUP BY m.nome_equipe
            ORDER BY m.nome_equipe;
            """
            cur.execute(query_barras_status_por_equipe)
            dados_barras_db = cur.fetchall()
        
        dados_barras = []
        if dados_barras_db:
            for row in dados_barras_db:
                dados_barras.append(dict(row))

        return jsonify({
            "pizza_data": dict(dados_pizza) if dados_pizza else {},
            "bar_chart_data": dados_barras
        })

    except Exception as e:
        print(f"Erro ao buscar dados para gráficos de adolescentes: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()


# --- Funções e Rotas para o Painel Hiperdia HAS ---

@app.route('/api/equipes_microareas_hiperdia')
def api_equipes_microareas_hiperdia():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        # Etapa 1: Buscar equipes, suas microáreas distintas e contagem de pacientes
        cur.execute("""
            SELECT 
                nome_equipe, 
                array_agg(DISTINCT microarea ORDER BY microarea ASC) as microareas_list,
                COUNT(DISTINCT cod_paciente) as total_pacientes_hipertensos
            FROM sistemaaps.mv_hiperdia_hipertensao
            WHERE nome_equipe IS NOT NULL AND microarea IS NOT NULL
            GROUP BY nome_equipe
            ORDER BY total_pacientes_hipertensos DESC, nome_equipe ASC;
        """)
        equipes_com_microareas = cur.fetchall()

        resultado_final = []
        for equipe_row in equipes_com_microareas:
            equipe_nome = equipe_row["nome_equipe"]
            
            # Etapa 2: Para cada equipe, buscar os agentes associados às suas microáreas
            # Usamos as microáreas da mv_hiperdia_hipertensao para garantir que só listamos agentes
            # de microáreas que de fato têm pacientes hipertensos naquela equipe.
            cur.execute("""
                SELECT DISTINCT ta.micro_area, ta.nome_agente
                FROM sistemaaps.tb_agentes ta
                WHERE ta.nome_equipe = %s AND ta.micro_area IN (
                    SELECT DISTINCT mh.microarea 
                    FROM sistemaaps.mv_hiperdia_hipertensao mh 
                    WHERE mh.nome_equipe = %s AND mh.microarea IS NOT NULL
                )
                ORDER BY ta.micro_area, ta.nome_agente;
            """, (equipe_nome, equipe_nome))
            agentes_db = cur.fetchall()
            agentes_formatados = [{"micro_area": ag["micro_area"], "nome_agente": ag["nome_agente"]} for ag in agentes_db]

            resultado_final.append({
                "nome_equipe": equipe_nome,
                "agentes": agentes_formatados, # Similar ao painel adolescentes
                "num_pacientes": equipe_row["total_pacientes_hipertensos"]
            })
        return jsonify(resultado_final)
    except Exception as e:
        print(f"Erro ao buscar equipes e microáreas para Hiperdia HAS: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()


def build_hiperdia_has_filters(args):
    equipe = args.get('equipe', 'Todas')
    microarea_selecionada = args.get('microarea', 'Todas')  # 'Todas' ou 'Todas as áreas' ou "Área X - Agente Y"
    search_term = args.get('search', None)
    sort_by = args.get('sort_by', 'proxima_acao_asc')  # Ordenação padrão
    status_filter = args.get('status', 'Todos')  # Mantido para compatibilidade

    query_params = {}
    where_clauses = []

    # Filtro por equipe
    if equipe != 'Todas':
        where_clauses.append("m.nome_equipe = %(equipe)s")
        query_params['equipe'] = equipe

    # Filtro por microárea (aceita formatos "Área X - Agente Y" ou "Área X")
    if microarea_selecionada not in (None, '', 'Todas', 'Todas as áreas'):
        if ' - ' in microarea_selecionada:
            micro_area_str = microarea_selecionada.split(' - ', 1)[0].replace('Área ', '').strip()
        else:
            micro_area_str = microarea_selecionada.replace('Área ', '').strip()
        if micro_area_str:
            where_clauses.append("m.microarea = %(microarea)s")
            query_params['microarea'] = micro_area_str

    # Busca por nome
    if search_term:
        where_clauses.append("UPPER(m.nome_paciente) LIKE UPPER(%(search)s)")
        query_params['search'] = f"%{search_term}%"

    # Filtro por status MRPA (Controlado/Descompensado)
    if status_filter == 'Controlado':
        where_clauses.append("EXISTS (SELECT 1 FROM sistemaaps.tb_hiperdia_mrpa mrpa JOIN sistemaaps.tb_hiperdia_has_acompanhamento ac ON mrpa.cod_acompanhamento = ac.cod_acompanhamento WHERE ac.cod_cidadao = m.cod_paciente AND mrpa.status_mrpa = 1)")
    elif status_filter == 'Descompensado':
        where_clauses.append("EXISTS (SELECT 1 FROM sistemaaps.tb_hiperdia_mrpa mrpa JOIN sistemaaps.tb_hiperdia_has_acompanhamento ac ON mrpa.cod_acompanhamento = ac.cod_acompanhamento WHERE ac.cod_cidadao = m.cod_paciente AND mrpa.status_mrpa = 0)")

    # Mapeamento de ordenação
    sort_mapping = {
        'nome_asc': 'm.nome_paciente ASC',
        'nome_desc': 'm.nome_paciente DESC',
        'idade_asc': 'm.idade_calculada ASC',
        'idade_desc': 'm.idade_calculada DESC',
        'proxima_acao_asc': 'data_proxima_acao_ordenacao ASC NULLS LAST, m.nome_paciente ASC',
        'proxima_acao_desc': 'data_proxima_acao_ordenacao DESC NULLS FIRST, m.nome_paciente DESC',
    }
    order_by_clause = " ORDER BY " + sort_mapping.get(
        sort_by, 'data_proxima_acao_ordenacao ASC NULLS LAST, m.nome_paciente ASC'
    )

    return where_clauses, order_by_clause, query_params, status_filter

@app.route('/api/pacientes_hiperdia_has')
def api_pacientes_hiperdia_has():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 10)) # Padrão de 10, mas pode ser alterado pelo frontend
        offset = (page - 1) * limit

        where_clauses, order_by_clause, query_params, status_filter = build_hiperdia_has_filters(request.args)
        
        fields = """
            m.cod_paciente, m.nome_paciente, m.cartao_sus, m.idade_calculada, m.nome_equipe, m.microarea,
            m.dt_nascimento, m.situacao_problema, ag.nome_agente,
            pa_futura.data_agendamento AS data_proxima_acao_ordenacao,
            pa_futura.cod_acao AS tipo_proxima_acao_ordenacao,
            tratamento.tratamento_atual,
            ultima_mrpa.mrpa_atual,
            exames.exames_atuais,
            ultimo_status_mrpa.status_mrpa,
            monitoramento.tem_monitoramento_ativo
        """
        from_join_clause = """
        FROM sistemaaps.mv_hiperdia_hipertensao m
        LEFT JOIN sistemaaps.tb_agentes ag ON m.nome_equipe = ag.nome_equipe AND m.microarea = ag.micro_area
        LEFT JOIN LATERAL (
            SELECT STRING_AGG(
                '<b>' || med.nome_medicamento || '</b> (' || 
                CASE 
                    WHEN med.dose IS NOT NULL AND med.frequencia IS NOT NULL 
                    THEN med.dose || ' comprimido(s) - ' || med.frequencia || 'x ao dia'
                    WHEN med.frequencia IS NOT NULL 
                    THEN med.frequencia || 'x ao dia'
                    ELSE 'Conforme prescrição'
                END || ')' || 
                CASE 
                    WHEN med.data_inicio IS NOT NULL 
                    THEN ' - ' || TO_CHAR(med.data_inicio, 'DD/MM/YYYY')
                    ELSE ''
                END,
                '<br>'
                ORDER BY 
                    CASE WHEN med.data_inicio IS NOT NULL THEN med.data_inicio ELSE '1900-01-01'::date END DESC, 
                    med.nome_medicamento
            ) AS tratamento_atual
            FROM sistemaaps.tb_hiperdia_has_medicamentos med
            WHERE med.codcidadao = m.cod_paciente 
              AND (med.data_fim IS NULL OR med.data_fim > CURRENT_DATE)
        ) tratamento ON TRUE
        LEFT JOIN LATERAL (
            SELECT mrpa.media_pa_sistolica || 'x' || mrpa.media_pa_diastolica || 'mmHg' AS mrpa_atual
            FROM sistemaaps.tb_hiperdia_mrpa mrpa
            JOIN sistemaaps.tb_hiperdia_has_acompanhamento ac ON mrpa.cod_acompanhamento = ac.cod_acompanhamento
            WHERE ac.cod_cidadao = m.cod_paciente
            ORDER BY mrpa.data_mrpa DESC
            LIMIT 1
        ) ultima_mrpa ON TRUE
        LEFT JOIN LATERAL (
            SELECT mrpa.status_mrpa
            FROM sistemaaps.tb_hiperdia_mrpa mrpa
            JOIN sistemaaps.tb_hiperdia_has_acompanhamento ac ON mrpa.cod_acompanhamento = ac.cod_acompanhamento
            WHERE ac.cod_cidadao = m.cod_paciente
            ORDER BY mrpa.data_mrpa DESC
            LIMIT 1
        ) ultimo_status_mrpa ON TRUE
        LEFT JOIN LATERAL (
            SELECT CASE WHEN COUNT(*) > 0 THEN true ELSE false END AS tem_monitoramento_ativo
            FROM sistemaaps.tb_hiperdia_has_acompanhamento ac
            WHERE ac.cod_cidadao = m.cod_paciente
              AND ac.status_acao = 'PENDENTE'
              AND ac.cod_acao IN (1, 2, 3, 4)
        ) monitoramento ON TRUE
        LEFT JOIN LATERAL (
            SELECT STRING_AGG(ex.nome_procedimento || ': ' || ex.resultado_exame, '<br>' ORDER BY ex.data_exame DESC) AS exames_atuais
            FROM (SELECT *, ROW_NUMBER() OVER(PARTITION BY co_proced ORDER BY data_exame DESC) as rn FROM sistemaaps.mv_hiperdia_exames WHERE co_seq_cidadao = m.cod_paciente AND co_proced IN (4500, 4507, 199)) ex
            WHERE ex.rn = 1
        ) exames ON TRUE
        LEFT JOIN LATERAL (
            SELECT data_agendamento, cod_acao  
            FROM sistemaaps.tb_hiperdia_has_acompanhamento 
            WHERE cod_cidadao = m.cod_paciente 
              AND status_acao = 'PENDENTE'
            ORDER BY data_agendamento ASC  
            LIMIT 1
        ) pa_futura ON TRUE
        """

        # Adiciona o filtro de status aqui, onde pa_futura está disponível
        if status_filter == 'AcoesPendentes':
            where_clauses.append("pa_futura.data_agendamento IS NOT NULL")
        elif status_filter == 'ComTratamento':
            where_clauses.append("tratamento.tratamento_atual IS NOT NULL AND TRIM(tratamento.tratamento_atual) != ''")

        where_clause_str = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""

        # Executa a contagem primeiro, com todos os filtros aplicados
        count_query = f"SELECT COUNT(*) {from_join_clause} {where_clause_str}"
        cur.execute(count_query, query_params)
        total_pacientes = cur.fetchone()[0] or 0

        # Agora, executa a busca paginada
        query_params_paginated = query_params.copy()
        query_params_paginated['limit'] = limit
        query_params_paginated['offset'] = offset
        final_query = f"SELECT {fields} {from_join_clause} {where_clause_str} {order_by_clause} LIMIT %(limit)s OFFSET %(offset)s"
        cur.execute(final_query, query_params_paginated)
        
        pacientes = [dict(row) for row in cur.fetchall()]

        for p in pacientes: # Formatar datas se necessário
            if p.get('dt_nascimento') and isinstance(p['dt_nascimento'], date):
                p['dt_nascimento'] = p['dt_nascimento'].strftime('%d/%m/%Y')

            # Formatar próxima ação para o frontend
            if p.get('data_proxima_acao_ordenacao') and isinstance(p['data_proxima_acao_ordenacao'], date):
                p['proxima_acao_data_formatada'] = p['data_proxima_acao_ordenacao'].strftime('%d/%m/%Y')
                p['proxima_acao_descricao'] = TIPO_ACAO_MAP_PY.get(p['tipo_proxima_acao_ordenacao'], 'Ação futura')
            else:
                p['proxima_acao_data_formatada'] = None
                p['proxima_acao_descricao'] = None


        return jsonify({
            'pacientes': pacientes,
            'total': total_pacientes,
            'page': page,
            'limit': limit,
            'pages': (total_pacientes + limit - 1) // limit if total_pacientes > 0 else 0
        })
    except Exception as e:
        print(f"Erro na API /api/pacientes_hiperdia_has: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()


@app.route('/api/get_total_hipertensos')
def api_get_total_hipertensos():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        where_clauses, _, query_params, status_filter = build_hiperdia_has_filters(request.args)

        # Para a contagem, precisamos do JOIN LATERAL se o filtro for 'AcoesPendentes' ou 'ComTratamento'
        from_join_clause = """
        FROM sistemaaps.mv_hiperdia_hipertensao m
        LEFT JOIN LATERAL (
            SELECT data_agendamento
            FROM sistemaaps.tb_hiperdia_has_acompanhamento
            WHERE cod_cidadao = m.cod_paciente
              AND status_acao = 'PENDENTE'
              AND data_agendamento >= CURRENT_DATE
            ORDER BY data_agendamento ASC
            LIMIT 1
        ) pa_futura ON TRUE
        LEFT JOIN LATERAL (
            SELECT STRING_AGG(
                '<b>' || med.nome_medicamento || '</b> (' || 
                CASE 
                    WHEN med.dose IS NOT NULL AND med.frequencia IS NOT NULL 
                    THEN med.dose || ' comprimido(s) - ' || med.frequencia || 'x ao dia'
                    WHEN med.frequencia IS NOT NULL 
                    THEN med.frequencia || 'x ao dia'
                    ELSE 'Conforme prescrição'
                END || ')' || 
                CASE 
                    WHEN med.data_inicio IS NOT NULL 
                    THEN ' - ' || TO_CHAR(med.data_inicio, 'DD/MM/YYYY')
                    ELSE ''
                END,
                '<br>'
                ORDER BY 
                    CASE WHEN med.data_inicio IS NOT NULL THEN med.data_inicio ELSE '1900-01-01'::date END DESC, 
                    med.nome_medicamento
            ) AS tratamento_atual
            FROM sistemaaps.tb_hiperdia_has_medicamentos med
            WHERE med.codcidadao = m.cod_paciente 
              AND (med.data_fim IS NULL OR med.data_fim > CURRENT_DATE)
        ) tratamento ON TRUE
        """
        if status_filter == 'AcoesPendentes':
            where_clauses.append("pa_futura.data_agendamento IS NOT NULL")
        elif status_filter == 'ComTratamento':
            where_clauses.append("tratamento.tratamento_atual IS NOT NULL AND TRIM(tratamento.tratamento_atual) != ''")
        where_clause = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""
        
        count_query = f"SELECT COUNT(*) {from_join_clause} {where_clause}"
        cur.execute(count_query, query_params)
        total_pacientes = cur.fetchone()[0] or 0

        return jsonify({'total_pacientes': total_pacientes})

    except Exception as e:
        print(f"Erro na API /api/get_total_hipertensos: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/get_hipertensos_mrpa_pendente')
def api_get_hipertensos_mrpa_pendente():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        where_clauses, _, query_params, status_filter = build_hiperdia_has_filters(request.args)
        # Condição para QUALQUER ação pendente
        pendente_condition = "EXISTS (SELECT 1 FROM sistemaaps.tb_hiperdia_has_acompanhamento ha WHERE ha.cod_cidadao = m.cod_paciente AND ha.status_acao = 'PENDENTE')"
        
        # Adiciona a condição de ação pendente à lista de condições existentes
        where_clauses.append(pendente_condition)

        # Constrói a string final da cláusula WHERE
        final_where_clause = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""
        
        base_query = "FROM sistemaaps.mv_hiperdia_hipertensao m"
        # Conta pacientes distintos para não contar o mesmo paciente várias vezes se ele tiver múltiplas ações pendentes
        count_query = "SELECT COUNT(DISTINCT m.cod_paciente) " + base_query + final_where_clause
        cur.execute(count_query, query_params)
        total_pacientes = cur.fetchone()[0] or 0

        return jsonify({'total_pacientes': total_pacientes})

    except Exception as e:
        app.logger.exception(f"Erro na API /api/get_hipertensos_acoes_pendentes: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()


@app.route('/api/get_hipertensos_controlados')
def api_get_hipertensos_controlados():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        where_clauses, _, query_params, status_filter = build_hiperdia_has_filters(request.args)
        # Condição para hipertensos controlados (status_mrpa = 1)
        controlados_condition = "EXISTS (SELECT 1 FROM sistemaaps.tb_hiperdia_mrpa mrpa JOIN sistemaaps.tb_hiperdia_has_acompanhamento ac ON mrpa.cod_acompanhamento = ac.cod_acompanhamento WHERE ac.cod_cidadao = m.cod_paciente AND mrpa.status_mrpa = 1)"

        # Adiciona a condição de controlados à lista de condições existentes
        where_clauses.append(controlados_condition)

        # Constrói a string final da cláusula WHERE
        final_where_clause = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""

        base_query = "FROM sistemaaps.mv_hiperdia_hipertensao m"
        # Conta pacientes distintos para não contar o mesmo paciente várias vezes
        count_query = "SELECT COUNT(DISTINCT m.cod_paciente) " + base_query + final_where_clause
        cur.execute(count_query, query_params)
        total_pacientes = cur.fetchone()[0] or 0

        return jsonify({'total_pacientes': total_pacientes})

    except Exception as e:
        app.logger.exception(f"Erro na API /api/get_hipertensos_controlados: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()


@app.route('/api/get_hipertensos_descompensados')
def api_get_hipertensos_descompensados():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        where_clauses, _, query_params, status_filter = build_hiperdia_has_filters(request.args)
        # Condição para hipertensos descompensados (status_mrpa = 0)
        descompensados_condition = "EXISTS (SELECT 1 FROM sistemaaps.tb_hiperdia_mrpa mrpa JOIN sistemaaps.tb_hiperdia_has_acompanhamento ac ON mrpa.cod_acompanhamento = ac.cod_acompanhamento WHERE ac.cod_cidadao = m.cod_paciente AND mrpa.status_mrpa = 0)"

        # Adiciona a condição de descompensados à lista de condições existentes
        where_clauses.append(descompensados_condition)

        # Constrói a string final da cláusula WHERE
        final_where_clause = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""

        base_query = "FROM sistemaaps.mv_hiperdia_hipertensao m"
        # Conta pacientes distintos para não contar o mesmo paciente várias vezes
        count_query = "SELECT COUNT(DISTINCT m.cod_paciente) " + base_query + final_where_clause
        cur.execute(count_query, query_params)
        total_pacientes = cur.fetchone()[0] or 0

        return jsonify({'total_pacientes': total_pacientes})

    except Exception as e:
        app.logger.exception(f"Erro na API /api/get_hipertensos_descompensados: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()


@app.route('/api/get_hipertensos_em_avaliacao')
def api_get_hipertensos_em_avaliacao():
    """
    Conta hipertensos em avaliação (ações pendentes: Iniciar MRPA (1), Avaliar MRPA (2), Solicitar Exames (4), Avaliar Exames (5))
    Um paciente pode estar controlado/descompensado E ter ações pendentes, então pode ser contado em múltiplos cards
    """
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        where_clauses, _, query_params, status_filter = build_hiperdia_has_filters(request.args)
        # Condição para hipertensos em avaliação (ações pendentes 1, 2, 4, 5)
        em_avaliacao_condition = """EXISTS (
            SELECT 1 FROM sistemaaps.tb_hiperdia_has_acompanhamento ha
            WHERE ha.cod_cidadao = m.cod_paciente
            AND ha.status_acao = 'PENDENTE'
            AND ha.cod_acao IN (1, 2, 4, 5)
        )"""

        where_clauses.append(em_avaliacao_condition)

        final_where_clause = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""
        base_query = "FROM sistemaaps.mv_hiperdia_hipertensao m"
        count_query = "SELECT COUNT(DISTINCT m.cod_paciente) " + base_query + final_where_clause
        cur.execute(count_query, query_params)
        total_pacientes = cur.fetchone()[0] or 0

        return jsonify({'total_pacientes': total_pacientes})

    except Exception as e:
        app.logger.exception(f"Erro na API /api/get_hipertensos_em_avaliacao: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()


@app.route('/api/get_hipertensos_aguardando_exames')
def api_get_hipertensos_aguardando_exames():
    """
    Conta hipertensos aguardando exames (ações pendentes: Solicitar Exames (4) ou Avaliar Exames (5))
    Um paciente pode estar controlado/descompensado E aguardando exames
    """
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        where_clauses, _, query_params, status_filter = build_hiperdia_has_filters(request.args)
        # Condição para hipertensos aguardando exames (ações pendentes 4, 5)
        aguardando_exames_condition = """EXISTS (
            SELECT 1 FROM sistemaaps.tb_hiperdia_has_acompanhamento ha
            WHERE ha.cod_cidadao = m.cod_paciente
            AND ha.status_acao = 'PENDENTE'
            AND ha.cod_acao IN (4, 5)
        )"""

        where_clauses.append(aguardando_exames_condition)

        final_where_clause = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""
        base_query = "FROM sistemaaps.mv_hiperdia_hipertensao m"
        count_query = "SELECT COUNT(DISTINCT m.cod_paciente) " + base_query + final_where_clause
        cur.execute(count_query, query_params)
        total_pacientes = cur.fetchone()[0] or 0

        return jsonify({'total_pacientes': total_pacientes})

    except Exception as e:
        app.logger.exception(f"Erro na API /api/get_hipertensos_aguardando_exames: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()


@app.route('/api/get_hipertensos_sem_avaliacao')
def api_get_hipertensos_sem_avaliacao():
    """
    Conta hipertensos sem avaliação (sem nenhuma ação de MRPA registrada - nem controlado nem descompensado)
    """
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        where_clauses, _, query_params, status_filter = build_hiperdia_has_filters(request.args)
        # Condição para hipertensos sem avaliação (sem status_mrpa definido)
        sem_avaliacao_condition = """NOT EXISTS (
            SELECT 1 FROM sistemaaps.tb_hiperdia_mrpa mrpa
            JOIN sistemaaps.tb_hiperdia_has_acompanhamento ac ON mrpa.cod_acompanhamento = ac.cod_acompanhamento
            WHERE ac.cod_cidadao = m.cod_paciente
        )"""

        where_clauses.append(sem_avaliacao_condition)

        final_where_clause = " WHERE " + " AND ".join(where_clauses) if where_clauses else ""
        base_query = "FROM sistemaaps.mv_hiperdia_hipertensao m"
        count_query = "SELECT COUNT(DISTINCT m.cod_paciente) " + base_query + final_where_clause
        cur.execute(count_query, query_params)
        total_pacientes = cur.fetchone()[0] or 0

        return jsonify({'total_pacientes': total_pacientes})

    except Exception as e:
        app.logger.exception(f"Erro na API /api/get_hipertensos_sem_avaliacao: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()


@app.route('/api/hiperdia/timeline/<int:cod_cidadao>')
def api_hiperdia_timeline(cod_cidadao):
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        period_filter = request.args.get('period', 'all') # Novo: Captura o filtro de período

        # Query base para buscar todas as ações de um paciente
        base_sql_query = """
            SELECT 
                ac.cod_acompanhamento, ac.cod_cidadao, ac.cod_acao, ta.dsc_acao,
                ac.status_acao,
                ac.responsavel_pela_acao,
                ac.data_realizacao,
                ac.data_agendamento,
                ac.observacoes,
                re.colesterol_total, re.hdl, re.ldl, re.triglicerideos, re.glicemia_jejum, 
                re.hemoglobina_glicada, re.ureia, re.creatinina, re.sodio, re.potassio, re.acido_urico,
                tr.tipo_ajuste, tr.medicamentos_novos,
                rcv.idade, rcv.sexo, rcv.tabagismo, rcv.diabetes, rcv.colesterol_total as rcv_colesterol, rcv.pressao_sistolica,
                mrpa.media_pa_sistolica, mrpa.media_pa_diastolica, mrpa.analise_mrpa, mrpa.status_mrpa,
                nu.peso, nu.imc, nu.circunferencia_abdominal, nu.orientacoes_nutricionais,
                mrg.data_mrg, mrg.g_jejum, mrg.g_apos_cafe, mrg.g_antes_almoco, mrg.g_apos_almoco, 
                mrg.g_antes_jantar, mrg.g_ao_deitar, mrg.analise_mrg,
                card.consulta_cardiologia, card.recomendacoes_cardiologia, 
                card.profissional_responsavel as card_profissional, card.tipo_consulta,
                next_ac.cod_acao AS next_action_cod,
                next_ac.data_agendamento AS next_action_date
            FROM 
                sistemaaps.tb_hiperdia_has_acompanhamento ac
            JOIN 
                sistemaaps.tb_hiperdia_tipos_acao ta ON ac.cod_acao = ta.cod_acao
            LEFT JOIN
                sistemaaps.tb_hiperdia_resultados_exames re ON ac.cod_acompanhamento = re.cod_acompanhamento
            LEFT JOIN
                sistemaaps.tb_hiperdia_tratamento tr ON ac.cod_acompanhamento = tr.cod_acompanhamento
            LEFT JOIN
                sistemaaps.tb_hiperdia_risco_cv rcv ON ac.cod_acompanhamento = rcv.cod_acompanhamento
            LEFT JOIN
                sistemaaps.tb_hiperdia_nutricao nu ON ac.cod_acompanhamento = nu.cod_acompanhamento
            LEFT JOIN
                sistemaaps.tb_hiperdia_mrpa mrpa ON ac.cod_acompanhamento = mrpa.cod_acompanhamento
            LEFT JOIN
                sistemaaps.tb_hiperdia_mrg mrg ON ac.cod_acompanhamento = mrg.cod_acompanhamento
            LEFT JOIN
                sistemaaps.tb_hiperdia_has_cardiologia card ON ac.cod_acompanhamento = card.cod_acompanhamento
            LEFT JOIN LATERAL (
                SELECT ha_next.cod_acao, ha_next.data_agendamento
                FROM sistemaaps.tb_hiperdia_has_acompanhamento ha_next
                WHERE ha_next.cod_acao_origem = ac.cod_acompanhamento
                  AND ha_next.status_acao = 'PENDENTE'
                  AND ha_next.data_agendamento >= CURRENT_DATE
                ORDER BY ha_next.data_agendamento ASC
                LIMIT 1
            ) next_ac ON TRUE
        """
        
        where_clauses = ["ac.cod_cidadao = %(cod_cidadao)s"]
        query_params = {'cod_cidadao': cod_cidadao}

        # Adiciona o filtro de período à cláusula WHERE
        if period_filter != 'all':
            try:
                days = int(period_filter.replace('d', ''))
                where_clauses.append(f"COALESCE(ac.data_realizacao, ac.data_agendamento) >= CURRENT_DATE - INTERVAL '{days} days'")
            except ValueError:
                # Se o filtro for inválido, ignora
                pass

        where_clause_str = " WHERE " + " AND ".join(where_clauses)
        order_by_clause = """
            ORDER BY
                COALESCE(ac.data_realizacao, ac.data_agendamento) DESC, ac.cod_acompanhamento DESC;
        """

        final_query = base_sql_query + where_clause_str + order_by_clause
        cur.execute(final_query, query_params)

        eventos = []
        for row in cur.fetchall():
            evento = {
                'cod_acompanhamento': row['cod_acompanhamento'],
                'cod_cidadao': row['cod_cidadao'],
                'cod_acao': row['cod_acao'],
                'dsc_acao': TIPO_ACAO_MAP_PY.get(row['cod_acao'], row['dsc_acao']),
                'status_acao': row['status_acao'],
                'responsavel_pela_acao': row['responsavel_pela_acao'], # Adicionado aqui
                'data_realizacao': row['data_realizacao'].strftime('%Y-%m-%d') if row['data_realizacao'] else None, # Adicionado aqui
                'data_agendamento': row['data_agendamento'].strftime('%Y-%m-%d') if row['data_agendamento'] else None,
                'observacoes': row['observacoes'],
            }
            # Se for uma avaliação de exames e houver resultados, agrupa-os
            # Checa se pelo menos um dos campos de exame foi preenchido
            if row['cod_acao'] == 5 and any([
                row['colesterol_total'] is not None,
                row['hdl'] is not None,
                row['ldl'] is not None,
                row['triglicerideos'] is not None,
                row['glicemia_jejum'] is not None,
                row['hemoglobina_glicada'] is not None,
                row['ureia'] is not None,
                row['creatinina'] is not None,
                row['sodio'] is not None,
                row['potassio'] is not None,
                row['acido_urico'] is not None
            ]):
                evento['resultados_exames'] = {
                    'colesterol_total': row['colesterol_total'], 'hdl': row['hdl'], 'ldl': row['ldl'],
                    'triglicerideos': row['triglicerideos'], 'glicemia_jejum': row['glicemia_jejum'],
                    'hemoglobina_glicada': safe_float_conversion(row['hemoglobina_glicada']),
                    'ureia': row['ureia'], 'creatinina': safe_float_conversion(row['creatinina']),
                    'sodio': row['sodio'], 'potassio': safe_float_conversion(row['potassio']),
                    'acido_urico': safe_float_conversion(row['acido_urico']),
                }
            # Se for uma modificação de tratamento e houver dados, agrupa-os
            if row['cod_acao'] == 3 and row['tipo_ajuste'] is not None:
                evento['treatment_modification'] = {
                    'tipo_ajuste': row['tipo_ajuste'],
                    'medicamentos_novos': row['medicamentos_novos'],
                }
            # Se for uma avaliação de RCV e houver dados, agrupa-os
            if row['cod_acao'] == 6 and row['idade'] is not None:
                evento['risk_assessment_details'] = {
                    'idade': row['idade'],
                    'sexo': row['sexo'],
                    'tabagismo': row['tabagismo'],
                    'diabetes': row['diabetes'],
                    'colesterol_total': row['rcv_colesterol'],
                    'pressao_sistolica': row['pressao_sistolica'],
                }
            # Se for um registro de consulta de nutrição e houver dados, agrupa-os
            if row['cod_acao'] == 8 and row['peso'] is not None:
                evento['nutrition_details'] = {
                    'peso': safe_float_conversion(row['peso']),
                    'imc': safe_float_conversion(row['imc']),
                    'circunferencia_abdominal': row['circunferencia_abdominal'],
                    'orientacoes_nutricionais': row['orientacoes_nutricionais'],
                }
            # Se for uma avaliação de MRPA e houver dados, agrupa-os
            if row['cod_acao'] == 2 and row['media_pa_sistolica'] is not None:
                evento['mrpa_details'] = {
                    'media_pa_sistolica': row['media_pa_sistolica'],
                    'media_pa_diastolica': row['media_pa_diastolica'],
                    'analise_mrpa': row['analise_mrpa'],
                    'status_mrpa': row['status_mrpa']
                }
            
            # Se for uma avaliação de MRG e houver dados, agrupa-os
            if row['cod_acao'] == 11 and any([
                row['g_jejum'] is not None,
                row['g_apos_cafe'] is not None,
                row['g_antes_almoco'] is not None,
                row['g_apos_almoco'] is not None,
                row['g_antes_jantar'] is not None,
                row['g_ao_deitar'] is not None
            ]):
                evento['mrg_details'] = {
                    'data_mrg': row['data_mrg'].strftime('%Y-%m-%d') if row['data_mrg'] else None,
                    'g_jejum': safe_float_conversion(row['g_jejum']),
                    'g_apos_cafe': safe_float_conversion(row['g_apos_cafe']),
                    'g_antes_almoco': safe_float_conversion(row['g_antes_almoco']),
                    'g_apos_almoco': safe_float_conversion(row['g_apos_almoco']),
                    'g_antes_jantar': safe_float_conversion(row['g_antes_jantar']),
                    'g_ao_deitar': safe_float_conversion(row['g_ao_deitar']),
                    'analise_mrg': row['analise_mrg']
                }

            # Se for um registro de cardiologia e houver dados, agrupa-os
            if (row['cod_acao'] == 10 or row['cod_acao'] == 11) and row['consulta_cardiologia'] is not None:
                evento['cardiologia_details'] = {
                    'consulta_cardiologia': row['consulta_cardiologia'],
                    'recomendacoes_cardiologia': row['recomendacoes_cardiologia'],
                    'profissional_responsavel': row['card_profissional'],
                    'tipo_consulta': row['tipo_consulta']
                }

            eventos.append(evento)
        
        return jsonify(eventos)

    except Exception as e:
        print(f"Erro na API /api/hiperdia/timeline: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/cancelar_acao/<int:cod_acompanhamento>', methods=['POST'])
def api_cancelar_acao_hiperdia(cod_acompanhamento):
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Primeiro, buscar informações da ação para identificar referências posteriores
        cur.execute("""
            SELECT cod_cidadao, cod_acao, status_acao, cod_acao_origem
            FROM sistemaaps.tb_hiperdia_has_acompanhamento
            WHERE cod_acompanhamento = %(cod_acompanhamento)s
        """, {'cod_acompanhamento': cod_acompanhamento})
        
        acao_info = cur.fetchone()
        if not acao_info:
            conn.rollback()
            return jsonify({"sucesso": False, "erro": "Ação não encontrada."}), 404

        cod_cidadao, cod_acao, status_acao, cod_acao_origem = acao_info

        # Se a ação não estiver pendente, não permitir cancelamento
        if status_acao != 'PENDENTE':
            conn.rollback()
            return jsonify({"sucesso": False, "erro": "Apenas ações pendentes podem ser canceladas."}), 400

        # Buscar todas as ações posteriores que referenciam esta ação
        cur.execute("""
            SELECT cod_acompanhamento, cod_acao, status_acao
            FROM sistemaaps.tb_hiperdia_has_acompanhamento
            WHERE cod_acao_origem = %(cod_acompanhamento)s
            ORDER BY cod_acompanhamento ASC
        """, {'cod_acompanhamento': cod_acompanhamento})
        
        acoes_posteriores = cur.fetchall()

        # Iniciar transação para remoção em cascata
        try:
            # 1. Remover dados relacionados à ação principal (se existirem)
            # Remover dados de MRPA
            cur.execute("""
                DELETE FROM sistemaaps.tb_hiperdia_mrpa
                WHERE cod_acompanhamento = %(cod_acompanhamento)s
            """, {'cod_acompanhamento': cod_acompanhamento})

            # Remover dados de tratamento
            cur.execute("""
                DELETE FROM sistemaaps.tb_hiperdia_tratamento
                WHERE cod_acompanhamento = %(cod_acompanhamento)s
            """, {'cod_acompanhamento': cod_acompanhamento})

            # Remover dados de resultados de exames
            cur.execute("""
                DELETE FROM sistemaaps.tb_hiperdia_resultados_exames
                WHERE cod_acompanhamento = %(cod_acompanhamento)s
            """, {'cod_acompanhamento': cod_acompanhamento})

            # Remover dados de risco cardiovascular
            cur.execute("""
                DELETE FROM sistemaaps.tb_hiperdia_risco_cv
                WHERE cod_acompanhamento = %(cod_acompanhamento)s
            """, {'cod_acompanhamento': cod_acompanhamento})

            # Remover dados de nutrição
            cur.execute("""
                DELETE FROM sistemaaps.tb_hiperdia_nutricao
                WHERE cod_acompanhamento = %(cod_acompanhamento)s
            """, {'cod_acompanhamento': cod_acompanhamento})

            # 2. Remover ações posteriores em cascata (recursivamente)
            for acao_posterior in acoes_posteriores:
                cod_acompanhamento_posterior = acao_posterior[0]
                
                # Remover dados relacionados à ação posterior
                cur.execute("""
                    DELETE FROM sistemaaps.tb_hiperdia_mrpa
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s
                """, {'cod_acompanhamento': cod_acompanhamento_posterior})

                cur.execute("""
                    DELETE FROM sistemaaps.tb_hiperdia_tratamento
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s
                """, {'cod_acompanhamento': cod_acompanhamento_posterior})

                cur.execute("""
                    DELETE FROM sistemaaps.tb_hiperdia_resultados_exames
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s
                """, {'cod_acompanhamento': cod_acompanhamento_posterior})

                cur.execute("""
                    DELETE FROM sistemaaps.tb_hiperdia_risco_cv
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s
                """, {'cod_acompanhamento': cod_acompanhamento_posterior})

                cur.execute("""
                    DELETE FROM sistemaaps.tb_hiperdia_nutricao
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s
                """, {'cod_acompanhamento': cod_acompanhamento_posterior})

                # Remover a ação posterior
                cur.execute("""
                    DELETE FROM sistemaaps.tb_hiperdia_has_acompanhamento
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s
                """, {'cod_acompanhamento': cod_acompanhamento_posterior})

            # 3. Finalmente, remover a ação principal
            cur.execute("""
                DELETE FROM sistemaaps.tb_hiperdia_has_acompanhamento
                WHERE cod_acompanhamento = %(cod_acompanhamento)s
            """, {'cod_acompanhamento': cod_acompanhamento})

            # Verificar se alguma linha foi realmente removida
            if cur.rowcount == 0:
                conn.rollback()
                return jsonify({"sucesso": False, "erro": "Ação não encontrada para remoção."}), 404

            conn.commit()
            
            # Contar quantas ações foram removidas
            total_removidas = 1 + len(acoes_posteriores)  # ação principal + ações posteriores
            
            mensagem = f"Ação e {len(acoes_posteriores)} ação(ões) posterior(es) removida(s) com sucesso!"
            if len(acoes_posteriores) == 0:
                mensagem = "Ação removida com sucesso!"
            
            return jsonify({
                "sucesso": True, 
                "mensagem": mensagem,
                "acoes_removidas": total_removidas
            })

        except Exception as e:
            conn.rollback()
            print(f"Erro durante a remoção em cascata: {e}")
            return jsonify({"sucesso": False, "erro": f"Erro durante a remoção: {str(e)}"}), 500

    except Exception as e:
        if conn: conn.rollback()
        print(f"Erro ao cancelar ação do Hiperdia: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {str(e)}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/atualizar_status_acao/<int:cod_acompanhamento>', methods=['POST'])
def api_atualizar_status_acao_hiperdia(cod_acompanhamento):
    """Endpoint para atualizar o status de uma ação (REALIZADA ou CANCELADA)"""
    conn = None
    cur = None
    try:
        data = request.get_json()
        if not data:
            return jsonify({"sucesso": False, "erro": "Dados não fornecidos"}), 400

        novo_status = data.get('status_acao')
        if novo_status not in ['REALIZADA', 'CANCELADA']:
            return jsonify({"sucesso": False, "erro": "Status inválido. Deve ser 'REALIZADA' ou 'CANCELADA'"}), 400

        conn = get_db_connection()
        cur = conn.cursor()

        # Primeiro, verificar se a ação existe e está pendente
        cur.execute("""
            SELECT cod_cidadao, cod_acao, status_acao
            FROM sistemaaps.tb_hiperdia_has_acompanhamento
            WHERE cod_acompanhamento = %(cod_acompanhamento)s
        """, {'cod_acompanhamento': cod_acompanhamento})
        
        acao_info = cur.fetchone()
        if not acao_info:
            return jsonify({"sucesso": False, "erro": "Ação não encontrada."}), 404

        cod_cidadao, cod_acao, status_atual = acao_info

        if status_atual != 'PENDENTE':
            return jsonify({"sucesso": False, "erro": "Apenas ações pendentes podem ter o status alterado."}), 400

        # Atualizar o status da ação
        data_realizacao = 'CURRENT_DATE' if novo_status == 'REALIZADA' else None
        
        if novo_status == 'REALIZADA':
            cur.execute("""
                UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                SET status_acao = %(status_acao)s,
                    data_realizacao = CURRENT_DATE
                WHERE cod_acompanhamento = %(cod_acompanhamento)s
            """, {
                'status_acao': novo_status,
                'cod_acompanhamento': cod_acompanhamento
            })
        else:  # CANCELADA
            cur.execute("""
                UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                SET status_acao = %(status_acao)s,
                    data_realizacao = NULL
                WHERE cod_acompanhamento = %(cod_acompanhamento)s
            """, {
                'status_acao': novo_status,
                'cod_acompanhamento': cod_acompanhamento
            })

        if cur.rowcount == 0:
            return jsonify({"sucesso": False, "erro": "Nenhuma ação foi atualizada."}), 404

        conn.commit()
        
        return jsonify({
            "sucesso": True,
            "mensagem": f"Ação marcada como {novo_status.lower()} com sucesso!"
        })

    except Exception as e:
        if conn: conn.rollback()
        print(f"Erro ao atualizar status da ação: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {str(e)}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/excluir_acao/<int:cod_acompanhamento>', methods=['DELETE'])
def api_excluir_acao_hiperdia(cod_acompanhamento):
    """Endpoint para excluir completamente uma ação"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Primeiro, buscar informações da ação para identificar referências posteriores
        cur.execute("""
            SELECT cod_cidadao, cod_acao, status_acao, cod_acao_origem
            FROM sistemaaps.tb_hiperdia_has_acompanhamento
            WHERE cod_acompanhamento = %(cod_acompanhamento)s
        """, {'cod_acompanhamento': cod_acompanhamento})
        
        acao_info = cur.fetchone()
        if not acao_info:
            return jsonify({"sucesso": False, "erro": "Ação não encontrada."}), 404

        cod_cidadao, cod_acao, status_acao, cod_acao_origem = acao_info

        # Buscar todas as ações posteriores que referenciam esta ação
        cur.execute("""
            SELECT cod_acompanhamento, cod_acao, status_acao
            FROM sistemaaps.tb_hiperdia_has_acompanhamento
            WHERE cod_acao_origem = %(cod_acompanhamento)s
            ORDER BY cod_acompanhamento ASC
        """, {'cod_acompanhamento': cod_acompanhamento})
        
        acoes_posteriores = cur.fetchall()

        # Iniciar transação para remoção em cascata
        try:
            # Tabelas relacionadas para limpeza
            tabelas_relacionadas = [
                'sistemaaps.tb_hiperdia_mrpa',
                'sistemaaps.tb_hiperdia_tratamento',
                'sistemaaps.tb_hiperdia_resultados_exames',
                'sistemaaps.tb_hiperdia_risco_cv',
                'sistemaaps.tb_hiperdia_nutricao',
                'sistemaaps.tb_hiperdia_has_cardiologia'
            ]

            # 1. Remover dados relacionados à ação principal
            for tabela in tabelas_relacionadas:
                cur.execute(f"""
                    DELETE FROM {tabela}
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s
                """, {'cod_acompanhamento': cod_acompanhamento})

            # 2. Remover ações posteriores em cascata
            for acao_posterior in acoes_posteriores:
                cod_acompanhamento_posterior = acao_posterior[0]
                
                # Remover dados relacionados à ação posterior
                for tabela in tabelas_relacionadas:
                    cur.execute(f"""
                        DELETE FROM {tabela}
                        WHERE cod_acompanhamento = %(cod_acompanhamento)s
                    """, {'cod_acompanhamento': cod_acompanhamento_posterior})

                # Remover a ação posterior
                cur.execute("""
                    DELETE FROM sistemaaps.tb_hiperdia_has_acompanhamento
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s
                """, {'cod_acompanhamento': cod_acompanhamento_posterior})

            # 3. Finalmente, remover a ação principal
            cur.execute("""
                DELETE FROM sistemaaps.tb_hiperdia_has_acompanhamento
                WHERE cod_acompanhamento = %(cod_acompanhamento)s
            """, {'cod_acompanhamento': cod_acompanhamento})

            # Verificar se alguma linha foi realmente removida
            if cur.rowcount == 0:
                conn.rollback()
                return jsonify({"sucesso": False, "erro": "Ação não encontrada para remoção."}), 404

            conn.commit()
            
            # Contar quantas ações foram removidas
            total_removidas = 1 + len(acoes_posteriores)  # ação principal + ações posteriores
            
            if len(acoes_posteriores) == 0:
                mensagem = "Ação excluída com sucesso!"
            else:
                mensagem = f"Ação e {len(acoes_posteriores)} ação(ões) posterior(es) excluída(s) com sucesso!"
            
            return jsonify({
                "sucesso": True, 
                "mensagem": mensagem,
                "acoes_removidas": total_removidas
            })

        except Exception as e:
            conn.rollback()
            print(f"Erro durante a exclusão em cascata: {e}")
            return jsonify({"sucesso": False, "erro": f"Erro durante a exclusão: {str(e)}"}), 500

    except Exception as e:
        if conn: conn.rollback()
        print(f"Erro ao excluir ação do Hiperdia: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {str(e)}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/pending_action/<int:cod_cidadao>/<int:cod_acao>', methods=['GET'])
def api_hiperdia_pending_action_by_type(cod_cidadao, cod_acao):
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        cur.execute("""
            SELECT cod_acompanhamento, cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao
            FROM sistemaaps.tb_hiperdia_has_acompanhamento
            WHERE cod_cidadao = %(cod_cidadao)s
              AND cod_acao = %(cod_acao)s
              AND status_acao = 'PENDENTE'
            ORDER BY data_agendamento ASC
            LIMIT 1;
        """, {'cod_cidadao': cod_cidadao, 'cod_acao': cod_acao})
        
        pending_action = cur.fetchone()

        if pending_action:
            # Format dates for consistency with frontend
            action_dict = dict(pending_action)
            if action_dict.get('data_agendamento') and isinstance(action_dict['data_agendamento'], date):
                action_dict['data_agendamento'] = action_dict['data_agendamento'].strftime('%Y-%m-%d')
            if action_dict.get('data_realizacao') and isinstance(action_dict['data_realizacao'], date):
                action_dict['data_realizacao'] = action_dict['data_realizacao'].strftime('%Y-%m-%d')
            return jsonify(action_dict)
        else:
            return jsonify({"message": "No pending action found."}), 404

    except Exception as e:
        print(f"Erro ao buscar ação pendente por tipo: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {str(e)}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/latest_pending_action/<int:cod_cidadao>/<int:cod_acao>', methods=['GET'])
def api_hiperdia_latest_pending_action_by_type(cod_cidadao, cod_acao):
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        cur.execute("""
            SELECT cod_acompanhamento, cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao
            FROM sistemaaps.tb_hiperdia_has_acompanhamento
            WHERE cod_cidadao = %(cod_cidadao)s
              AND cod_acao = %(cod_acao)s
              AND status_acao = 'PENDENTE'
            ORDER BY data_agendamento DESC, cod_acompanhamento DESC
            LIMIT 1;
        """, {'cod_cidadao': cod_cidadao, 'cod_acao': cod_acao})
        
        latest_pending_action = cur.fetchone()

        if latest_pending_action:
            # Format dates for consistency with frontend
            action_dict = dict(latest_pending_action)
            if action_dict.get('data_agendamento') and isinstance(action_dict['data_agendamento'], date):
                action_dict['data_agendamento'] = action_dict['data_agendamento'].strftime('%Y-%m-%d')
            if action_dict.get('data_realizacao') and isinstance(action_dict['data_realizacao'], date):
                action_dict['data_realizacao'] = action_dict['data_realizacao'].strftime('%Y-%m-%d')
            return jsonify(action_dict)
        else:
            return jsonify({"message": "No latest pending action found."}), 404

    except Exception as e:
        print(f"Erro ao buscar a ação pendente mais recente por tipo: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {str(e)}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/update_acao/<int:cod_acompanhamento>', methods=['PUT'])
def api_hiperdia_update_acao(cod_acompanhamento):
    data = request.get_json()
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        cod_cidadao = data.get('cod_cidadao')
        cod_acao_atual = data.get('cod_acao_atual')
        data_acao_atual_str = data.get('data_acao_atual')
        observacoes_atuais = data.get('observacoes')
        responsavel_pela_acao = data.get('responsavel_pela_acao')
        status_acao = data.get('status_acao', 'REALIZADA') # Default to REALIZADA for updates

        print(f"[LOG] Iniciando api_hiperdia_update_acao - cod_acompanhamento: {cod_acompanhamento}, cod_acao_atual: {cod_acao_atual}")
        print(f"[LOG] Dados recebidos: {data}")

        if not all([cod_cidadao, cod_acao_atual, data_acao_atual_str]):
            print(f"[LOG] Dados incompletos - cod_cidadao: {cod_cidadao}, cod_acao_atual: {cod_acao_atual}, data_acao_atual_str: {data_acao_atual_str}")
            return jsonify({"sucesso": False, "erro": "Dados incompletos para atualização."}), 400

        data_realizacao_acao = datetime.strptime(data_acao_atual_str, '%Y-%m-%d').date()

        # Update the main accompaniment record
        sql_update_acompanhamento = """
            UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
            SET status_acao = %(status_acao)s,
                data_realizacao = %(data_realizacao)s,
                observacoes = %(observacoes)s,
                responsavel_pela_acao = %(responsavel_pela_acao)s
            WHERE cod_acompanhamento = %(cod_acompanhamento)s
            RETURNING cod_acompanhamento;
        """
        print(f"[LOG] Executando UPDATE na tabela tb_hiperdia_has_acompanhamento")
        cur.execute(sql_update_acompanhamento, {
            'status_acao': status_acao,
            'data_realizacao': data_realizacao_acao,
            'observacoes': observacoes_atuais,
            'responsavel_pela_acao': responsavel_pela_acao,
            'cod_acompanhamento': cod_acompanhamento
        })
        updated_row = cur.fetchone()
        if not updated_row:
            print(f"[LOG] ERRO: Ação de acompanhamento não encontrada para atualização")
            return jsonify({"sucesso": False, "erro": "Ação de acompanhamento não encontrada para atualização."}), 404
        
        cod_acompanhamento_atualizado = updated_row[0]
        print(f"[LOG] Ação atualizada com sucesso - cod_acompanhamento: {cod_acompanhamento_atualizado}")

        # Handle specific data based on action type
        if int(cod_acao_atual) == 3: # Modificar tratamento
            print(f"[LOG] Processando Modificar tratamento (cod_acao = 3)")
            treatment_data = data.get('treatment_modification')
            if not treatment_data:
                return jsonify({"sucesso": False, "erro": "Dados da modificação de tratamento não fornecidos."}), 400
            
            # Check if a treatment record already exists for this cod_acompanhamento
            cur.execute("SELECT 1 FROM sistemaaps.tb_hiperdia_tratamento WHERE cod_acompanhamento = %s", (cod_acompanhamento_atualizado,))
            exists = cur.fetchone()

            if exists:
                sql_upsert_tratamento = """
                    UPDATE sistemaaps.tb_hiperdia_tratamento
                    SET tipo_ajuste = %(tipo_ajuste)s,
                        medicamentos_novos = %(medicamentos_novos)s,
                        data_modificacao = %(data_modificacao)s
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                """
            else:
                sql_upsert_tratamento = """
                    INSERT INTO sistemaaps.tb_hiperdia_tratamento
                    (cod_acompanhamento, tipo_ajuste, medicamentos_novos, data_modificacao)
                    VALUES (%(cod_acompanhamento)s, %(tipo_ajuste)s, %(medicamentos_novos)s, %(data_modificacao)s);
                """
            cur.execute(sql_upsert_tratamento, {
                'cod_acompanhamento': cod_acompanhamento_atualizado,
                'tipo_ajuste': treatment_data.get('tipo_ajuste'),
                'medicamentos_novos': treatment_data.get('medicamentos_novos'),
                'data_modificacao': data_realizacao_acao
            })
            
            # Após completar "Modificar tratamento", criar automaticamente "Iniciar MRPA" para 30 dias
            print(f"[LOG] Criando ação futura 'Iniciar MRPA' automaticamente para 30 dias após modificação de tratamento")
            cod_proxima_acao_pendente = 1 # Iniciar MRPA
            data_agendamento_proxima = data_realizacao_acao + timedelta(days=30) # 30 dias
            sql_insert_pendente_mrpa = '''
                INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'MRPA de controle após modificação de tratamento (30 dias).', %(responsavel_pela_acao)s);
            '''
            cur.execute(sql_insert_pendente_mrpa, {
                'cod_cidadao': cod_cidadao,
                'cod_acao': cod_proxima_acao_pendente,
                'data_agendamento': data_agendamento_proxima,
                'cod_acao_origem': cod_acompanhamento_atualizado,
                'responsavel_pela_acao': responsavel_pela_acao
            })
            print(f"[LOG] Ação futura 'Iniciar MRPA' criada automaticamente para 30 dias ({data_agendamento_proxima})")

        elif int(cod_acao_atual) == 2: # Avaliar MRPA
            print(f"[LOG] Processando Avaliar MRPA (cod_acao = 2)")
            
            # Verificar se há dados da avaliação do MRPA (aceitar ambas as chaves)
            mrpa_assessment_data = data.get('mrpa_assessment_data') or data.get('mrpa_results')
            if not mrpa_assessment_data:
                print(f"[LOG] ERRO: Dados da avaliação do MRPA não fornecidos")
                return jsonify({"sucesso": False, "erro": "Dados da avaliação do MRPA não fornecidos."}), 400
            
            print(f"[LOG] Dados da avaliação do MRPA recebidos: {mrpa_assessment_data}")
            
            # Buscar ação pendente de "Avaliar MRPA" para este paciente
            cur.execute('''
                SELECT cod_acompanhamento
                FROM sistemaaps.tb_hiperdia_has_acompanhamento
                WHERE cod_cidadao = %(cod_cidadao)s
                    AND cod_acao = 2
                    AND status_acao = 'PENDENTE'
                ORDER BY data_agendamento ASC
                LIMIT 1;
            ''', {'cod_cidadao': cod_cidadao})
            pending_mrpa = cur.fetchone()

            if pending_mrpa:
                # Se existe uma ação pendente, atualizar ela
                cod_acompanhamento_realizado = pending_mrpa[0]
                print(f"[LOG] Atualizando ação pendente de Avaliar MRPA - cod_acompanhamento: {cod_acompanhamento_realizado}")
                cur.execute('''
                    UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                    SET status_acao = 'REALIZADA',
                        data_realizacao = %(data_realizacao)s,
                        observacoes = %(observacoes)s,
                        responsavel_pela_acao = %(responsavel_pela_acao)s
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                ''', {
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes_atuais,
                    'responsavel_pela_acao': responsavel_pela_acao,
                    'cod_acompanhamento': cod_acompanhamento_realizado
                })
                print(f"[LOG] Ação pendente de Avaliar MRPA atualizada para REALIZADA")
            else:
                # Se não existe ação pendente, criar uma nova ação realizada
                print(f"[LOG] Criando nova ação de Avaliar MRPA")
                cur.execute('''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, 2, 'REALIZADA', %(data_realizacao)s, %(data_realizacao)s, %(observacoes)s, %(responsavel_pela_acao)s)
                    RETURNING cod_acompanhamento;
                ''', {
                    'cod_cidadao': cod_cidadao,
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes_atuais,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                cod_acompanhamento_realizado = cur.fetchone()[0]
                print(f"[LOG] Nova ação de Avaliar MRPA criada - cod_acompanhamento: {cod_acompanhamento_realizado}")

            # Mapear os campos do frontend para o formato esperado pelo banco
            mapped_data = {
                'pressao_sistolica': mrpa_assessment_data.get('media_pa_sistolica') or mrpa_assessment_data.get('pressao_sistolica'),
                'pressao_diastolica': mrpa_assessment_data.get('media_pa_diastolica') or mrpa_assessment_data.get('pressao_diastolica'),
                'frequencia_cardiaca': mrpa_assessment_data.get('frequencia_cardiaca', 0),
                'decisao_tratamento': mrpa_assessment_data.get('decision') or mrpa_assessment_data.get('decisao_tratamento', ''),
                'observacoes': mrpa_assessment_data.get('analise_mrpa') or mrpa_assessment_data.get('observacoes', '')
            }
            
            # Salvar os dados da avaliação do MRPA
            print(f"[LOG] Salvando dados da avaliação do MRPA na tabela tb_hiperdia_mrpa")
            
            # Determinar status_mrpa baseado na decisão de tratamento
            decisao_tratamento = mrpa_assessment_data.get('decision', '').lower()
            # 0 = Modificar Tratamento (Hipertensão descompensada)
            # 1 = Manter Tratamento (Hipertensão controlada)
            # 2 = Encaminhar Cardiologia
            if 'cardiology' in decisao_tratamento or 'cardiologia' in decisao_tratamento:
                status_mrpa = 2
            elif 'manter' in decisao_tratamento or 'maintain' in decisao_tratamento or 'manter tratamento' in decisao_tratamento:
                status_mrpa = 1
            else:
                status_mrpa = 0
            print(f"[LOG] Decisão de tratamento: {decisao_tratamento}, status_mrpa: {status_mrpa}")
            
            sql_insert_mrpa = """
                INSERT INTO sistemaaps.tb_hiperdia_mrpa
                (cod_acompanhamento, data_mrpa, media_pa_sistolica, media_pa_diastolica, analise_mrpa, status_mrpa)
                VALUES (%(cod_acompanhamento)s, %(data_mrpa)s, %(media_pa_sistolica)s, %(media_pa_diastolica)s, %(analise_mrpa)s, %(status_mrpa)s);
            """
            params = {
                'cod_acompanhamento': cod_acompanhamento_realizado,
                'data_mrpa': data_realizacao_acao,
                'media_pa_sistolica': mapped_data.get('pressao_sistolica'),
                'media_pa_diastolica': mapped_data.get('pressao_diastolica'),
                'analise_mrpa': mapped_data.get('observacoes'),
                'status_mrpa': status_mrpa
            }
            cur.execute(sql_insert_mrpa, params)
            print(f"[LOG] Dados da avaliação do MRPA salvos com sucesso com status_mrpa: {status_mrpa}")
            
            # Verificar se a decisão é "Manter tratamento" e criar "Solicitar Exames" se for
            if 'manter' in decisao_tratamento or 'maintain' in decisao_tratamento or 'manter tratamento' in decisao_tratamento:
                print(f"[LOG] Decisão é 'Manter tratamento'. Criando ação 'Solicitar Exames' automaticamente.")
                
                # Criar ação futura "Solicitar Exames" (pendente) para hoje
                cod_proxima_acao_pendente = 4 # Solicitar Exames
                data_agendamento_proxima = data_realizacao_acao # Mesma data (hoje)
                sql_insert_pendente = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Solicitação de exames após avaliação do MRPA com decisão de manter tratamento.', %(responsavel_pela_acao)s);
                '''
                cur.execute(sql_insert_pendente, {
                    'cod_cidadao': cod_cidadao,
                    'cod_acao': cod_proxima_acao_pendente,
                    'data_agendamento': data_agendamento_proxima,
                    'cod_acao_origem': cod_acompanhamento_realizado,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                print(f"[LOG] Ação futura 'Solicitar Exames' criada automaticamente")
            elif 'modificar' in decisao_tratamento or 'modify' in decisao_tratamento or 'modificar tratamento' in decisao_tratamento:
                print(f"[LOG] Decisão é 'Modificar tratamento'. Criando ação 'Modificar tratamento' automaticamente.")
                
                # Criar ação futura "Modificar tratamento" (pendente) para hoje
                cod_proxima_acao_pendente = 3 # Modificar tratamento
                data_agendamento_proxima = data_realizacao_acao # Mesma data (hoje)
                sql_insert_pendente = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Tratamento a ser modificado conforme avaliação do MRPA.', %(responsavel_pela_acao)s);
                '''
                cur.execute(sql_insert_pendente, {
                    'cod_cidadao': cod_cidadao,
                    'cod_acao': cod_proxima_acao_pendente,
                    'data_agendamento': data_agendamento_proxima,
                    'cod_acao_origem': cod_acompanhamento_realizado,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                print(f"[LOG] Ação futura 'Modificar tratamento' criada automaticamente")
            elif 'cardiology' in decisao_tratamento or 'cardiologia' in decisao_tratamento:
                print(f"[LOG] Decisão é 'Encaminhar Cardiologia'. Criando ação 'Encaminhar Cardiologia' automaticamente.")
                
                # Criar ação futura "Encaminhar Cardiologia" (pendente) para hoje
                cod_proxima_acao_pendente = 10 # Encaminhar Cardiologia
                data_agendamento_proxima = data_realizacao_acao # Mesma data (hoje)
                sql_insert_pendente = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Encaminhamento para cardiologia conforme avaliação do MRPA.', %(responsavel_pela_acao)s);
                '''
                cur.execute(sql_insert_pendente, {
                    'cod_cidadao': cod_cidadao,
                    'cod_acao': cod_proxima_acao_pendente,
                    'data_agendamento': data_agendamento_proxima,
                    'cod_acao_origem': cod_acompanhamento_realizado,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                print(f"[LOG] Ação futura 'Encaminhar Cardiologia' criada automaticamente")
                
                # Também registrar imediatamente o encaminhamento na tabela de cardiologia
                print(f"[LOG] Registrando encaminhamento na tabela tb_hiperdia_has_cardiologia")
                sql_insert_cardiologia = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_cardiologia
                    (cod_acompanhamento, cod_cidadao, tipo_acao, data_acao, profissional_responsavel, observacoes)
                    VALUES (%(cod_acompanhamento)s, %(cod_cidadao)s, 10, %(data_acao)s, %(profissional_responsavel)s, %(observacoes)s);
                '''
                cur.execute(sql_insert_cardiologia, {
                    'cod_acompanhamento': cod_acompanhamento_realizado,
                    'cod_cidadao': cod_cidadao,
                    'data_acao': data_realizacao_acao,
                    'profissional_responsavel': responsavel_pela_acao,
                    'observacoes': 'Encaminhamento para cardiologia baseado na avaliação do MRPA com status_mrpa=2'
                })
                print(f"[LOG] Encaminhamento para cardiologia registrado com sucesso")
            else:
                print(f"[LOG] Decisão não reconhecida: '{decisao_tratamento}'. Não será criada ação automática.")

        elif int(cod_acao_atual) == 5: # Avaliar Exames
            print(f"[LOG] Processando Avaliar Exames (cod_acao = 5) - buscando ação pendente existente")
            
            # Verificar se há dados dos resultados dos exames
            lab_exam_results = data.get('lab_exam_results')
            if not lab_exam_results:
                print(f"[LOG] ERRO: Resultados dos exames não fornecidos")
                return jsonify({"sucesso": False, "erro": "Resultados dos exames não fornecidos."}), 400
            
            print(f"[LOG] Dados dos resultados dos exames recebidos: {lab_exam_results}")
            
            # Validação dos campos numéricos para evitar overflow
            if lab_exam_results.get('potassio') and float(lab_exam_results.get('potassio', 0)) > 99.9:
                return jsonify({"sucesso": False, "erro": "Valor do potássio não pode ser maior que 99.9."}), 400
            
            if lab_exam_results.get('hemoglobina_glicada') and float(lab_exam_results.get('hemoglobina_glicada', 0)) > 99.99:
                return jsonify({"sucesso": False, "erro": "Valor da hemoglobina glicada não pode ser maior que 99.99."}), 400
            
            if lab_exam_results.get('creatinina') and float(lab_exam_results.get('creatinina', 0)) > 99.99:
                return jsonify({"sucesso": False, "erro": "Valor da creatinina não pode ser maior que 99.99."}), 400
            
            if lab_exam_results.get('acido_urico') and float(lab_exam_results.get('acido_urico', 0)) > 99.99:
                return jsonify({"sucesso": False, "erro": "Valor do ácido úrico não pode ser maior que 99.99."}), 400
            
            # Buscar ação pendente de "Avaliar Exames" para este paciente
            cur.execute('''
                SELECT cod_acompanhamento
                FROM sistemaaps.tb_hiperdia_has_acompanhamento
                WHERE cod_cidadao = %(cod_cidadao)s
                    AND cod_acao = 5
                    AND status_acao = 'PENDENTE'
                ORDER BY data_agendamento ASC
                LIMIT 1;
            ''', {'cod_cidadao': cod_cidadao})
            pending_exames = cur.fetchone()

            if pending_exames:
                # Se existe uma ação pendente, atualizar ela
                cod_acompanhamento_realizado = pending_exames[0]
                print(f"[LOG] Atualizando ação pendente de Avaliar Exames - cod_acompanhamento: {cod_acompanhamento_realizado}")
                cur.execute('''
                    UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                    SET status_acao = 'REALIZADA',
                        data_realizacao = %(data_realizacao)s,
                        observacoes = %(observacoes)s,
                        responsavel_pela_acao = %(responsavel_pela_acao)s
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                ''', {
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes_atuais,
                    'responsavel_pela_acao': responsavel_pela_acao,
                    'cod_acompanhamento': cod_acompanhamento_realizado
                })
                print(f"[LOG] Ação pendente de Avaliar Exames atualizada para REALIZADA")
            else:
                # Se não existe ação pendente, criar uma nova ação realizada
                print(f"[LOG] Criando nova ação de Avaliar Exames")
                cur.execute('''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, 5, 'REALIZADA', %(data_realizacao)s, %(data_realizacao)s, %(observacoes)s, %(responsavel_pela_acao)s)
                    RETURNING cod_acompanhamento;
                ''', {
                    'cod_cidadao': cod_cidadao,
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes_atuais,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                cod_acompanhamento_realizado = cur.fetchone()[0]
                print(f"[LOG] Nova ação de Avaliar Exames criada - cod_acompanhamento: {cod_acompanhamento_realizado}")

            # Salvar os dados dos resultados dos exames
            print(f"[LOG] Salvando dados dos resultados dos exames na tabela tb_hiperdia_resultados_exames")
            cur.execute("SELECT 1 FROM sistemaaps.tb_hiperdia_resultados_exames WHERE cod_acompanhamento = %s", (cod_acompanhamento_realizado,))
            exists = cur.fetchone()

            if exists:
                sql_upsert_exames = """
                    UPDATE sistemaaps.tb_hiperdia_resultados_exames
                    SET data_avaliacao = %(data_avaliacao)s, colesterol_total = %(colesterol_total)s, hdl = %(hdl)s, 
                        ldl = %(ldl)s, triglicerideos = %(triglicerideos)s, glicemia_jejum = %(glicemia_jejum)s, 
                        hemoglobina_glicada = %(hemoglobina_glicada)s, ureia = %(ureia)s, creatinina = %(creatinina)s, 
                        sodio = %(sodio)s, potassio = %(potassio)s, acido_urico = %(acido_urico)s
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                """
            else:
                sql_upsert_exames = """
                    INSERT INTO sistemaaps.tb_hiperdia_resultados_exames
                    (cod_acompanhamento, data_avaliacao, colesterol_total, hdl, ldl, triglicerideos, glicemia_jejum, hemoglobina_glicada, ureia, creatinina, sodio, potassio, acido_urico)
                    VALUES (%(cod_acompanhamento)s, %(data_avaliacao)s, %(colesterol_total)s, %(hdl)s, %(ldl)s, %(triglicerideos)s, %(glicemia_jejum)s, %(hemoglobina_glicada)s, %(ureia)s, %(creatinina)s, %(sodio)s, %(potassio)s, %(acido_urico)s);
                """
            params = {'cod_acompanhamento': cod_acompanhamento_realizado, 'data_avaliacao': data_realizacao_acao}
            params.update(lab_exam_results)
            cur.execute(sql_upsert_exames, params)
            print(f"[LOG] Dados dos resultados dos exames salvos com sucesso")

            # Criar automaticamente "Avaliar RCV" (cod_acao = 6) como PENDENTE
            print(f"[LOG] Criando ação 'Avaliar RCV' automaticamente")
            cod_proxima_acao_pendente = 6 # Avaliar RCV
            data_agendamento_proxima = data_realizacao_acao # Mesma data (hoje)
            sql_insert_pendente = '''
                INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Avaliação de Risco Cardiovascular após análise dos exames.', %(responsavel_pela_acao)s);
            '''
            cur.execute(sql_insert_pendente, {
                'cod_cidadao': cod_cidadao,
                'cod_acao': cod_proxima_acao_pendente,
                'data_agendamento': data_agendamento_proxima,
                'cod_acao_origem': cod_acompanhamento_realizado,
                'responsavel_pela_acao': responsavel_pela_acao
            })
            print(f"[LOG] Ação futura 'Avaliar RCV' criada automaticamente")

        elif int(cod_acao_atual) == 8: # Registrar consulta nutrição
            print(f"[LOG] Processando Registrar consulta nutrição (cod_acao = 8)")
            nutrition_data = data.get('nutrition_data')
            if not nutrition_data:
                return jsonify({"sucesso": False, "erro": "Dados da consulta de nutrição não fornecidos."}), 400
            # Buscar ação pendente de "Registrar consulta nutrição" para este paciente
            cur.execute('''
                SELECT cod_acompanhamento
                FROM sistemaaps.tb_hiperdia_has_acompanhamento
                WHERE cod_cidadao = %(cod_cidadao)s
                  AND cod_acao = 8
                  AND status_acao = 'PENDENTE'
                ORDER BY data_agendamento ASC
                LIMIT 1;
            ''', {'cod_cidadao': cod_cidadao})
            pending_nutricao = cur.fetchone()
            if pending_nutricao:
                # Se existe uma ação pendente, atualizar ela
                cod_acompanhamento_realizado = pending_nutricao[0]
                cur.execute('''
                    UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                    SET status_acao = 'REALIZADA',
                        data_realizacao = %(data_realizacao)s,
                        observacoes = %(observacoes)s,
                        responsavel_pela_acao = %(responsavel_pela_acao)s
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                ''', {
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes_atuais,
                    'responsavel_pela_acao': responsavel_pela_acao,
                    'cod_acompanhamento': cod_acompanhamento_realizado
                })
                # Atualiza/insere dados na tabela de nutrição
                cur.execute("SELECT 1 FROM sistemaaps.tb_hiperdia_nutricao WHERE cod_acompanhamento = %s", (cod_acompanhamento_realizado,))
                exists = cur.fetchone()
                if exists:
                    sql_upsert_nutricao = """
                        UPDATE sistemaaps.tb_hiperdia_nutricao
                        SET data_avaliacao = %(data_avaliacao)s, peso = %(peso)s, imc = %(imc)s, 
                            circunferencia_abdominal = %(circunferencia_abdominal)s, orientacoes_nutricionais = %(orientacoes_nutricionais)s
                        WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                    """
                else:
                    sql_upsert_nutricao = """
                    INSERT INTO sistemaaps.tb_hiperdia_nutricao
                    (cod_acompanhamento, data_avaliacao, peso, imc, circunferencia_abdominal, orientacoes_nutricionais)
                    VALUES (%(cod_acompanhamento)s, %(data_avaliacao)s, %(peso)s, %(imc)s, %(circunferencia_abdominal)s, %(orientacoes_nutricionais)s);
                """
                params = {'cod_acompanhamento': cod_acompanhamento_realizado, 'data_avaliacao': data_realizacao_acao}
                params.update(nutrition_data)
                cur.execute(sql_upsert_nutricao, params)
                # Criar ação futura "Agendar Hiperdia" (pendente) para daqui a 1 ano
                cod_proxima_acao_pendente = 9 # Agendar Hiperdia
                data_agendamento_proxima = data_realizacao_acao + timedelta(days=365)
                sql_insert_pendente = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, '<span style="color: #1e40af; font-weight: bold;">Hiperdia concluído!</span> Agendar novo hiperdia para daqui 6 meses há 1 ano, conforme seja necessário.', %(responsavel_pela_acao)s);
                '''
                cur.execute(sql_insert_pendente, {
                    'cod_cidadao': cod_cidadao,
                    'cod_acao': cod_proxima_acao_pendente,
                    'data_agendamento': data_agendamento_proxima,
                    'cod_acao_origem': cod_acompanhamento_realizado,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
            else:
                print(f"[LOG] Nenhuma ação pendente de Registrar Nutrição encontrada para atualizar. Nenhuma nova ação será criada.")

        elif int(cod_acao_atual) == 4: # Solicitar Exames
            print(f"[LOG] Processando Solicitar Exames (cod_acao = 4)")
            
            # Buscar ação pendente de "Solicitar Exames" para este paciente
            cur.execute('''
                SELECT cod_acompanhamento
                FROM sistemaaps.tb_hiperdia_has_acompanhamento
                WHERE cod_cidadao = %(cod_cidadao)s
                    AND cod_acao = 4
                    AND status_acao = 'PENDENTE'
                ORDER BY data_agendamento ASC
                LIMIT 1;
            ''', {'cod_cidadao': cod_cidadao})
            pending_exames = cur.fetchone()

            if pending_exames:
                # Se existe uma ação pendente, atualizar ela
                cod_acompanhamento_realizado = pending_exames[0]
                print(f"[LOG] Atualizando ação pendente de Solicitar Exames - cod_acompanhamento: {cod_acompanhamento_realizado}")
                cur.execute('''
                    UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                    SET status_acao = 'REALIZADA',
                        data_realizacao = %(data_realizacao)s,
                        observacoes = %(observacoes)s,
                        responsavel_pela_acao = %(responsavel_pela_acao)s
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                ''', {
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes,
                    'responsavel_pela_acao': responsavel_pela_acao,
                    'cod_acompanhamento': cod_acompanhamento_realizado
                })
                print(f"[LOG] Ação pendente de Solicitar Exames atualizada para REALIZADA")
            else:
                # Se não existe ação pendente, criar uma nova ação realizada
                print(f"[LOG] Criando nova ação de Solicitar Exames")
                cur.execute('''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, 4, 'REALIZADA', %(data_realizacao)s, %(data_realizacao)s, %(observacoes)s, %(responsavel_pela_acao)s)
                    RETURNING cod_acompanhamento;
                ''', {
                    'cod_cidadao': cod_cidadao,
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                cod_acompanhamento_realizado = cur.fetchone()[0]
                print(f"[LOG] Nova ação de Solicitar Exames criada - cod_acompanhamento: {cod_acompanhamento_realizado}")

            # Criar automaticamente "Avaliar Exames" (cod_acao = 5) como PENDENTE
            print(f"[LOG] Criando ação 'Avaliar Exames' automaticamente")
            cod_proxima_acao_pendente = 5 # Avaliar Exames
            data_agendamento_proxima = data_realizacao_acao + timedelta(days=7) # 7 dias para aguardar os exames
            sql_insert_pendente = '''
                INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Avaliação dos exames solicitados após 7 dias.', %(responsavel_pela_acao)s);
            '''
            cur.execute(sql_insert_pendente, {
                'cod_cidadao': cod_cidadao,
                'cod_acao': cod_proxima_acao_pendente,
                'data_agendamento': data_agendamento_proxima,
                'cod_acao_origem': cod_acompanhamento_realizado,
                'responsavel_pela_acao': responsavel_pela_acao
            })
            print(f"[LOG] Ação futura 'Avaliar Exames' criada automaticamente para 7 dias")

        elif int(cod_acao_atual) == 6: # Avaliar RCV
            print(f"[LOG] Processando Avaliar RCV (cod_acao = 6) - atualizando ação existente")
            
            # Verificar se há dados da avaliação de RCV
            risk_assessment_data = data.get('risk_assessment_data')
            if not risk_assessment_data:
                print(f"[LOG] ERRO: Dados da avaliação de RCV não fornecidos")
                return jsonify({"sucesso": False, "erro": "Dados da avaliação de RCV não fornecidos."}), 400
            
            print(f"[LOG] Dados da avaliação de RCV recebidos: {risk_assessment_data}")
            
            # CORREÇÃO: Usar diretamente o cod_acompanhamento passado como parâmetro
            # em vez de fazer uma nova busca por ações pendentes
            cod_acompanhamento_realizado = cod_acompanhamento_atualizado
            print(f"[LOG] Usando cod_acompanhamento existente: {cod_acompanhamento_realizado}")

            # Salvar os dados do RCV
            print(f"[LOG] Salvando dados do RCV na tabela tb_hiperdia_risco_cv")
            cur.execute("SELECT 1 FROM sistemaaps.tb_hiperdia_risco_cv WHERE cod_acompanhamento = %s", (cod_acompanhamento_realizado,))
            exists = cur.fetchone()
            if exists:
                print(f"[LOG] Atualizando dados existentes do RCV")
                sql_upsert_rcv = """
                    UPDATE sistemaaps.tb_hiperdia_risco_cv
                    SET data_avaliacao = %(data_avaliacao)s, idade = %(idade)s, sexo = %(sexo)s, 
                        tabagismo = %(tabagismo)s, diabetes = %(diabetes)s, colesterol_total = %(colesterol_total)s, 
                        pressao_sistolica = %(pressao_sistolica)s
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                """
            else:
                print(f"[LOG] Inserindo novos dados do RCV")
                sql_upsert_rcv = """
                    INSERT INTO sistemaaps.tb_hiperdia_risco_cv
                    (cod_acompanhamento, data_avaliacao, idade, sexo, tabagismo, diabetes, colesterol_total, pressao_sistolica)
                    VALUES (%(cod_acompanhamento)s, %(data_avaliacao)s, %(idade)s, %(sexo)s, %(tabagismo)s, %(diabetes)s, %(colesterol_total)s, %(pressao_sistolica)s);
                """
            params = {'cod_acompanhamento': cod_acompanhamento_realizado, 'data_avaliacao': data_realizacao_acao}
            params.update(risk_assessment_data)
            cur.execute(sql_upsert_rcv, params)
            print(f"[LOG] Dados do RCV salvos com sucesso")

            # Criar ação futura "Encaminhar para nutrição" (pendente)
            cod_proxima_acao_pendente = 7 # Encaminhar para nutrição
            data_agendamento_proxima = data_realizacao_acao # pode ser para hoje mesmo
            sql_insert_pendente = '''
                INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Encaminhamento para avaliação nutricional.', %(responsavel_pela_acao)s);
            '''
            cur.execute(sql_insert_pendente, {
                'cod_cidadao': cod_cidadao,
                'cod_acao': cod_proxima_acao_pendente,
                'data_agendamento': data_agendamento_proxima,
                'cod_acao_origem': cod_acompanhamento_realizado,
                'responsavel_pela_acao': responsavel_pela_acao
            })
            print(f"[LOG] Ação futura 'Encaminhar para nutrição' criada")

        elif int(cod_acao_atual) == 1: # Solicitar MRPA
            print(f"[LOG] Processando Solicitar MRPA (cod_acao = 1)")
            
            # 1. Marcar "Agendar Hiperdia" (cod_acao = 9) pendente como REALIZADA
            print(f"[LOG] Marcando ações 'Agendar Hiperdia' pendentes como realizadas")
            sql_update_agendar_hiperdia = """
                UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                SET status_acao = 'REALIZADA',
                    data_agendamento = %(data_realizacao)s,
                    data_realizacao = %(data_realizacao)s,
                    responsavel_pela_acao = %(responsavel_pela_acao)s
                WHERE cod_cidadao = %(cod_cidadao)s
                  AND cod_acao = 9
                  AND status_acao = 'PENDENTE';
            """
            cur.execute(sql_update_agendar_hiperdia, {
                'data_realizacao': data_realizacao_acao,
                'responsavel_pela_acao': responsavel_pela_acao,
                'cod_cidadao': cod_cidadao
            })
            print(f"[LOG] Ações 'Agendar Hiperdia' pendentes marcadas como realizadas")
            
            # 2. Criar "Solicitar MRPA" como REALIZADA com a data de hoje
            print(f"[LOG] Criando ação 'Solicitar MRPA' como REALIZADA")
            sql_insert_solicitar_mrpa = """
                INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
                VALUES (%(cod_cidadao)s, 1, 'REALIZADA', %(data_realizacao)s, %(data_realizacao)s, 'Solicitação de MRPA para monitorização da pressão arterial.', %(responsavel_pela_acao)s)
                RETURNING cod_acompanhamento;
            """
            cur.execute(sql_insert_solicitar_mrpa, {
                'cod_cidadao': cod_cidadao,
                'data_realizacao': data_realizacao_acao,
                'responsavel_pela_acao': responsavel_pela_acao
            })
            cod_acompanhamento_criado = cur.fetchone()[0]
            print(f"[LOG] Ação 'Solicitar MRPA' criada como REALIZADA - cod_acompanhamento: {cod_acompanhamento_criado}")
            
            # 3. Criar automaticamente "Avaliar MRPA" (cod_acao = 2) como PENDENTE para 7 dias
            print(f"[LOG] Criando ação 'Avaliar MRPA' automaticamente para 7 dias")
            cod_proxima_acao_pendente = 2 # Avaliar MRPA
            data_agendamento_proxima = data_realizacao_acao + timedelta(days=7) # 7 dias
            sql_insert_pendente = '''
                INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Avaliação do MRPA após 7 dias de monitorização.', %(responsavel_pela_acao)s);
            '''
            cur.execute(sql_insert_pendente, {
                'cod_cidadao': cod_cidadao,
                'cod_acao': cod_proxima_acao_pendente,
                'data_agendamento': data_agendamento_proxima,
                'cod_acao_origem': cod_acompanhamento_criado,
                'responsavel_pela_acao': responsavel_pela_acao
            })
            print(f"[LOG] Ação futura 'Avaliar MRPA' criada automaticamente para 7 dias")

        elif int(cod_acao_atual) == 7: # Encaminhar para Nutrição
            print(f"[LOG] Criando ação pendente Registrar Nutrição (cod_acao=8)")
            try:
                cur.execute('''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, 8, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Aguardando registro da consulta de nutrição.', %(responsavel_pela_acao)s)
                ''', {
                    'cod_cidadao': cod_cidadao,
                    'data_agendamento': data_realizacao_acao,
                    'cod_acao_origem': cod_acompanhamento_atualizado,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                print(f"[LOG] Ação pendente Registrar Nutrição criada com sucesso")
            except Exception as e:
                print(f"[LOG] ERRO ao criar ação pendente Registrar Nutrição: {e}")

        conn.commit()
        print(f"[LOG] Commit realizado com sucesso - Ação {cod_acao_atual} atualizada")
        return jsonify({"sucesso": True, "mensagem": "Ação atualizada com sucesso!"})

    except Exception as e:
        if conn: conn.rollback()
        print(f"[LOG] ERRO ao atualizar ação do Hiperdia: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {str(e)}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()


@app.route('/api/hiperdia/medicamentos/<int:cod_cidadao>')
def api_get_medicamentos_hiperdia(cod_cidadao):
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        sql_query = """
            SELECT 
                codmedicamento, 
                codCidadao as cod_cidadao, 
                medicamento, 
                posologia, 
                dt_inicio_tratamento
            FROM 
                sistemaaps.mv_hiperdia_hipertensao_medicamentos
            WHERE 
                codCidadao = %(cod_cidadao)s
            ORDER BY 
                dt_inicio_tratamento DESC, medicamento ASC;
        """
        
        cur.execute(sql_query, {'cod_cidadao': cod_cidadao})
        medicamentos_db = cur.fetchall()
        
        medicamentos = []
        for row in medicamentos_db:
            med_dict = dict(row)
            if med_dict.get('dt_inicio_tratamento') and isinstance(med_dict['dt_inicio_tratamento'], date):
                med_dict['dt_inicio_tratamento'] = med_dict['dt_inicio_tratamento'].strftime('%d/%m/%Y')
            medicamentos.append(med_dict)
            
        return jsonify(medicamentos)

    except Exception as e:
        print(f"Erro na API /api/hiperdia/medicamentos: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/registrar_acao', methods=['POST'])
def api_registrar_acao_hiperdia():
    data = request.get_json()
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        cod_cidadao = data.get('cod_cidadao')
        cod_acao_atual = data.get('cod_acao_atual')
        data_acao_atual_str = data.get('data_acao_atual')
        observacoes = data.get('observacoes')
        responsavel_pela_acao = data.get('responsavel_pela_acao')
        status_acao = data.get('status_acao', 'REALIZADA')
        skip_auto_next_actions = data.get('skip_auto_next_actions', False)

        print(f"[LOG] Iniciando api_registrar_acao_hiperdia - cod_cidadao: {cod_cidadao}, cod_acao_atual: {cod_acao_atual}")
        print(f"[LOG] Dados recebidos: {data}")

        if not all([cod_cidadao, cod_acao_atual, data_acao_atual_str]):
            print(f"[LOG] Dados incompletos - cod_cidadao: {cod_cidadao}, cod_acao_atual: {cod_acao_atual}, data_acao_atual_str: {data_acao_atual_str}")
            return jsonify({"sucesso": False, "erro": "Dados incompletos para registro."}), 400

        data_realizacao_acao = datetime.strptime(data_acao_atual_str, '%Y-%m-%d').date()

        # Tratamento especial para "Iniciar MRPA" (cod_acao = 1)
        if int(cod_acao_atual) == 1: # Iniciar MRPA
            print(f"[LOG] Processando Iniciar MRPA (cod_acao = 1)")
            
            # Verificar se é uma ação futura (PENDENTE) ou ação normal (REALIZADA)
            if status_acao == 'PENDENTE' and skip_auto_next_actions:
                # Caso 1: Criação de ação futura PENDENTE (ex: vinda de "Modificar tratamento")
                print(f"[LOG] Criando 'Iniciar MRPA' como PENDENTE (ação futura)")
                sql_insert_iniciar_mrpa_pendente = """
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, 1, 'PENDENTE', %(data_agendamento)s, %(observacoes)s, %(responsavel_pela_acao)s)
                    RETURNING cod_acompanhamento;
                """
                cur.execute(sql_insert_iniciar_mrpa_pendente, {
                    'cod_cidadao': cod_cidadao,
                    'data_agendamento': data_realizacao_acao,  # Para PENDENTE, data_realizacao_acao é usada como data_agendamento
                    'observacoes': observacoes,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                cod_acompanhamento_criado = cur.fetchone()[0]
                print(f"[LOG] Ação 'Iniciar MRPA' PENDENTE criada - cod_acompanhamento: {cod_acompanhamento_criado}, data_agendamento: {data_realizacao_acao}")
            else:
                # Caso 2: Verificar se existe "Iniciar MRPA" PENDENTE para atualizar ou criar novo
                print(f"[LOG] Verificando se existe 'Iniciar MRPA' PENDENTE para atualizar")
                
                # Verificar se existe ação "Iniciar MRPA" PENDENTE para este cidadão
                sql_check_pendente = """
                    SELECT cod_acompanhamento FROM sistemaaps.tb_hiperdia_has_acompanhamento
                    WHERE cod_cidadao = %(cod_cidadao)s AND cod_acao = 1 AND status_acao = 'PENDENTE'
                    ORDER BY data_agendamento DESC LIMIT 1;
                """
                cur.execute(sql_check_pendente, {'cod_cidadao': cod_cidadao})
                existing_pendente = cur.fetchone()
                
                if existing_pendente:
                    # Caso 2a: Atualizar "Iniciar MRPA" PENDENTE existente para REALIZADA
                    cod_acompanhamento_existente = existing_pendente[0]
                    print(f"[LOG] Encontrada ação 'Iniciar MRPA' PENDENTE existente (cod_acompanhamento: {cod_acompanhamento_existente}). Atualizando para REALIZADA.")
                    
                    sql_update_iniciar_mrpa = """
                        UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                        SET status_acao = 'REALIZADA',
                            data_realizacao = %(data_realizacao)s,
                            responsavel_pela_acao = %(responsavel_pela_acao)s
                        WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                    """
                    cur.execute(sql_update_iniciar_mrpa, {
                        'data_realizacao': data_realizacao_acao,
                        'responsavel_pela_acao': responsavel_pela_acao,
                        'cod_acompanhamento': cod_acompanhamento_existente
                    })
                    cod_acompanhamento_criado = cod_acompanhamento_existente
                    print(f"[LOG] Ação 'Iniciar MRPA' atualizada para REALIZADA - cod_acompanhamento: {cod_acompanhamento_criado}")
                else:
                    # Caso 2b: Fluxo normal - marcar "Agendar Hiperdia" como realizado e criar novo "Iniciar MRPA" como REALIZADA
                    print(f"[LOG] Fluxo normal: marcando 'Agendar Hiperdia' como realizada e criando novo 'Iniciar MRPA' como REALIZADA")
                    
                    # 1. Marcar "Agendar Hiperdia" (cod_acao = 9) pendente como REALIZADA
                    print(f"[LOG] Marcando ações 'Agendar Hiperdia' pendentes como realizadas")
                    sql_update_agendar_hiperdia = """
                        UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                        SET status_acao = 'REALIZADA',
                            data_agendamento = %(data_realizacao)s,
                            data_realizacao = %(data_realizacao)s,
                            responsavel_pela_acao = %(responsavel_pela_acao)s
                        WHERE cod_cidadao = %(cod_cidadao)s
                          AND cod_acao = 9
                          AND status_acao = 'PENDENTE';
                    """
                    cur.execute(sql_update_agendar_hiperdia, {
                        'data_realizacao': data_realizacao_acao,
                        'responsavel_pela_acao': responsavel_pela_acao,
                        'cod_cidadao': cod_cidadao
                    })
                    print(f"[LOG] Ações 'Agendar Hiperdia' pendentes marcadas como realizadas")
                    
                    # 2. Criar "Iniciar MRPA" como REALIZADA com a data de hoje
                    print(f"[LOG] Criando ação 'Iniciar MRPA' como REALIZADA")
                    # Usar observação do usuário se fornecida, senão usar texto padrão
                    observacao_iniciar_mrpa = observacoes_atuais if observacoes_atuais and observacoes_atuais.strip() else 'Iniciar MRPA para monitorização da pressão arterial.'
                    sql_insert_iniciar_mrpa_realizada = """
                        INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                        (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
                        VALUES (%(cod_cidadao)s, 1, 'REALIZADA', %(data_realizacao)s, %(data_realizacao)s, %(observacoes)s, %(responsavel_pela_acao)s)
                        RETURNING cod_acompanhamento;
                    """
                    cur.execute(sql_insert_iniciar_mrpa_realizada, {
                        'cod_cidadao': cod_cidadao,
                        'data_realizacao': data_realizacao_acao,
                        'observacoes': observacao_iniciar_mrpa,
                        'responsavel_pela_acao': responsavel_pela_acao
                    })
                    cod_acompanhamento_criado = cur.fetchone()[0]
                    print(f"[LOG] Ação 'Iniciar MRPA' criada como REALIZADA - cod_acompanhamento: {cod_acompanhamento_criado}")
            
            # 3. Criar automaticamente "Avaliar MRPA" (cod_acao = 2) como PENDENTE para 7 dias
            # (apenas se não foi solicitado para pular ações automáticas)
            if not skip_auto_next_actions:
                print(f"[LOG] Criando ação 'Avaliar MRPA' automaticamente para 7 dias")
                cod_proxima_acao_pendente = 2 # Avaliar MRPA
                data_agendamento_proxima = data_realizacao_acao + timedelta(days=7) # 7 dias
                sql_insert_pendente = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Avaliação do MRPA após 7 dias de monitorização.', %(responsavel_pela_acao)s);
                '''
                cur.execute(sql_insert_pendente, {
                    'cod_cidadao': cod_cidadao,
                    'cod_acao': cod_proxima_acao_pendente,
                    'data_agendamento': data_agendamento_proxima,
                    'cod_acao_origem': cod_acompanhamento_criado,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                print(f"[LOG] Ação futura 'Avaliar MRPA' criada automaticamente para 7 dias")
            else:
                print(f"[LOG] Pulando criação automática de 'Avaliar MRPA' devido a skip_auto_next_actions=True")

        # Tratamento especial para "Avaliar MRPA" - buscar ação pendente existente
        elif int(cod_acao_atual) == 2: # Avaliar MRPA
            print(f"[LOG] Processando Avaliar MRPA (cod_acao = 2) - buscando ação pendente existente")
            
            # Verificar se há dados da avaliação do MRPA (aceitar ambas as chaves)
            mrpa_assessment_data = data.get('mrpa_assessment_data') or data.get('mrpa_results')
            if not mrpa_assessment_data:
                print(f"[LOG] ERRO: Dados da avaliação do MRPA não fornecidos")
                return jsonify({"sucesso": False, "erro": "Dados da avaliação do MRPA não fornecidos."}), 400
            
            print(f"[LOG] Dados da avaliação do MRPA recebidos: {mrpa_assessment_data}")
            
            # Buscar ação pendente de "Avaliar MRPA" para este paciente
            cur.execute('''
                SELECT cod_acompanhamento
                FROM sistemaaps.tb_hiperdia_has_acompanhamento
                WHERE cod_cidadao = %(cod_cidadao)s
                    AND cod_acao = 2
                    AND status_acao = 'PENDENTE'
                ORDER BY data_agendamento ASC
                LIMIT 1;
            ''', {'cod_cidadao': cod_cidadao})
            pending_mrpa = cur.fetchone()

            if pending_mrpa:
                # Se existe uma ação pendente, atualizar ela
                cod_acompanhamento_realizado = pending_mrpa[0]
                print(f"[LOG] Atualizando ação pendente de Avaliar MRPA - cod_acompanhamento: {cod_acompanhamento_realizado}")
                cur.execute('''
                    UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                    SET status_acao = 'REALIZADA',
                        data_realizacao = %(data_realizacao)s,
                        observacoes = %(observacoes)s,
                        responsavel_pela_acao = %(responsavel_pela_acao)s
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                ''', {
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes,
                    'responsavel_pela_acao': responsavel_pela_acao,
                    'cod_acompanhamento': cod_acompanhamento_realizado
                })
                print(f"[LOG] Ação pendente de Avaliar MRPA atualizada para REALIZADA")
            else:
                # Se não existe ação pendente, criar uma nova ação realizada
                print(f"[LOG] Criando nova ação de Avaliar MRPA")
                cur.execute('''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, 2, 'REALIZADA', %(data_realizacao)s, %(data_realizacao)s, %(observacoes)s, %(responsavel_pela_acao)s)
                    RETURNING cod_acompanhamento;
                ''', {
                    'cod_cidadao': cod_cidadao,
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                cod_acompanhamento_realizado = cur.fetchone()[0]
                print(f"[LOG] Nova ação de Avaliar MRPA criada - cod_acompanhamento: {cod_acompanhamento_realizado}")

            # Mapear os campos do frontend para o formato esperado pelo banco
            mapped_data = {
                'pressao_sistolica': mrpa_assessment_data.get('media_pa_sistolica') or mrpa_assessment_data.get('pressao_sistolica'),
                'pressao_diastolica': mrpa_assessment_data.get('media_pa_diastolica') or mrpa_assessment_data.get('pressao_diastolica'),
                'frequencia_cardiaca': mrpa_assessment_data.get('frequencia_cardiaca', 0),
                'decisao_tratamento': mrpa_assessment_data.get('decision') or mrpa_assessment_data.get('decisao_tratamento', ''),
                'observacoes': mrpa_assessment_data.get('analise_mrpa') or mrpa_assessment_data.get('observacoes', '')
            }
            
            # Salvar os dados da avaliação do MRPA
            print(f"[LOG] Salvando dados da avaliação do MRPA na tabela tb_hiperdia_mrpa")
            
            # Determinar status_mrpa baseado na decisão de tratamento
            decisao_tratamento = mrpa_assessment_data.get('decision', '').lower()
            # 0 = Modificar Tratamento (Hipertensão descompensada)
            # 1 = Manter Tratamento (Hipertensão controlada)
            # 2 = Encaminhar Cardiologia
            if 'cardiology' in decisao_tratamento or 'cardiologia' in decisao_tratamento:
                status_mrpa = 2
            elif 'manter' in decisao_tratamento or 'maintain' in decisao_tratamento or 'manter tratamento' in decisao_tratamento:
                status_mrpa = 1
            else:
                status_mrpa = 0
            print(f"[LOG] Decisão de tratamento: {decisao_tratamento}, status_mrpa: {status_mrpa}")
            
            sql_insert_mrpa = """
                INSERT INTO sistemaaps.tb_hiperdia_mrpa
                (cod_acompanhamento, data_mrpa, media_pa_sistolica, media_pa_diastolica, analise_mrpa, status_mrpa)
                VALUES (%(cod_acompanhamento)s, %(data_mrpa)s, %(media_pa_sistolica)s, %(media_pa_diastolica)s, %(analise_mrpa)s, %(status_mrpa)s);
            """
            params = {
                'cod_acompanhamento': cod_acompanhamento_realizado,
                'data_mrpa': data_realizacao_acao,
                'media_pa_sistolica': mapped_data.get('pressao_sistolica'),
                'media_pa_diastolica': mapped_data.get('pressao_diastolica'),
                'analise_mrpa': mapped_data.get('observacoes'),
                'status_mrpa': status_mrpa
            }
            cur.execute(sql_insert_mrpa, params)
            print(f"[LOG] Dados da avaliação do MRPA salvos com sucesso com status_mrpa: {status_mrpa}")
            
            # Verificar a decisão e criar ações futuras conforme necessário
            if 'manter' in decisao_tratamento or 'maintain' in decisao_tratamento or 'manter tratamento' in decisao_tratamento:
                print(f"[LOG] Decisão é 'Manter tratamento'. Criando ação 'Solicitar Exames' automaticamente.")
                
                # Criar ação futura "Solicitar Exames" (pendente) para hoje
                cod_proxima_acao_pendente = 4 # Solicitar Exames
                data_agendamento_proxima = data_realizacao_acao # Mesma data (hoje)
                sql_insert_pendente = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Solicitação de exames após avaliação do MRPA com decisão de manter tratamento.', %(responsavel_pela_acao)s);
                '''
                cur.execute(sql_insert_pendente, {
                    'cod_cidadao': cod_cidadao,
                    'cod_acao': cod_proxima_acao_pendente,
                    'data_agendamento': data_agendamento_proxima,
                    'cod_acao_origem': cod_acompanhamento_realizado,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                print(f"[LOG] Ação futura 'Solicitar Exames' criada automaticamente")
            elif 'modificar' in decisao_tratamento or 'modify' in decisao_tratamento or 'modificar tratamento' in decisao_tratamento:
                print(f"[LOG] Decisão é 'Modificar tratamento'. Criando ação 'Modificar tratamento' automaticamente.")
                
                # Criar ação futura "Modificar tratamento" (pendente) para hoje
                cod_proxima_acao_pendente = 3 # Modificar tratamento
                data_agendamento_proxima = data_realizacao_acao # Mesma data (hoje)
                sql_insert_pendente = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Tratamento a ser modificado conforme avaliação do MRPA.', %(responsavel_pela_acao)s);
                '''
                cur.execute(sql_insert_pendente, {
                    'cod_cidadao': cod_cidadao,
                    'cod_acao': cod_proxima_acao_pendente,
                    'data_agendamento': data_agendamento_proxima,
                    'cod_acao_origem': cod_acompanhamento_realizado,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                print(f"[LOG] Ação futura 'Modificar tratamento' criada automaticamente")
            elif 'cardiology' in decisao_tratamento or 'cardiologia' in decisao_tratamento:
                print(f"[LOG] Decisão é 'Encaminhar Cardiologia'. Criando ação 'Encaminhar Cardiologia' automaticamente.")
                
                # Criar ação futura "Encaminhar Cardiologia" (pendente) para hoje
                cod_proxima_acao_pendente = 10 # Encaminhar Cardiologia
                data_agendamento_proxima = data_realizacao_acao # Mesma data (hoje)
                sql_insert_pendente = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Encaminhamento para cardiologia conforme avaliação do MRPA.', %(responsavel_pela_acao)s);
                '''
                cur.execute(sql_insert_pendente, {
                    'cod_cidadao': cod_cidadao,
                    'cod_acao': cod_proxima_acao_pendente,
                    'data_agendamento': data_agendamento_proxima,
                    'cod_acao_origem': cod_acompanhamento_realizado,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                print(f"[LOG] Ação futura 'Encaminhar Cardiologia' criada automaticamente")
                
                # Também registrar imediatamente o encaminhamento na tabela de cardiologia
                print(f"[LOG] Registrando encaminhamento na tabela tb_hiperdia_has_cardiologia")
                sql_insert_cardiologia = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_cardiologia
                    (cod_acompanhamento, cod_cidadao, tipo_acao, data_acao, profissional_responsavel, observacoes)
                    VALUES (%(cod_acompanhamento)s, %(cod_cidadao)s, 10, %(data_acao)s, %(profissional_responsavel)s, %(observacoes)s);
                '''
                cur.execute(sql_insert_cardiologia, {
                    'cod_acompanhamento': cod_acompanhamento_realizado,
                    'cod_cidadao': cod_cidadao,
                    'data_acao': data_realizacao_acao,
                    'profissional_responsavel': responsavel_pela_acao,
                    'observacoes': 'Encaminhamento para cardiologia baseado na avaliação do MRPA com status_mrpa=2'
                })
                print(f"[LOG] Encaminhamento para cardiologia registrado com sucesso")
            else:
                print(f"[LOG] Decisão não reconhecida: '{decisao_tratamento}'. Não será criada ação automática.")

        elif int(cod_acao_atual) == 4: # Solicitar Exames
            print(f"[LOG] Processando Solicitar Exames (cod_acao = 4)")
            
            # Buscar ação pendente de "Solicitar Exames" para este paciente
            cur.execute('''
                SELECT cod_acompanhamento
                FROM sistemaaps.tb_hiperdia_has_acompanhamento
                WHERE cod_cidadao = %(cod_cidadao)s
                    AND cod_acao = 4
                    AND status_acao = 'PENDENTE'
                ORDER BY data_agendamento ASC
                LIMIT 1;
            ''', {'cod_cidadao': cod_cidadao})
            pending_exames = cur.fetchone()

            if pending_exames:
                # Se existe uma ação pendente, atualizar ela
                cod_acompanhamento_realizado = pending_exames[0]
                print(f"[LOG] Atualizando ação pendente de Solicitar Exames - cod_acompanhamento: {cod_acompanhamento_realizado}")
                cur.execute('''
                    UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                    SET status_acao = 'REALIZADA',
                        data_realizacao = %(data_realizacao)s,
                        observacoes = %(observacoes)s,
                        responsavel_pela_acao = %(responsavel_pela_acao)s
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                ''', {
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes,
                    'responsavel_pela_acao': responsavel_pela_acao,
                    'cod_acompanhamento': cod_acompanhamento_realizado
                })
                print(f"[LOG] Ação pendente de Solicitar Exames atualizada para REALIZADA")
            else:
                # Se não existe ação pendente, criar uma nova ação realizada
                print(f"[LOG] Criando nova ação de Solicitar Exames")
                cur.execute('''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, 4, 'REALIZADA', %(data_realizacao)s, %(data_realizacao)s, %(observacoes)s, %(responsavel_pela_acao)s)
                    RETURNING cod_acompanhamento;
                ''', {
                    'cod_cidadao': cod_cidadao,
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                cod_acompanhamento_realizado = cur.fetchone()[0]
                print(f"[LOG] Nova ação de Solicitar Exames criada - cod_acompanhamento: {cod_acompanhamento_realizado}")

            # Criar automaticamente "Avaliar Exames" (cod_acao = 5) como PENDENTE
            print(f"[LOG] Criando ação 'Avaliar Exames' automaticamente")
            cod_proxima_acao_pendente = 5 # Avaliar Exames
            data_agendamento_proxima = data_realizacao_acao + timedelta(days=7) # 7 dias para aguardar os exames
            sql_insert_pendente = '''
                INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Avaliação dos exames solicitados após 7 dias.', %(responsavel_pela_acao)s);
            '''
            cur.execute(sql_insert_pendente, {
                'cod_cidadao': cod_cidadao,
                'cod_acao': cod_proxima_acao_pendente,
                'data_agendamento': data_agendamento_proxima,
                'cod_acao_origem': cod_acompanhamento_realizado,
                'responsavel_pela_acao': responsavel_pela_acao
            })
            print(f"[LOG] Ação futura 'Avaliar Exames' criada automaticamente para 7 dias")

        elif int(cod_acao_atual) == 5: # Avaliar Exames
            print(f"[LOG] Processando Avaliar Exames (cod_acao = 5) - buscando ação pendente existente")
            
            # Verificar se há dados dos resultados dos exames
            lab_exam_results = data.get('lab_exam_results')
            if not lab_exam_results:
                print(f"[LOG] ERRO: Resultados dos exames não fornecidos")
                return jsonify({"sucesso": False, "erro": "Resultados dos exames não fornecidos."}), 400
            
            print(f"[LOG] Dados dos resultados dos exames recebidos: {lab_exam_results}")
            
            # Validação dos campos numéricos para evitar overflow
            if lab_exam_results.get('potassio') and float(lab_exam_results.get('potassio', 0)) > 99.9:
                return jsonify({"sucesso": False, "erro": "Valor do potássio não pode ser maior que 99.9."}), 400
            
            if lab_exam_results.get('hemoglobina_glicada') and float(lab_exam_results.get('hemoglobina_glicada', 0)) > 99.99:
                return jsonify({"sucesso": False, "erro": "Valor da hemoglobina glicada não pode ser maior que 99.99."}), 400
            
            if lab_exam_results.get('creatinina') and float(lab_exam_results.get('creatinina', 0)) > 99.99:
                return jsonify({"sucesso": False, "erro": "Valor da creatinina não pode ser maior que 99.99."}), 400
            
            if lab_exam_results.get('acido_urico') and float(lab_exam_results.get('acido_urico', 0)) > 99.99:
                return jsonify({"sucesso": False, "erro": "Valor do ácido úrico não pode ser maior que 99.99."}), 400
            
            # Buscar ação pendente de "Avaliar Exames" para este paciente
            cur.execute('''
                SELECT cod_acompanhamento
                FROM sistemaaps.tb_hiperdia_has_acompanhamento
                WHERE cod_cidadao = %(cod_cidadao)s
                    AND cod_acao = 5
                    AND status_acao = 'PENDENTE'
                ORDER BY data_agendamento ASC
                LIMIT 1;
            ''', {'cod_cidadao': cod_cidadao})
            pending_exames = cur.fetchone()

            if pending_exames:
                # Se existe uma ação pendente, atualizar ela
                cod_acompanhamento_realizado = pending_exames[0]
                print(f"[LOG] Atualizando ação pendente de Avaliar Exames - cod_acompanhamento: {cod_acompanhamento_realizado}")
                cur.execute('''
                    UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                    SET status_acao = 'REALIZADA',
                        data_realizacao = %(data_realizacao)s,
                        observacoes = %(observacoes)s,
                        responsavel_pela_acao = %(responsavel_pela_acao)s
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                ''', {
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes,
                    'responsavel_pela_acao': responsavel_pela_acao,
                    'cod_acompanhamento': cod_acompanhamento_realizado
                })
                print(f"[LOG] Ação pendente de Avaliar Exames atualizada para REALIZADA")
            else:
                # Se não existe ação pendente, criar uma nova ação realizada
                print(f"[LOG] Criando nova ação de Avaliar Exames")
                cur.execute('''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, 5, 'REALIZADA', %(data_realizacao)s, %(data_realizacao)s, %(observacoes)s, %(responsavel_pela_acao)s)
                    RETURNING cod_acompanhamento;
                ''', {
                    'cod_cidadao': cod_cidadao,
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                cod_acompanhamento_realizado = cur.fetchone()[0]
                print(f"[LOG] Nova ação de Avaliar Exames criada - cod_acompanhamento: {cod_acompanhamento_realizado}")

            # Salvar os dados dos resultados dos exames
            print(f"[LOG] Salvando dados dos resultados dos exames na tabela tb_hiperdia_resultados_exames")
            cur.execute("SELECT 1 FROM sistemaaps.tb_hiperdia_resultados_exames WHERE cod_acompanhamento = %s", (cod_acompanhamento_realizado,))
            exists = cur.fetchone()

            if exists:
                sql_upsert_exames = """
                    UPDATE sistemaaps.tb_hiperdia_resultados_exames
                    SET data_avaliacao = %(data_avaliacao)s, colesterol_total = %(colesterol_total)s, hdl = %(hdl)s, 
                        ldl = %(ldl)s, triglicerideos = %(triglicerideos)s, glicemia_jejum = %(glicemia_jejum)s, 
                        hemoglobina_glicada = %(hemoglobina_glicada)s, ureia = %(ureia)s, creatinina = %(creatinina)s, 
                        sodio = %(sodio)s, potassio = %(potassio)s, acido_urico = %(acido_urico)s
                    WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                """
            else:
                sql_upsert_exames = """
                    INSERT INTO sistemaaps.tb_hiperdia_resultados_exames
                    (cod_acompanhamento, data_avaliacao, colesterol_total, hdl, ldl, triglicerideos, glicemia_jejum, hemoglobina_glicada, ureia, creatinina, sodio, potassio, acido_urico)
                    VALUES (%(cod_acompanhamento)s, %(data_avaliacao)s, %(colesterol_total)s, %(hdl)s, %(ldl)s, %(triglicerideos)s, %(glicemia_jejum)s, %(hemoglobina_glicada)s, %(ureia)s, %(creatinina)s, %(sodio)s, %(potassio)s, %(acido_urico)s);
                """
            params = {'cod_acompanhamento': cod_acompanhamento_realizado, 'data_avaliacao': data_realizacao_acao}
            params.update(lab_exam_results)
            cur.execute(sql_upsert_exames, params)
            print(f"[LOG] Dados dos resultados dos exames salvos com sucesso")

            # Criar automaticamente "Avaliar RCV" (cod_acao = 6) como PENDENTE
            print(f"[LOG] Criando ação 'Avaliar RCV' automaticamente")
            cod_proxima_acao_pendente = 6 # Avaliar RCV
            data_agendamento_proxima = data_realizacao_acao # Mesma data (hoje)
            sql_insert_pendente = '''
                INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'Avaliação de Risco Cardiovascular após análise dos exames.', %(responsavel_pela_acao)s);
            '''
            cur.execute(sql_insert_pendente, {
                'cod_cidadao': cod_cidadao,
                'cod_acao': cod_proxima_acao_pendente,
                'data_agendamento': data_agendamento_proxima,
                'cod_acao_origem': cod_acompanhamento_realizado,
                'responsavel_pela_acao': responsavel_pela_acao
            })
            print(f"[LOG] Ação futura 'Avaliar RCV' criada automaticamente")

        elif int(cod_acao_atual) == 1: # Solicitar MRPA
            print(f"[LOG] Processando Solicitar MRPA (cod_acao = 1)")
            
            # 1. Marcar "Agendar Hiperdia" (cod_acao = 9) pendente como REALIZADA
            print(f"[LOG] Marcando ações 'Agendar Hiperdia' pendentes como realizadas")
            sql_update_agendar_hiperdia = """
                UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                SET status_acao = 'REALIZADA',
                    data_agendamento = %(data_realizacao)s,
                    data_realizacao = %(data_realizacao)s,
                    responsavel_pela_acao = %(responsavel_pela_acao)s
                WHERE cod_cidadao = %(cod_cidadao)s
                  AND cod_acao = 9
                  AND status_acao = 'PENDENTE';
            """
            cur.execute(sql_update_agendar_hiperdia, {
                'data_realizacao': data_realizacao_acao,
                'responsavel_pela_acao': responsavel_pela_acao,
                'cod_cidadao': cod_cidadao
            })
            print(f"[LOG] Ações 'Agendar Hiperdia' pendentes marcadas como realizadas")
            
            # 2. Criar "Solicitar MRPA" como REALIZADA com a data de hoje
            print(f"[LOG] Criando ação 'Solicitar MRPA' como REALIZADA")
            sql_insert_solicitar_mrpa = """
                INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
                VALUES (%(cod_cidadao)s, 1, 'REALIZADA', %(data_realizacao)s, %(data_realizacao)s, 'Solicitação de MRPA para monitorização da pressão arterial.', %(responsavel_pela_acao)s)
                RETURNING cod_acompanhamento;
            """
            cur.execute(sql_insert_solicitar_mrpa, {
                'cod_cidadao': cod_cidadao,
                'data_realizacao': data_realizacao_acao,
                'responsavel_pela_acao': responsavel_pela_acao
            })
            cod_acompanhamento_criado = cur.fetchone()[0]
            print(f"[LOG] Ação 'Solicitar MRPA' criada como REALIZADA - cod_acompanhamento: {cod_acompanhamento_criado}")

        else:
            # Criar nova ação para outras ações (tratamento diferenciado para PENDENTE vs REALIZADA)
            if status_acao == 'PENDENTE':
                sql_insert_acompanhamento = """
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, %(cod_acao)s, %(status_acao)s, %(data_agendamento)s, %(observacoes)s, %(responsavel_pela_acao)s)
                    RETURNING cod_acompanhamento;
                """
                print(f"[LOG] Executando INSERT PENDENTE na tabela tb_hiperdia_has_acompanhamento")
                cur.execute(sql_insert_acompanhamento, {
                    'cod_cidadao': cod_cidadao,
                    'cod_acao': cod_acao_atual,
                    'status_acao': status_acao,
                    'data_agendamento': data_realizacao_acao,  # Para PENDENTE, data_realizacao_acao é usada como data_agendamento
                    'observacoes': observacoes,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                cod_acompanhamento_criado = cur.fetchone()[0]
                print(f"[LOG] Nova ação PENDENTE criada - cod_acompanhamento: {cod_acompanhamento_criado}, data_agendamento: {data_realizacao_acao}")
            else:
                sql_insert_acompanhamento = """
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, %(cod_acao)s, %(status_acao)s, %(data_realizacao)s, %(data_realizacao)s, %(observacoes)s, %(responsavel_pela_acao)s)
                    RETURNING cod_acompanhamento;
                """
                print(f"[LOG] Executando INSERT REALIZADA na tabela tb_hiperdia_has_acompanhamento")
                cur.execute(sql_insert_acompanhamento, {
                    'cod_cidadao': cod_cidadao,
                    'cod_acao': cod_acao_atual,
                    'status_acao': status_acao,
                    'data_realizacao': data_realizacao_acao,
                    'observacoes': observacoes,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                cod_acompanhamento_criado = cur.fetchone()[0]
                print(f"[LOG] Nova ação REALIZADA criada - cod_acompanhamento: {cod_acompanhamento_criado}, data_realizacao: {data_realizacao_acao}")

            # Handle specific data based on action type
            # Tratamento especial para "Encaminhar Cardiologia"
            if int(cod_acao_atual) == 10: # Encaminhar Cardiologia
                print(f"[LOG] Processando Encaminhar Cardiologia (cod_acao = 10)")
                
                cardiologia_data = data.get('cardiologia_data', {})
                profissional_responsavel = cardiologia_data.get('profissional_responsavel', responsavel_pela_acao)
                observacoes_cardiologia = cardiologia_data.get('observacoes', observacoes)
                
                # Inserir na tabela específica de cardiologia
                sql_insert_cardiologia = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_cardiologia
                    (cod_acompanhamento, cod_cidadao, tipo_acao, data_acao, profissional_responsavel, observacoes)
                    VALUES (%(cod_acompanhamento)s, %(cod_cidadao)s, %(tipo_acao)s, %(data_acao)s, %(profissional_responsavel)s, %(observacoes)s);
                '''
                cur.execute(sql_insert_cardiologia, {
                    'cod_acompanhamento': cod_acompanhamento_criado,
                    'cod_cidadao': cod_cidadao,
                    'tipo_acao': cod_acao_atual,
                    'data_acao': data_realizacao_acao,
                    'profissional_responsavel': profissional_responsavel,
                    'observacoes': observacoes_cardiologia
                })
                print(f"[LOG] Encaminhamento para cardiologia registrado - cod_acompanhamento: {cod_acompanhamento_criado}")

            # Tratamento especial para "Registrar Cardiologia"
            elif int(cod_acao_atual) == 11: # Registrar Cardiologia
                print(f"[LOG] Processando Registrar Cardiologia (cod_acao = 11)")
                
                cardiologia_data = data.get('cardiologia_data', {})
                profissional_responsavel = cardiologia_data.get('profissional_responsavel', responsavel_pela_acao)
                consulta_cardiologia = cardiologia_data.get('consulta_cardiologia', '')
                recomendacoes_cardiologia = cardiologia_data.get('recomendacoes_cardiologia', '')
                observacoes_cardiologia = cardiologia_data.get('observacoes', observacoes)
                
                if not consulta_cardiologia:
                    print(f"[LOG] ERRO: Consulta cardiológica não fornecida")
                    return jsonify({"sucesso": False, "erro": "Dados da consulta cardiológica são obrigatórios."}), 400
                
                # Inserir na tabela específica de cardiologia
                sql_insert_cardiologia = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_cardiologia
                    (cod_acompanhamento, cod_cidadao, tipo_acao, data_acao, profissional_responsavel, 
                     consulta_cardiologia, recomendacoes_cardiologia, observacoes)
                    VALUES (%(cod_acompanhamento)s, %(cod_cidadao)s, %(tipo_acao)s, %(data_acao)s, %(profissional_responsavel)s, 
                            %(consulta_cardiologia)s, %(recomendacoes_cardiologia)s, %(observacoes)s);
                '''
                cur.execute(sql_insert_cardiologia, {
                    'cod_acompanhamento': cod_acompanhamento_criado,
                    'cod_cidadao': cod_cidadao,
                    'tipo_acao': cod_acao_atual,
                    'data_acao': data_realizacao_acao,
                    'profissional_responsavel': profissional_responsavel,
                    'consulta_cardiologia': consulta_cardiologia,
                    'recomendacoes_cardiologia': recomendacoes_cardiologia,
                    'observacoes': observacoes_cardiologia
                })
                print(f"[LOG] Consulta cardiológica registrada - cod_acompanhamento: {cod_acompanhamento_criado}")
                
                # Após registrar a consulta cardiológica, marcar ação "Encaminhar Cardiologia" (cod_acao = 10) como REALIZADA
                print(f"[LOG] Buscando ação 'Encaminhar Cardiologia' pendente para finalizar...")
                cur.execute('''
                    SELECT cod_acompanhamento 
                    FROM sistemaaps.tb_hiperdia_has_acompanhamento 
                    WHERE cod_cidadao = %(cod_cidadao)s 
                        AND cod_acao = 10 
                        AND status_acao = 'PENDENTE'
                    ORDER BY data_agendamento DESC
                    LIMIT 1;
                ''', {'cod_cidadao': cod_cidadao})
                
                encaminhamento_pendente = cur.fetchone()
                if encaminhamento_pendente:
                    cod_acompanhamento_encaminhamento = encaminhamento_pendente[0]
                    cur.execute('''
                        UPDATE sistemaaps.tb_hiperdia_has_acompanhamento
                        SET status_acao = 'REALIZADA',
                            data_realizacao = %(data_realizacao)s,
                            responsavel_pela_acao = %(responsavel_pela_acao)s
                        WHERE cod_acompanhamento = %(cod_acompanhamento)s;
                    ''', {
                        'data_realizacao': data_realizacao_acao,
                        'responsavel_pela_acao': responsavel_pela_acao,
                        'cod_acompanhamento': cod_acompanhamento_encaminhamento
                    })
                    print(f"[LOG] Ação 'Encaminhar Cardiologia' (cod_acompanhamento: {cod_acompanhamento_encaminhamento}) finalizada automaticamente")
                else:
                    print(f"[LOG] Nenhuma ação 'Encaminhar Cardiologia' pendente encontrada para finalizar")

            # Tratamento especial para "Modificar Tratamento"
            elif int(cod_acao_atual) == 3: # Modificar Tratamento
                print(f"[LOG] Processando Modificar Tratamento (cod_acao = 3)")
                
                # Após registrar "Modificar tratamento", criar automaticamente "Iniciar MRPA" para 30 dias
                print(f"[LOG] Criando ação futura 'Iniciar MRPA' automaticamente para 30 dias após modificação de tratamento")
                cod_proxima_acao_pendente = 1 # Iniciar MRPA
                data_agendamento_proxima = data_realizacao_acao + timedelta(days=30) # 30 dias
                sql_insert_pendente_mrpa = '''
                    INSERT INTO sistemaaps.tb_hiperdia_has_acompanhamento
                    (cod_cidadao, cod_acao, status_acao, data_agendamento, cod_acao_origem, observacoes, responsavel_pela_acao)
                    VALUES (%(cod_cidadao)s, %(cod_acao)s, 'PENDENTE', %(data_agendamento)s, %(cod_acao_origem)s, 'MRPA de controle após modificação de tratamento (30 dias).', %(responsavel_pela_acao)s);
                '''
                cur.execute(sql_insert_pendente_mrpa, {
                    'cod_cidadao': cod_cidadao,
                    'cod_acao': cod_proxima_acao_pendente,
                    'data_agendamento': data_agendamento_proxima,
                    'cod_acao_origem': cod_acompanhamento_criado,
                    'responsavel_pela_acao': responsavel_pela_acao
                })
                print(f"[LOG] Ação futura 'Iniciar MRPA' criada automaticamente para 30 dias ({data_agendamento_proxima})")

        conn.commit()
        print(f"[LOG] Commit realizado com sucesso - Ação {cod_acao_atual} registrada")
        return jsonify({"sucesso": True, "mensagem": "Ação registrada com sucesso!"})

    except Exception as e:
        if conn: conn.rollback()
        print(f"[LOG] ERRO ao registrar ação do Hiperdia: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {str(e)}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/graficos_painel_plafam')
def api_graficos_painel_plafam():
    equipe_req = request.args.get('equipe', 'Todas')
    # Não há filtro de idade aqui!
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        # --- Gráfico de Pizza (Distribuição da Equipe) ---
        base_where_pizza = []
        params_pizza = {}
        if equipe_req != 'Todas':
            base_where_pizza.append("m.nome_equipe = %(equipe)s")
            params_pizza['equipe'] = equipe_req
        where_clause_pizza_str = " WHERE " + " AND ".join(base_where_pizza) if base_where_pizza else ""

        query_pizza_status = f"""
            SELECT 
                SUM(CASE WHEN m.status_gravidez = 'Grávida' THEN 1 ELSE 0 END) as gestantes,
                SUM(CASE WHEN (m.metodo IS NULL OR m.metodo = '') AND (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') THEN 1 ELSE 0 END) as sem_metodo,
                SUM(CASE WHEN 
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days'))
                            ))
                        )
                    THEN 1 ELSE 0 END) as metodo_atraso,
                SUM(CASE WHEN
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                            )) OR
                            (m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%')
                        )
                    THEN 1 ELSE 0 END) as metodo_em_dia
            FROM sistemaaps.mv_plafam m
            {where_clause_pizza_str};
        """
        cur.execute(query_pizza_status, params_pizza)
        dados_pizza = cur.fetchone()

        # --- Gráfico de Barras (Distribuição por Micro-área ou Equipe) ---
        dados_barras_db = []
        if equipe_req != 'Todas':
            base_where_barras_agente = ["m.nome_equipe = %(equipe)s"]
            params_barras_agente = {'equipe': equipe_req}
            where_clause_barras_agente_str = " WHERE " + " AND ".join(base_where_barras_agente)
            query_barras_status_por_agente = f"""
            SELECT 
                m.microarea, ag.nome_agente,
                SUM(CASE WHEN m.status_gravidez = 'Grávida' THEN 1 ELSE 0 END) as gestantes,
                SUM(CASE WHEN (m.metodo IS NULL OR m.metodo = '') AND (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') THEN 1 ELSE 0 END) as sem_metodo,
                SUM(CASE WHEN 
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days'))
                            ))
                        )
                    THEN 1 ELSE 0 END) as metodo_atraso,
                SUM(CASE WHEN
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                            )) OR
                            (m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%')
                        )
                    THEN 1 ELSE 0 END) as metodo_em_dia
            FROM sistemaaps.mv_plafam m
            LEFT JOIN sistemaaps.tb_agentes ag ON m.nome_equipe = ag.nome_equipe AND m.microarea = ag.micro_area
            {where_clause_barras_agente_str}
            GROUP BY m.microarea, ag.nome_agente
            ORDER BY m.microarea, ag.nome_agente;
            """
            cur.execute(query_barras_status_por_agente, params_barras_agente)
            dados_barras_db = cur.fetchall()
        else:
            query_barras_status_por_equipe = f"""
            SELECT 
                m.nome_equipe,
                SUM(CASE WHEN m.status_gravidez = 'Grávida' THEN 1 ELSE 0 END) as gestantes,
                SUM(CASE WHEN (m.metodo IS NULL OR m.metodo = '') AND (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') THEN 1 ELSE 0 END) as sem_metodo,
                SUM(CASE WHEN 
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days'))
                            ))
                        )
                    THEN 1 ELSE 0 END) as metodo_atraso,
                SUM(CASE WHEN
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                            )) OR
                            (m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%')
                        )
                    THEN 1 ELSE 0 END) as metodo_em_dia
            FROM sistemaaps.mv_plafam m
            WHERE m.nome_equipe IS NOT NULL
            GROUP BY m.nome_equipe
            ORDER BY m.nome_equipe;
            """
            cur.execute(query_barras_status_por_equipe)
            dados_barras_db = cur.fetchall()
        dados_barras = []
        if dados_barras_db:
            for row in dados_barras_db:
                dados_barras.append(dict(row))
        return jsonify({
            "pizza_data": dict(dados_pizza) if dados_pizza else {},
            "bar_chart_data": dados_barras
        })
    except Exception as e:
        print(f"Erro ao buscar dados para gráficos do painel plafam: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/plafam/analytics/status_snapshot')
def api_plafam_status_snapshot():
    equipe_req = request.args.get('equipe', 'Todas')
    microarea_req = request.args.get('microarea', 'Todas as áreas')
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        where_clauses = []
        params = {}
        if equipe_req != 'Todas':
            where_clauses.append("m.nome_equipe = %(equipe)s")
            params['equipe'] = equipe_req
        if microarea_req != 'Todas as áreas':
            if ' - ' in microarea_req:
                parts = microarea_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else:
                micro_area_str = microarea_req.replace('Área ', '').strip()
            if micro_area_str:
                where_clauses.append("m.microarea = %(microarea)s")
                params['microarea'] = micro_area_str
        where_str = (" WHERE " + " AND ".join(where_clauses)) if where_clauses else ""

        query = f"""
            SELECT 
                SUM(CASE WHEN m.status_gravidez = 'Grávida' THEN 1 ELSE 0 END) as gestantes,
                SUM(CASE WHEN (m.metodo IS NULL OR m.metodo = '') AND (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') THEN 1 ELSE 0 END) as sem_metodo,
                SUM(CASE WHEN 
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days'))
                            ))
                        )
                    THEN 1 ELSE 0 END) as metodo_atraso,
                SUM(CASE WHEN
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                            )) OR
                            (m.metodo ILIKE '%%diu%%' OR m.metodo ILIKE '%%implante%%' OR m.metodo ILIKE '%%laqueadura%%')
                        )
                    THEN 1 ELSE 0 END) as metodo_em_dia
            FROM sistemaaps.mv_plafam m
            {where_str};
        """
        cur.execute(query, params)
        row = cur.fetchone()
        return jsonify(dict(row) if row else {})
    except Exception as e:
        print(f"Erro ao buscar status_snapshot do plafam: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/plafam/analytics/actions_overview')
def api_plafam_actions_overview():
    equipe_req = request.args.get('equipe', 'Todas')
    microarea_req = request.args.get('microarea', 'Todas as áreas')
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        where_clauses = []
        params = {}
        if equipe_req != 'Todas':
            where_clauses.append("m.nome_equipe = %(equipe)s")
            params['equipe'] = equipe_req
        if microarea_req != 'Todas as áreas':
            if ' - ' in microarea_req:
                parts = microarea_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else:
                micro_area_str = microarea_req.replace('Área ', '').strip()
            if micro_area_str:
                where_clauses.append("m.microarea = %(microarea)s")
                params['microarea'] = micro_area_str
        where_str = (" WHERE " + " AND ".join(where_clauses)) if where_clauses else ""

        query = f"""
        WITH base AS (
            SELECT m.cod_paciente
            FROM sistemaaps.mv_plafam m
            {where_str}
        ), joined AS (
            SELECT b.cod_paciente, COALESCE(pa.status_acompanhamento, 0) AS status
            FROM base b
            LEFT JOIN sistemaaps.tb_plafam_acompanhamento pa ON pa.co_cidadao = b.cod_paciente
        )
        SELECT status, COUNT(*) AS total
        FROM joined
        GROUP BY status
        ORDER BY status;
        """
        cur.execute(query, params)
        rows = cur.fetchall() or []
        data = {str(r['status']): int(r['total']) for r in rows}
        return jsonify({"counts": data})
    except Exception as e:
        print(f"Erro ao buscar actions_overview do plafam: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/plafam/analytics/actions_timeseries')
def api_plafam_actions_timeseries():
    equipe_req = request.args.get('equipe', 'Todas')
    microarea_req = request.args.get('microarea', 'Todas as áreas')
    granularity = request.args.get('granularity', 'month')  # 'day'|'week'|'month'
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        trunc_expr = {
            'day': "DATE(pa.data_acompanhamento)",
            'week': "DATE_TRUNC('week', pa.data_acompanhamento)",
            'month': "DATE_TRUNC('month', pa.data_acompanhamento)"
        }.get(granularity, "DATE_TRUNC('month', pa.data_acompanhamento)")

        where_clauses = ["pa.data_acompanhamento IS NOT NULL"]
        params = {}
        if equipe_req != 'Todas':
            where_clauses.append("m.nome_equipe = %(equipe)s")
            params['equipe'] = equipe_req
        if microarea_req != 'Todas as áreas':
            if ' - ' in microarea_req:
                parts = microarea_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else:
                micro_area_str = microarea_req.replace('Área ', '').strip()
            if micro_area_str:
                where_clauses.append("m.microarea = %(microarea)s")
                params['microarea'] = micro_area_str
        where_str = (" WHERE " + " AND ".join(where_clauses)) if where_clauses else ""

        query = f"""
        SELECT {trunc_expr} AS periodo, COALESCE(pa.status_acompanhamento, 0) AS status, COUNT(*) AS total
        FROM sistemaaps.mv_plafam m
        JOIN sistemaaps.tb_plafam_acompanhamento pa ON pa.co_cidadao = m.cod_paciente
        {where_str}
        GROUP BY periodo, status
        ORDER BY periodo ASC, status ASC;
        """
        cur.execute(query, params)
        rows = cur.fetchall() or []
        series = {}
        for r in rows:
            status_key = str(r['status'])
            periodo_val = r['periodo']
            if isinstance(periodo_val, (datetime, date)):
                periodo_str = periodo_val.strftime('%Y-%m-%d')
            else:
                periodo_str = str(periodo_val)
            series.setdefault(status_key, []).append({"period": periodo_str, "count": int(r['total'])})
        return jsonify({"series": series, "granularity": granularity})
    except Exception as e:
        print(f"Erro ao buscar actions_timeseries do plafam: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/plafam/analytics/method_mix')
def api_plafam_method_mix():
    equipe_req = request.args.get('equipe', 'Todas')
    microarea_req = request.args.get('microarea', 'Todas as áreas')
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        where_clauses = []
        params = {}
        if equipe_req != 'Todas':
            where_clauses.append("m.nome_equipe = %(equipe)s")
            params['equipe'] = equipe_req
        if microarea_req != 'Todas as áreas':
            if ' - ' in microarea_req:
                parts = microarea_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else:
                micro_area_str = microarea_req.replace('Área ', '').strip()
            if micro_area_str:
                where_clauses.append("m.microarea = %(microarea)s")
                params['microarea'] = micro_area_str
        where_str = (" WHERE " + " AND ".join(where_clauses)) if where_clauses else ""

        query = f"""
        SELECT
            CASE
                WHEN m.metodo ILIKE '%%laqueadura%%' THEN 'LAQUEADURA'
                WHEN m.metodo ILIKE '%%histerectomia%%' THEN 'HISTERECTOMIA'
                WHEN m.metodo ILIKE '%%diu%%' THEN 'DIU'
                WHEN m.metodo ILIKE '%%implante%%' THEN 'IMPLANTE SUBDÉRMICO'
                WHEN m.metodo ILIKE '%%mensal%%' THEN 'MENSAL'
                WHEN m.metodo ILIKE '%%trimestral%%' THEN 'TRIMESTRAL'
                WHEN m.metodo ILIKE '%%pílula%%' OR m.metodo ILIKE '%%p\u00edlula%%' OR m.metodo ILIKE '%%pilula%%' THEN 'PÍLULA'
                WHEN m.metodo IS NULL OR m.metodo = '' THEN 'SEM MÉTODO'
                ELSE 'OUTROS'
            END AS categoria,
            COUNT(*) AS total
        FROM sistemaaps.mv_plafam m
        {where_str}
        GROUP BY categoria
        ORDER BY total DESC;
        """
        cur.execute(query, params)
        rows = cur.fetchall() or []
        data = [{"categoria": r['categoria'], "total": int(r['total'])} for r in rows]
        return jsonify({"mix": data})
    except Exception as e:
        print(f"Erro ao buscar method_mix do plafam: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/indicadores_plafam')
def api_indicadores_plafam():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # 1. Mulheres de 19 a 45 anos em uso de métodos seguros
        query_mulheres_metodo_seguro = """
            SELECT COUNT(DISTINCT m.cod_paciente)
        FROM sistemaaps.mv_plafam m
            WHERE m.idade_calculada BETWEEN 19 AND 45
              AND (
                m.metodo ILIKE '%diu%' OR m.metodo ILIKE '%implante%' OR m.metodo ILIKE '%laqueadura%' OR
                m.metodo ILIKE '%injet%' OR m.metodo ILIKE '%pílula%' OR m.metodo ILIKE '%oral%'
              )
              AND (
                (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                    ((m.metodo ILIKE '%mensal%' OR m.metodo ILIKE '%pílula%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                    (m.metodo ILIKE '%trimestral%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                ))
                OR (m.metodo ILIKE '%diu%' OR m.metodo ILIKE '%implante%' OR m.metodo ILIKE '%laqueadura%')
              )
        """
        cur.execute(query_mulheres_metodo_seguro)
        mulheres_19_45_metodo_seguro = cur.fetchone()[0] or 0

        # 2. Adolescentes de 14 a 18 anos em uso regular de métodos
        query_adolescentes_metodo_regular = """
            SELECT COUNT(DISTINCT m.cod_paciente)
            FROM sistemaaps.mv_plafam m
            WHERE m.idade_calculada BETWEEN 14 AND 18
              AND m.metodo IS NOT NULL AND m.metodo != ''
              AND (
                (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                    ((m.metodo ILIKE '%mensal%' OR m.metodo ILIKE '%pílula%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                    (m.metodo ILIKE '%trimestral%' AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days'))
                ))
                OR (m.metodo ILIKE '%diu%' OR m.metodo ILIKE '%implante%' OR m.metodo ILIKE '%laqueadura%')
              )
        """
        cur.execute(query_adolescentes_metodo_regular)
        adolescentes_14_18_metodo_regular = cur.fetchone()[0] or 0

        # 3. Número de gestantes adolescentes (14 a 18 anos)
        query_gestantes_adolescentes = """
            SELECT COUNT(DISTINCT m.cod_paciente)
            FROM sistemaaps.mv_plafam m
            WHERE m.idade_calculada BETWEEN 14 AND 18
              AND m.status_gravidez = 'Grávida'
        """
        cur.execute(query_gestantes_adolescentes)
        gestantes_adolescentes_14_18 = cur.fetchone()[0] or 0

        return jsonify({
            'mulheres_19_45_metodo_seguro': mulheres_19_45_metodo_seguro,
            'adolescentes_14_18_metodo_regular': adolescentes_14_18_metodo_regular,
            'gestantes_adolescentes_14_18': gestantes_adolescentes_14_18
        })
    except Exception as e:
        print(f"Erro ao buscar indicadores do plafam: {e}")
        return jsonify({'erro': f'Erro no servidor: {e}'}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/equipes_com_agentes_plafam')
def api_equipes_com_agentes_plafam():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        # Pega todas as equipes distintas de mv_plafam ou tb_agentes
        cur.execute('''
            SELECT DISTINCT nome_equipe 
            FROM (
                SELECT nome_equipe FROM sistemaaps.mv_plafam
                UNION
                SELECT nome_equipe FROM sistemaaps.tb_agentes
            ) AS equipes_unidas
            WHERE nome_equipe IS NOT NULL 
            ORDER BY nome_equipe
        ''')
        equipes_db = cur.fetchall()
        resultado_final = []
        for eq_row in equipes_db:
            equipe_nome = eq_row['nome_equipe']
            # Contar mulheres 14-45 anos para a equipe atual
            cur.execute('''
                SELECT COUNT(DISTINCT m.cod_paciente)
                FROM sistemaaps.mv_plafam m
                WHERE m.nome_equipe = %s AND m.idade_calculada BETWEEN 14 AND 45
            ''', (equipe_nome,))
            num_mulheres = cur.fetchone()[0] or 0
            cur.execute('''
                SELECT micro_area, nome_agente 
                FROM sistemaaps.tb_agentes 
                WHERE nome_equipe = %s 
                ORDER BY micro_area, nome_agente
            ''', (equipe_nome,))
            agentes_db = cur.fetchall()
            agentes = []
            if agentes_db:
                agentes = [{"micro_area": ag['micro_area'], "nome_agente": ag['nome_agente']} for ag in agentes_db]
            resultado_final.append({"nome_equipe": equipe_nome, "agentes": agentes, "num_mulheres": num_mulheres})
        resultado_final_ordenado = sorted(resultado_final, key=lambda x: x.get('num_mulheres', 0), reverse=True)
        return jsonify(resultado_final_ordenado)
    except Exception as e:
        print(f"Erro ao buscar equipes com agentes plafam: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/medicamentos_atuais/<int:cod_cidadao>')
def api_get_medicamentos_atuais_hiperdia(cod_cidadao):
    """Busca medicamentos atuais de um paciente combinando view existente e tabela manual"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # Busca medicamentos da tabela manual
        sql_manual = """
            SELECT 
                cod_seq_medicamentos as cod_seq_medicamento,
                'manual' as origem,
                codcidadao, 
                nome_medicamento, 
                CASE 
                    WHEN COALESCE(dose, 1) = 1 THEN '1 comprimido'
                    WHEN COALESCE(dose, 1) = 2 THEN '2 comprimidos'
                    WHEN COALESCE(dose, 1) = 3 THEN '3 comprimidos'
                    WHEN COALESCE(dose, 1) = 4 THEN '4 comprimidos'
                    WHEN COALESCE(dose, 1) = 5 THEN '5 comprimidos'
                    ELSE COALESCE(dose, 1)::text || ' comprimidos'
                END as dose_texto,
                COALESCE(dose, 1) as dose,
                CASE 
                    WHEN COALESCE(frequencia, 1) = 1 THEN '1x ao dia'
                    WHEN COALESCE(frequencia, 1) = 2 THEN '2x ao dia'
                    WHEN COALESCE(frequencia, 1) = 3 THEN '3x ao dia'
                    WHEN COALESCE(frequencia, 1) = 4 THEN '4x ao dia'
                    ELSE COALESCE(frequencia, 1)::text || 'x ao dia'
                END as frequencia_texto,
                COALESCE(frequencia, 1) as frequencia,
                data_inicio,
                data_fim,
                null as motivo_interrupcao,
                observacao as observacoes
            FROM 
                sistemaaps.tb_hiperdia_has_medicamentos
            WHERE 
                codcidadao = %(cod_cidadao)s
                AND (data_fim IS NULL OR data_fim > CURRENT_DATE)
        """
        
        medicamentos = []
        
        # Busca apenas da tabela manual
        try:
            cur.execute(sql_manual, {'cod_cidadao': cod_cidadao})
            medicamentos_manual = cur.fetchall()
            
            for row in medicamentos_manual:
                med_dict = dict(row)
                if med_dict.get('data_inicio') and isinstance(med_dict['data_inicio'], date):
                    med_dict['data_inicio'] = med_dict['data_inicio'].strftime('%Y-%m-%d')
                if med_dict.get('data_fim') and isinstance(med_dict['data_fim'], date):
                    med_dict['data_fim'] = med_dict['data_fim'].strftime('%Y-%m-%d')
                medicamentos.append(med_dict)
        except Exception as e:
            print(f"Erro ao buscar medicamentos manuais: {e}")
            # Se houver erro, retorna lista vazia
        
        # Ordena por data de início (mais recente primeiro), tratando valores None
        medicamentos.sort(key=lambda x: x.get('data_inicio') or '', reverse=True)
            
        return jsonify(medicamentos)

    except Exception as e:
        print(f"Erro na API /api/hiperdia/medicamentos_atuais: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/medicamentos', methods=['POST'])
def api_adicionar_medicamento_hiperdia():
    """Adiciona um novo medicamento para um paciente"""
    data = request.get_json()
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        cod_cidadao = data.get('codcidadao')
        nome_medicamento = data.get('nome_medicamento')
        dose = data.get('dose')
        frequencia = data.get('frequencia')
        data_inicio = data.get('data_inicio')
        observacao = data.get('observacoes', '')  # Mapear observacoes para observacao

        if not all([cod_cidadao, nome_medicamento, dose, frequencia]):
            return jsonify({"sucesso": False, "erro": "Dados incompletos para adicionar medicamento. Campos obrigatórios: nome, dose e frequência."}), 400

        # Converter dose e frequência para inteiros
        try:
            dose_num = int(dose)
            frequencia_num = int(frequencia)
        except (ValueError, TypeError):
            return jsonify({"sucesso": False, "erro": "Dose e frequência devem ser números válidos."}), 400

        # Se data_inicio não foi fornecida ou está vazia, usar data atual
        if not data_inicio or data_inicio.strip() == '':
            from datetime import date
            data_inicio = date.today()

        sql_insert = """
            INSERT INTO sistemaaps.tb_hiperdia_has_medicamentos
            (codcidadao, nome_medicamento, dose, frequencia, data_inicio, observacao)
            VALUES (%(codcidadao)s, %(nome_medicamento)s, %(dose)s, %(frequencia)s, %(data_inicio)s, %(observacao)s)
            RETURNING cod_seq_medicamentos;
        """

        cur.execute(sql_insert, {
            'codcidadao': cod_cidadao,
            'nome_medicamento': nome_medicamento,
            'dose': dose_num,
            'frequencia': frequencia_num,
            'data_inicio': data_inicio,
            'observacao': observacao
        })

        cod_seq_medicamento = cur.fetchone()[0]
        conn.commit()

        return jsonify({
            "sucesso": True, 
            "mensagem": "Medicamento adicionado com sucesso",
            "cod_seq_medicamento": cod_seq_medicamento
        })

    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro na API adicionar medicamento: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/medicamentos/<int:cod_seq_medicamento>', methods=['PUT'])
def api_atualizar_medicamento_hiperdia(cod_seq_medicamento):
    """Atualiza um medicamento existente"""
    data = request.get_json()
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Campos que podem ser atualizados
        nome_medicamento = data.get('nome_medicamento')
        dose = data.get('dose')
        frequencia = data.get('frequencia')
        observacao = data.get('observacoes')  # Mapear para observacao

        # Construir query dinâmica baseada nos campos fornecidos
        update_fields = []
        params = {'cod_seq_medicamentos': cod_seq_medicamento}

        if nome_medicamento is not None:
            update_fields.append("nome_medicamento = %(nome_medicamento)s")
            params['nome_medicamento'] = nome_medicamento

        if dose is not None:
            try:
                dose_num = int(dose)
                update_fields.append("dose = %(dose)s")
                params['dose'] = dose_num
            except (ValueError, TypeError):
                return jsonify({"sucesso": False, "erro": "Dose deve ser um número válido."}), 400

        if frequencia is not None:
            try:
                frequencia_num = int(frequencia)
                update_fields.append("frequencia = %(frequencia)s")
                params['frequencia'] = frequencia_num
            except (ValueError, TypeError):
                return jsonify({"sucesso": False, "erro": "Frequência deve ser um número válido."}), 400

        if observacao is not None:
            update_fields.append("observacao = %(observacao)s")
            params['observacao'] = observacao

        if not update_fields:
            return jsonify({"sucesso": False, "erro": "Nenhum campo para atualizar fornecido."}), 400

        sql_update = f"""
            UPDATE sistemaaps.tb_hiperdia_has_medicamentos
            SET {', '.join(update_fields)}
            WHERE cod_seq_medicamentos = %(cod_seq_medicamentos)s;
        """

        cur.execute(sql_update, params)
        
        if cur.rowcount == 0:
            return jsonify({"sucesso": False, "erro": "Medicamento não encontrado."}), 404

        conn.commit()

        return jsonify({
            "sucesso": True, 
            "mensagem": "Medicamento atualizado com sucesso"
        })

    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro na API atualizar medicamento: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/medicamentos/<int:cod_seq_medicamento>/interromper', methods=['PUT'])
def api_interromper_medicamento_hiperdia(cod_seq_medicamento):
    """Interrompe um medicamento definindo data_fim e motivo"""
    data = request.get_json()
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        data_fim = data.get('data_fim')
        motivo = data.get('motivo', 'Interrompido pelo profissional')

        if not data_fim:
            return jsonify({"sucesso": False, "erro": "Data de fim é obrigatória para interromper medicamento."}), 400

        sql_update = """
            UPDATE sistemaaps.tb_hiperdia_has_medicamentos
            SET data_fim = %(data_fim)s
            WHERE cod_seq_medicamentos = %(cod_seq_medicamentos)s;
        """

        cur.execute(sql_update, {
            'data_fim': data_fim,
            'cod_seq_medicamentos': cod_seq_medicamento
        })

        if cur.rowcount == 0:
            return jsonify({"sucesso": False, "erro": "Medicamento não encontrado."}), 404

        conn.commit()

        return jsonify({
            "sucesso": True, 
            "mensagem": "Medicamento interrompido com sucesso"
        })

    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro na API interromper medicamento: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/export_plafam')
def api_export_plafam():
    """API para exportação completa de dados do Plafam sem limitação de paginação"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Usar a mesma lógica de filtros da API principal, mas sem paginação
        where_clause, order_by_clause, query_params = build_filtered_query(request.args)
        
        base_query = """
        SELECT
            m.cod_paciente, m.nome_paciente, m.cartao_sus, m.idade_calculada, m.microarea,
            m.metodo, m.nome_equipe, m.data_aplicacao, m.status_gravidez, m.data_provavel_parto,
            pa.status_acompanhamento, pa.data_acompanhamento,
            ag.nome_agente
        FROM sistemaaps.mv_plafam m
        LEFT JOIN sistemaaps.tb_plafam_acompanhamento pa ON m.cod_paciente = pa.co_cidadao
        LEFT JOIN sistemaaps.tb_agentes ag ON m.microarea = ag.micro_area AND m.nome_equipe = ag.nome_equipe
        """
        
        # Construir query final sem limit e offset
        final_query = base_query + where_clause + " " + order_by_clause
        
        print(f"DEBUG Export Plafam Query: {final_query}")
        print(f"DEBUG Export Plafam Params: {query_params}")
        
        cur.execute(final_query, query_params)
        pacientes = cur.fetchall()
        
        # Converter para formato JSON
        pacientes_list = []
        for pac in pacientes:
            pac_dict = {
                'cod_paciente': pac[0],
                'nome_paciente': pac[1],
                'cartao_sus': pac[2],
                'idade_calculada': pac[3],
                'microarea': pac[4],
                'metodo': pac[5],
                'nome_equipe': pac[6],
                'data_aplicacao': pac[7],
                'status_gravidez': pac[8],
                'data_provavel_parto': pac[9],
                'status_acompanhamento': pac[10],
                'data_acompanhamento': pac[11],
                'nome_agente': pac[12]
            }
            pacientes_list.append(pac_dict)
        
        print(f"DEBUG Export Plafam: Retornando {len(pacientes_list)} registros")
        return jsonify(pacientes_list)
        
    except Exception as e:
        print(f"Erro na API export_plafam: {e}")
        return jsonify({'erro': f'Erro no servidor: {e}'}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/plano_semanal_plafam')
def api_plano_semanal_plafam():
    """API para gerar Plano Semanal - seleciona 2 mulheres por microárea (>=19 anos, sem método, sem acompanhamento) ordenadas por idade (mais novas primeiro)"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Parâmetro opcional para equipe específica
        equipe_selecionada = request.args.get('equipe_selecionada', None)
        
        # Construir filtros
        where_clauses = []
        query_params = {}
        
        # Se há equipe específica, usar apenas ela
        if equipe_selecionada and equipe_selecionada != 'Todas':
            where_clauses.append("m.nome_equipe = %(equipe_selecionada)s")
            query_params['equipe_selecionada'] = equipe_selecionada
        
        # Filtros fixos do Plano Semanal
        where_clauses.extend([
            "m.microarea NOT IN (0, 7, 8)",  # Excluir microáreas específicas
            "(m.metodo = '' OR m.metodo IS NULL)",  # Sem método
            "m.status_gravidez != 'Grávida'",  # Não grávidas
            "m.idade_calculada >= 19",  # Apenas mulheres com 19 anos ou mais
            "(pa.status_acompanhamento IS NULL OR pa.status_acompanhamento = 0)"  # Sem acompanhamento
        ])
        
        where_clause = "WHERE " + " AND ".join(where_clauses) if where_clauses else ""
        
        # Query para buscar mulheres elegíveis, limitando 2 por microárea
        query = f"""
        WITH MulheresElegiveis AS (
            SELECT
                m.cod_paciente, m.nome_paciente, m.cartao_sus, m.idade_calculada, 
                m.microarea, m.metodo, m.nome_equipe, m.data_aplicacao, 
                m.status_gravidez, m.data_provavel_parto,
                pa.status_acompanhamento, pa.data_acompanhamento,
                ag.nome_agente,
                ROW_NUMBER() OVER (PARTITION BY m.nome_equipe, m.microarea ORDER BY m.idade_calculada ASC, m.nome_paciente) as rn
            FROM sistemaaps.mv_plafam m
            LEFT JOIN sistemaaps.tb_plafam_acompanhamento pa ON m.cod_paciente = pa.co_cidadao
            LEFT JOIN sistemaaps.tb_agentes ag ON m.microarea = ag.micro_area AND m.nome_equipe = ag.nome_equipe
            {where_clause}
        )
        SELECT 
            cod_paciente, nome_paciente, cartao_sus, idade_calculada, 
            microarea, metodo, nome_equipe, data_aplicacao, 
            status_gravidez, data_provavel_parto,
            status_acompanhamento, data_acompanhamento, nome_agente
        FROM MulheresElegiveis
        WHERE rn <= 2
        ORDER BY nome_equipe, microarea, idade_calculada ASC, nome_paciente
        """
        
        print(f"DEBUG Plano Semanal Query: {query}")
        print(f"DEBUG Plano Semanal Params: {query_params}")
        
        cur.execute(query, query_params)
        pacientes_selecionadas = cur.fetchall()
        
        # Se não há pacientes elegíveis
        if not pacientes_selecionadas:
            return jsonify({
                'pacientes': [],
                'total': 0,
                'mensagem': 'Nenhuma paciente elegível encontrada para o Plano Semanal.'
            })
        
        # Inserir registros de acompanhamento para as pacientes selecionadas
        pacientes_inseridos = []
        resumo_equipes = set()
        total_microareas = set()
        
        for pac in pacientes_selecionadas:
            cod_paciente = pac[0]
            
            try:
                # Inserir registro de acompanhamento com status "Convite com o agente" (status = 1)
                sql_insert_acompanhamento = """
                    INSERT INTO sistemaaps.tb_plafam_acompanhamento (co_cidadao, status_acompanhamento, data_acompanhamento)
                    VALUES (%(co_cidadao)s, 1, CURRENT_DATE)
                    ON CONFLICT (co_cidadao) DO UPDATE
                    SET status_acompanhamento = EXCLUDED.status_acompanhamento,
                        data_acompanhamento = EXCLUDED.data_acompanhamento;
                """
                
                cur.execute(sql_insert_acompanhamento, {'co_cidadao': cod_paciente})
                print(f"DEBUG: Inserido acompanhamento para paciente {cod_paciente}")
                
                # Adicionar à lista de pacientes para PDF
                pac_dict = {
                    'cod_paciente': pac[0],
                    'nome_paciente': pac[1],
                    'cartao_sus': pac[2],
                    'idade_calculada': pac[3],
                    'microarea': pac[4],
                    'metodo': pac[5] or '',
                    'nome_equipe': pac[6],
                    'data_aplicacao': pac[7],
                    'status_gravidez': pac[8],
                    'data_provavel_parto': pac[9],
                    'status_acompanhamento': 1,  # Agora com status "Convite com o agente"
                    'data_acompanhamento': datetime.now().date().strftime('%d/%m/%Y'),
                    'nome_agente': pac[12] or 'A definir'
                }
                pacientes_inseridos.append(pac_dict)
                resumo_equipes.add(pac[6])
                total_microareas.add(f"{pac[6]}-{pac[4]}")
                
            except Exception as e_insert:
                print(f"Erro ao inserir acompanhamento para paciente {cod_paciente}: {e_insert}")
                # Continua com os outros pacientes mesmo se um falhar
        
        # Commit das inserções
        conn.commit()
        
        # Preparar resposta
        resumo_equipes_str = ', '.join(sorted(resumo_equipes)) if resumo_equipes else 'Nenhuma'
        
        response_data = {
            'pacientes': pacientes_inseridos,
            'total': len(pacientes_inseridos),
            'resumo_equipes': resumo_equipes_str,
            'total_microareas': len(total_microareas),
            'mensagem': f'Plano Semanal gerado com sucesso! {len(pacientes_inseridos)} convites selecionados.'
        }
        
        print(f"DEBUG Plano Semanal: Retornando {len(pacientes_inseridos)} pacientes")
        return jsonify(response_data)
        
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro na API plano_semanal_plafam: {e}")
        return jsonify({'erro': f'Erro no servidor: {e}'}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/debug_table_structure')
def api_debug_table_structure():
    """Debug: Verifica a estrutura da tabela tb_hiperdia_has_medicamentos"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # Verificar se a tabela existe e suas colunas
        sql_query = """
            SELECT column_name, data_type 
            FROM information_schema.columns 
            WHERE table_schema = 'sistemaaps' 
            AND table_name = 'tb_hiperdia_has_medicamentos'
            ORDER BY ordinal_position;
        """
        
        cur.execute(sql_query)
        columns = cur.fetchall()
        
        return jsonify({
            "tabela": "tb_hiperdia_has_medicamentos",
            "colunas": [dict(row) for row in columns]
        })

    except Exception as e:
        print(f"Erro ao verificar estrutura da tabela: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/medicamentos_hipertensao')
def api_get_medicamentos_hipertensao():
    """Busca medicamentos anti-hipertensivos da tabela tb_medicamento"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        sql_query = """
            SELECT 
                co_seq_medicamento,
                no_principio_ativo_filtro as nome_medicamento
            FROM 
                public.tb_medicamento
            WHERE 
                co_seq_medicamento BETWEEN 195 AND 196 OR co_seq_medicamento BETWEEN 222 AND 224 OR
                co_seq_medicamento BETWEEN 276 AND 288 OR co_seq_medicamento BETWEEN 317 AND 323 OR
                co_seq_medicamento BETWEEN 372 AND 374 OR co_seq_medicamento BETWEEN 429 AND 431 OR
                co_seq_medicamento BETWEEN 454 AND 457 OR co_seq_medicamento BETWEEN 471 AND 472 OR
                co_seq_medicamento BETWEEN 542 AND 547 OR co_seq_medicamento BETWEEN 551 AND 554 OR
                co_seq_medicamento BETWEEN 600 AND 603 OR co_seq_medicamento BETWEEN 700 AND 702 OR
                co_seq_medicamento BETWEEN 787 AND 790 OR co_seq_medicamento BETWEEN 848 AND 850 OR
                co_seq_medicamento BETWEEN 1027 AND 1035 OR co_seq_medicamento BETWEEN 1114 AND 1118 OR
                co_seq_medicamento BETWEEN 1169 AND 1172 OR co_seq_medicamento BETWEEN 1228 AND 1229 OR
                co_seq_medicamento BETWEEN 1362 AND 1367 OR co_seq_medicamento BETWEEN 1479 AND 1484 OR
                co_seq_medicamento BETWEEN 1572 AND 1573 OR co_seq_medicamento BETWEEN 1614 AND 1617 OR
                co_seq_medicamento BETWEEN 1682 AND 1683 OR co_seq_medicamento BETWEEN 1768 AND 1769 OR
                co_seq_medicamento BETWEEN 1785 AND 1790 OR co_seq_medicamento BETWEEN 1867 AND 1868 OR
                co_seq_medicamento BETWEEN 1890 AND 1897 OR co_seq_medicamento BETWEEN 2111 AND 2118 OR
                co_seq_medicamento BETWEEN 2386 AND 2391 OR co_seq_medicamento BETWEEN 2411 AND 2417 OR
                co_seq_medicamento BETWEEN 2670 AND 2674 OR co_seq_medicamento BETWEEN 2723 AND 2724 OR
                co_seq_medicamento BETWEEN 2783 AND 2784 OR co_seq_medicamento BETWEEN 2846 AND 2862 OR
                co_seq_medicamento BETWEEN 2889 AND 2892 OR co_seq_medicamento BETWEEN 2937 AND 2938 OR
                co_seq_medicamento BETWEEN 3004 AND 3006 OR co_seq_medicamento BETWEEN 3156 AND 3157 OR
                co_seq_medicamento BETWEEN 3249 AND 3250 OR co_seq_medicamento BETWEEN 3268 AND 3269 OR
                co_seq_medicamento IN (199, 447, 516, 1085, 1165, 1668, 1688, 1803, 2258, 2934, 2942, 2957, 2987, 3068, 3179, 3218, 3301, 3343)
            ORDER BY 
                no_principio_ativo_filtro ASC;
        """
        
        cur.execute(sql_query)
        medicamentos_db = cur.fetchall()
        
        medicamentos = []
        for row in medicamentos_db:
            med_dict = dict(row)
            medicamentos.append(med_dict)
            
        return jsonify(medicamentos)

    except Exception as e:
        print(f"Erro na API /api/hiperdia/medicamentos_hipertensao: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/generate_prescriptions_pdf_old', methods=['POST'])
def api_generate_prescriptions_pdf_old():
    """Versão antiga usando ReportLab - mantida como backup"""
    # Código antigo mantido...
    pass

@app.route('/api/hiperdia/generate_prescriptions_pdf', methods=['POST'])
def api_generate_prescriptions_pdf():
    """Gera PDF com receituários usando template Word - versão aprimorada"""
    import unicodedata
    import tempfile
    import os
    from datetime import datetime
    import calendar
    
    
    data = request.get_json()
    patients = data.get('patients', [])
    
    if not patients:
        return jsonify({"erro": "Nenhum paciente selecionado"}), 400
    
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # Caminho para o template dinâmico
        template_path = os.path.join(os.path.dirname(__file__), 'modelos', 'template_receituario_dynamic.docx')
        
        print(f"DEBUG: Template path: {template_path}")
        print(f"DEBUG: Template exists: {os.path.exists(template_path)}")
        if os.path.exists(template_path):
            print(f"DEBUG: Template size: {os.path.getsize(template_path)} bytes")
        
        # Criar diretório temporário
        with tempfile.TemporaryDirectory() as temp_dir:
            print(f"DEBUG: Temp directory: {temp_dir}")
            output_files = []
            
            for i, patient in enumerate(patients):
                # Buscar dados do paciente
                sql_paciente = """
                    SELECT 
                        nome_paciente,
                        cartao_sus,
                        dt_nascimento,
                        sexo
                    FROM sistemaaps.mv_hiperdia_hipertensao
                    WHERE cod_paciente = %(cod_paciente)s
                """
                
                cur.execute(sql_paciente, {'cod_paciente': patient['cod_paciente']})
                paciente_dados = cur.fetchone()
                
                if not paciente_dados:
                    continue
                    
                paciente_dict = dict(paciente_dados)
                
                # Buscar medicamentos
                sql_medicamentos = """
                    SELECT 
                        nome_medicamento,
                        dose,
                        frequencia,
                        data_inicio,
                        observacao,
                        updated_at,
                        created_at
                    FROM sistemaaps.tb_hiperdia_has_medicamentos
                    WHERE codcidadao = %(cod_paciente)s
                    AND (data_fim IS NULL OR data_fim > CURRENT_DATE)
                    ORDER BY created_at DESC, nome_medicamento
                """
                
                cur.execute(sql_medicamentos, {'cod_paciente': patient['cod_paciente']})
                medicamentos = cur.fetchall()
                
                if not medicamentos:
                    continue
                
                # Preparar dados para o template
                hoje = datetime.now()
                
                # Preparar lista de medicamentos
                medicamentos_lista = []
                for idx, med in enumerate(medicamentos, 1):
                    med_dict = dict(med)
                    nome = med_dict['nome_medicamento'].upper()
                    dose = med_dict['dose']
                    freq = med_dict['frequencia']
                    total_comprimidos = dose * freq * 30
                    
                    # Preparar instruções - um horário por linha
                    instrucoes = []
                    if freq == 1:
                        instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 06:00 horas")
                    elif freq == 2:
                        instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 06:00 horas")
                        instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 18:00 horas")
                    elif freq == 3:
                        instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 06:00 horas")
                        instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 14:00 horas")
                        instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 22:00 horas")
                    elif freq == 4:
                        instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 06:00 horas")
                        instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 12:00 horas")
                        instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 18:00 horas")
                        instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 24:00 horas")
                    
                    medicamentos_lista.append({
                        'numero': idx,
                        'nome': nome,
                        'quantidade': total_comprimidos,
                        'instrucoes': instrucoes
                    })
                
                # Calcular tamanho da fonte baseado na quantidade de medicamentos
                num_medicamentos = len(medicamentos_lista)
                
                # Definir tamanho da fonte baseado na quantidade de medicamentos
                if num_medicamentos <= 2:  # 1-2 medicamentos
                    font_size = 14
                elif num_medicamentos == 3:  # 3 medicamentos
                    font_size = 14
                elif num_medicamentos == 4:  # 4 medicamentos
                    font_size = 12
                elif num_medicamentos == 5:  # 5 medicamentos
                    font_size = 12
                elif num_medicamentos == 6:  # 6 medicamentos
                    font_size = 11
                elif num_medicamentos == 7:  # 7 medicamentos
                    font_size = 10
                else:  # 8+ medicamentos (diminui 1pt para cada aumento)
                    font_size = max(6, 16 - num_medicamentos)  # 8pt para 8 med, 7pt para 9 med, etc. Mínimo 6pt
                
                print(f"DEBUG: {num_medicamentos} medicamentos, fonte {font_size}pt")
                
                # Definir quantidade de traços baseada no tamanho da fonte
                if num_medicamentos <= 3:  # Fonte maior (14pt)
                    tracos = "--------------------------------"  # 32 traços
                else:  # Fonte menor (12pt, 11pt, 10pt, 8pt, 7pt, etc)
                    tracos = "---------------------------------------------"  # 45 traços
                
                # Gerar texto completo de medicamentos dinamicamente com espaçamento otimizado
                medicamentos_texto = ""
                for idx, med in enumerate(medicamentos_lista, 1):
                    # Nome e quantidade do medicamento (em negrito no template) - traços ajustados
                    medicamentos_texto += f"{idx}) {med['nome']} {tracos} {med['quantidade']} comprimidos\n"
                    
                    # Adicionar todas as instruções
                    for instrucao in med['instrucoes']:
                        medicamentos_texto += f"{instrucao}\n"
                    
                    # Espaçamento mínimo entre medicamentos (apenas se não for o último)
                    if idx < len(medicamentos_lista):
                        medicamentos_texto += "\n"
                
                # Se não há medicamentos, usar texto padrão
                if not medicamentos_texto:
                    medicamentos_texto = "1) MEDICAMENTO CONFORME ORIENTAÇÃO MÉDICA -------------------------------- 30 comprimidos\n\nConforme orientação médica"
                    font_size = 11
                
                # Contexto com medicamentos dinâmicos (sem idade)
                context = {
                    'nome_paciente': remove_acentos(paciente_dict['nome_paciente'].upper()),
                    'data_nascimento': paciente_dict['dt_nascimento'].strftime('%d/%m/%Y') if paciente_dict['dt_nascimento'] else "xx/xx/xxxx",
                    'sexo': paciente_dict.get('sexo', 'Não informado'),
                    'cns': paciente_dict['cartao_sus'] if paciente_dict['cartao_sus'] else "CNS não registrado no PEC",
                    'ultima_atualizacao': medicamentos[0]['created_at'].strftime('%d/%m/%Y') if medicamentos[0]['created_at'] else "Não disponível",
                    'medicamentos_texto': medicamentos_texto,
                    'font_size': font_size
                }
                
                print(f"DEBUG: Context for patient {i}: {context}")
                
                # Carregar template e gerar documento
                print(f"DEBUG: Loading template...")
                doc = DocxTemplate(template_path)
                print(f"DEBUG: Rendering context...")
                doc.render(context)
                print(f"DEBUG: Render successful")
                
                # Salvar documento temporário
                temp_docx = os.path.join(temp_dir, f'receituario_{i}.docx')
                temp_pdf = os.path.join(temp_dir, f'receituario_{i}.pdf')
                
                print(f"DEBUG: Saving to {temp_docx}")
                doc.save(temp_docx)
                
                # Aplicar formatação específica aos medicamentos após renderização
                print(f"DEBUG: Aplicando formatação de fonte {font_size}pt e negrito...")
                from docx import Document as DocDocument
                from docx.shared import Pt
                rendered_doc = DocDocument(temp_docx)
                
                for paragraph in rendered_doc.paragraphs:
                    text = paragraph.text.strip()
                    
                    # Remover qualquer referência restante de idade no texto
                    if "( anos)" in text or "(  anos)" in text or "{{ idade }}" in text:
                        # Limpar o texto de referências de idade
                        new_text = text.replace("( anos)", "").replace("(  anos)", "").replace("{{ idade }}", "")
                        new_text = new_text.replace("  ", " ").strip()  # Remover espaços duplos
                        
                        # Limpar o parágrafo e reescrever
                        paragraph.clear()
                        new_run = paragraph.add_run(new_text)
                        print(f"  -> Limpeza de idade: '{text}' -> '{new_text}'")
                    
                    # Se é uma linha de medicamento (contém ") nome -------- quantidade comprimidos")
                    if (') ' in text and ('--------' in text or '-----' in text) and 'comprimidos' in text.lower()):
                        # Aplicar negrito e tamanho de fonte à linha do medicamento
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(font_size)
                        print(f"  -> Medicamento em negrito: {text[:50]}...")
                        
                    # Se é uma linha de instrução (começa com "Tomar")
                    elif text.startswith('Tomar '):
                        # Aplicar apenas tamanho de fonte (sem negrito) às instruções
                        for run in paragraph.runs:
                            run.font.bold = False
                            run.font.size = Pt(font_size)
                        
                # Duplicar o conteúdo para criar segunda página idêntica
                print(f"DEBUG: Duplicando receituário para segunda página...")
                
                # Coletar todos os parágrafos da primeira página ANTES de adicionar quebra
                original_paragraphs = []
                for paragraph in rendered_doc.paragraphs:
                    # Capturar o parágrafo completo
                    para_info = {
                        'text': paragraph.text,
                        'alignment': paragraph.alignment,
                        'runs': []
                    }
                    
                    # Capturar informações de formatação de cada run
                    for run in paragraph.runs:
                        run_info = {
                            'text': run.text,
                            'bold': run.font.bold,
                            'size': run.font.size,
                            'underline': run.font.underline,
                            'italic': run.font.italic
                        }
                        para_info['runs'].append(run_info)
                    
                    original_paragraphs.append(para_info)
                
                print(f"DEBUG: Capturados {len(original_paragraphs)} parágrafos da primeira página")
                
                # Adicionar quebra de página
                rendered_doc.add_page_break()
                
                # Recriar cada parágrafo da primeira página na segunda página
                for para_info in original_paragraphs:
                    new_p = rendered_doc.add_paragraph()
                    new_p.alignment = para_info['alignment']
                    
                    # Se não há runs, adicionar o texto simples
                    if not para_info['runs']:
                        if para_info['text']:
                            new_run = new_p.add_run(para_info['text'])
                            new_run.font.size = Pt(font_size)
                    else:
                        # Recriar cada run com sua formatação original
                        for run_info in para_info['runs']:
                            new_run = new_p.add_run(run_info['text'])
                            new_run.font.bold = run_info['bold']
                            new_run.font.size = run_info['size']
                            new_run.font.underline = run_info['underline']
                            new_run.font.italic = run_info['italic']
                
                # Salvar documento com receituário duplicado
                rendered_doc.save(temp_docx)
                print(f"DEBUG: Receituário duplicado criado - 2 páginas idênticas")
                
                # Verificar se arquivo foi criado
                if os.path.exists(temp_docx):
                    file_size = os.path.getsize(temp_docx)
                    print(f"DEBUG: File created successfully. Size: {file_size} bytes")
                    
                    # Testar se o arquivo pode ser reaberto
                    try:
                        from docx import Document
                        test_doc = Document(temp_docx)
                        print(f"DEBUG: File can be reopened. Paragraphs: {len(test_doc.paragraphs)}")
                    except Exception as e:
                        print(f"DEBUG: ERROR - File cannot be reopened: {e}")
                        continue
                else:
                    print(f"DEBUG: ERROR - File was not created!")
                    continue
                
                # Tentar conversão para PDF
                try:
                    print(f"DEBUG: Tentando conversão para PDF...")
                    
                    # Método 1: Usando docx2pdf com COM inicializado corretamente
                    import pythoncom
                    import threading
                    
                    # Inicializar COM de forma thread-safe
                    pythoncom.CoInitializeEx(pythoncom.COINIT_APARTMENTTHREADED)
                    
                    try:
                        from docx2pdf import convert
                        convert(temp_docx, temp_pdf)
                        
                        if os.path.exists(temp_pdf) and os.path.getsize(temp_pdf) > 0:
                            print(f"DEBUG: PDF criado com sucesso: {os.path.getsize(temp_pdf)} bytes")
                            output_files.append(temp_pdf)
                        else:
                            print(f"DEBUG: PDF não foi criado corretamente, usando DOCX")
                            output_files.append(temp_docx)
                            
                    finally:
                        pythoncom.CoUninitialize()
                        
                except Exception as e:
                    print(f"DEBUG: Erro na conversão PDF: {e}")
                    print(f"DEBUG: Tentando método alternativo...")
                    
                    # Método 2: Usar LibreOffice se disponível
                    try:
                        import subprocess
                        result = subprocess.run([
                            'soffice', '--headless', '--convert-to', 'pdf',
                            '--outdir', temp_dir, temp_docx
                        ], capture_output=True, text=True, timeout=30)
                        
                        if result.returncode == 0 and os.path.exists(temp_pdf):
                            print(f"DEBUG: PDF criado via LibreOffice: {os.path.getsize(temp_pdf)} bytes")
                            output_files.append(temp_pdf)
                        else:
                            print(f"DEBUG: LibreOffice falhou, usando DOCX")
                            output_files.append(temp_docx)
                            
                    except Exception as e2:
                        print(f"DEBUG: LibreOffice não disponível: {e2}")
                        print(f"DEBUG: Retornando DOCX como fallback")
                        output_files.append(temp_docx)
                
                print(f"DEBUG: Arquivo final adicionado: {output_files[-1]}")
            
            print(f"DEBUG: Total output files: {len(output_files)}")
            
            if not output_files:
                print("DEBUG: No files generated!")
                return jsonify({"erro": "Nenhum receituário foi gerado"}), 400
            
            # Se apenas um arquivo, retorna diretamente
            if len(output_files) == 1:
                file_path = output_files[0]
                print(f"DEBUG: Returning file: {file_path}")
                print(f"DEBUG: File size: {os.path.getsize(file_path)} bytes")
                
                with open(file_path, 'rb') as f:
                    file_data = f.read()
                
                print(f"DEBUG: File data read: {len(file_data)} bytes")
                
                mimetype = 'application/pdf' if file_path.endswith('.pdf') else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                extension = 'pdf' if file_path.endswith('.pdf') else 'docx'
                
                print(f"DEBUG: Returning {extension} file with mimetype {mimetype}")
                
                return Response(
                    file_data,
                    mimetype=mimetype,
                    headers={
                        'Content-Disposition': f'attachment; filename=receituario_hipertensao_{datetime.now().strftime("%Y%m%d")}.{extension}'
                    }
                )
            
            # Se múltiplos arquivos, seria necessário criar um ZIP
            # Por simplicidade, retornamos apenas o primeiro por enquanto
            file_path = output_files[0]
            with open(file_path, 'rb') as f:
                file_data = f.read()
            
            return Response(
                file_data,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={
                    'Content-Disposition': f'attachment; filename=receituarios_hipertensao_{datetime.now().strftime("%Y%m%d")}.docx'
                }
            )
        
    except Exception as e:
        print(f"Erro ao gerar receituário: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/generate_prescription_pdf_individual', methods=['POST'])
def api_generate_prescription_pdf_individual():
    """Gera PDF individual para um único paciente"""
    import unicodedata
    import tempfile
    import os
    from datetime import datetime
    
    
    data = request.get_json()
    patient = data.get('patient')
    
    if not patient:
        return jsonify({"erro": "Dados do paciente não fornecidos"}), 400
    
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # Caminho para o template
        template_path = os.path.join(os.path.dirname(__file__), 'modelos', 'template_receituario_dynamic.docx')
        
        if not os.path.exists(template_path):
            return jsonify({"erro": "Template não encontrado"}), 500
        
        # Buscar dados do paciente
        sql_paciente = """
            SELECT 
                nome_paciente,
                cartao_sus,
                dt_nascimento,
                sexo
            FROM sistemaaps.mv_hiperdia_hipertensao
            WHERE cod_paciente = %(cod_paciente)s
        """
        
        cur.execute(sql_paciente, {'cod_paciente': patient['cod_paciente']})
        paciente_dados = cur.fetchone()
        
        if not paciente_dados:
            return jsonify({"erro": "Paciente não encontrado"}), 404
            
        paciente_dict = dict(paciente_dados)
        
        # Buscar medicamentos
        sql_medicamentos = """
            SELECT 
                nome_medicamento,
                dose,
                frequencia,
                data_inicio,
                observacao,
                updated_at,
                created_at
            FROM sistemaaps.tb_hiperdia_has_medicamentos
            WHERE codcidadao = %(cod_paciente)s
            AND (data_fim IS NULL OR data_fim > CURRENT_DATE)
            ORDER BY created_at DESC, nome_medicamento
        """
        
        cur.execute(sql_medicamentos, {'cod_paciente': patient['cod_paciente']})
        medicamentos = cur.fetchall()
        
        if not medicamentos:
            return jsonify({"erro": "Nenhum medicamento ativo encontrado para este paciente"}), 404
        
        # Preparar dados para o template
        hoje = datetime.now()
        
        # Preparar lista de medicamentos
        medicamentos_lista = []
        for idx, med in enumerate(medicamentos, 1):
            med_dict = dict(med)
            nome = med_dict['nome_medicamento'].upper()
            dose = med_dict['dose'] or 1
            freq = med_dict['frequencia'] or 1
            total_comprimidos = dose * freq * 30
            
            # Preparar instruções
            instrucoes = []
            if freq == 1:
                instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 06:00 horas")
            elif freq == 2:
                instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 06:00 horas")
                instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 18:00 horas")
            elif freq == 3:
                instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 06:00 horas")
                instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 14:00 horas")
                instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 22:00 horas")
            elif freq == 4:
                instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 06:00 horas")
                instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 12:00 horas")
                instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 18:00 horas")
                instrucoes.append(f"Tomar {dose:02d} comprimido{'s' if dose > 1 else ''} as 24:00 horas")
            
            medicamentos_lista.append({
                'numero': idx,
                'nome': nome,
                'quantidade': total_comprimidos,
                'instrucoes': instrucoes
            })
        
        # Calcular tamanho da fonte baseado na quantidade de medicamentos
        num_medicamentos = len(medicamentos_lista)
        if num_medicamentos <= 2:  # 1-2 medicamentos
            font_size = 14
        elif num_medicamentos == 3:  # 3 medicamentos
            font_size = 14
        elif num_medicamentos == 4:  # 4 medicamentos
            font_size = 12
        elif num_medicamentos == 5:  # 5 medicamentos
            font_size = 12
        elif num_medicamentos == 6:  # 6 medicamentos
            font_size = 11
        elif num_medicamentos == 7:  # 7 medicamentos
            font_size = 10
        else:  # 8+ medicamentos (diminui 1pt para cada aumento)
            font_size = max(6, 16 - num_medicamentos)  # 8pt para 8 med, 7pt para 9 med, etc. Mínimo 6pt
        
        # Definir quantidade de traços baseada no tamanho da fonte
        if num_medicamentos <= 3:  # Fonte maior (14pt)
            tracos = "--------------------------------"  # 32 traços
        else:  # Fonte menor (12pt, 11pt, 10pt, 8pt, 7pt, etc)
            tracos = "---------------------------------------------"  # 45 traços
        
        # Gerar texto completo de medicamentos dinamicamente
        medicamentos_texto = ""
        for idx, med in enumerate(medicamentos_lista, 1):
            # Nome e quantidade do medicamento
            medicamentos_texto += f"{idx}) {med['nome']} {tracos} {med['quantidade']} comprimidos\n"
            
            # Adicionar todas as instruções
            for instrucao in med['instrucoes']:
                medicamentos_texto += f"{instrucao}\n"
            
            # Espaçamento entre medicamentos (se não for o último)
            if idx < len(medicamentos_lista):
                medicamentos_texto += "\n"
        
        # Se não há medicamentos, usar texto padrão
        if not medicamentos_texto:
            medicamentos_texto = "1) MEDICAMENTO CONFORME ORIENTAÇÃO MÉDICA -------------------------------- 30 comprimidos\n\nConforme orientação médica"
            font_size = 11
        
        # Contexto completo para o template (sem idade)
        context = {
            'nome_paciente': remove_acentos(paciente_dict['nome_paciente'].upper()),
            'data_nascimento': paciente_dict['dt_nascimento'].strftime('%d/%m/%Y') if paciente_dict['dt_nascimento'] else "xx/xx/xxxx",
            'sexo': paciente_dict.get('sexo', 'Não informado'),
            'cns': paciente_dict['cartao_sus'] if paciente_dict['cartao_sus'] else "CNS não registrado no PEC",
            'ultima_atualizacao': medicamentos[0]['created_at'].strftime('%d/%m/%Y') if medicamentos[0]['created_at'] else "Não disponível",
            'medicamentos_texto': medicamentos_texto,
            'font_size': font_size
        }
        
        # Criar arquivo temporário
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_docx = os.path.join(temp_dir, f"receituario_{patient['cod_paciente']}.docx")
            temp_pdf = os.path.join(temp_dir, f"receituario_{patient['cod_paciente']}.pdf")
            
            # Processar template usando DocxTemplate (mesma lógica do endpoint principal)
            from docxtpl import DocxTemplate
            from docx import Document as DocDocument
            from docx.shared import Pt
            
            # Carregar e renderizar template
            doc = DocxTemplate(template_path)
            doc.render(context)
            
            # Salvar documento temporário
            doc.save(temp_docx)
            
            # Aplicar formatação específica aos medicamentos após renderização
            rendered_doc = DocDocument(temp_docx)
            
            for paragraph in rendered_doc.paragraphs:
                text = paragraph.text.strip()
                
                # Remover qualquer referência restante de idade no texto
                if "( anos)" in text or "(  anos)" in text or "{{ idade }}" in text:
                    # Limpar o texto de referências de idade
                    new_text = text.replace("( anos)", "").replace("(  anos)", "").replace("{{ idade }}", "")
                    new_text = new_text.replace("  ", " ").strip()  # Remover espaços duplos
                    
                    # Limpar o parágrafo e reescrever
                    paragraph.clear()
                    new_run = paragraph.add_run(new_text)
                    print(f"  -> Limpeza de idade individual: '{text}' -> '{new_text}'")
                
                # Se é uma linha de medicamento (contém ") nome -------- quantidade comprimidos")
                if ') ' in text and ('--------' in text or '-----' in text) and 'comprimidos' in text.lower():
                    # Aplicar negrito e tamanho de fonte à linha do medicamento
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(font_size)
                        
                # Se é uma linha de instrução (começa com "Tomar")
                elif text.startswith('Tomar '):
                    # Aplicar apenas tamanho de fonte (sem negrito) às instruções
                    for run in paragraph.runs:
                        run.font.bold = False
                        run.font.size = Pt(font_size)
            
            # Duplicar o conteúdo para criar segunda página idêntica
            print(f"DEBUG: Duplicando receituário individual para segunda página...")
            
            # Coletar todos os parágrafos da primeira página ANTES de adicionar quebra
            original_paragraphs = []
            for paragraph in rendered_doc.paragraphs:
                # Capturar o parágrafo completo
                para_info = {
                    'text': paragraph.text,
                    'alignment': paragraph.alignment,
                    'runs': []
                }
                
                # Capturar informações de formatação de cada run
                for run in paragraph.runs:
                    run_info = {
                        'text': run.text,
                        'bold': run.font.bold,
                        'size': run.font.size,
                        'underline': run.font.underline,
                        'italic': run.font.italic
                    }
                    para_info['runs'].append(run_info)
                
                original_paragraphs.append(para_info)
            
            print(f"DEBUG: Capturados {len(original_paragraphs)} parágrafos da primeira página")
            
            # Adicionar quebra de página
            rendered_doc.add_page_break()
            
            # Recriar cada parágrafo da primeira página na segunda página
            for para_info in original_paragraphs:
                new_p = rendered_doc.add_paragraph()
                new_p.alignment = para_info['alignment']
                
                # Se não há runs, adicionar o texto simples
                if not para_info['runs']:
                    if para_info['text']:
                        new_run = new_p.add_run(para_info['text'])
                        new_run.font.size = Pt(font_size)
                else:
                    # Recriar cada run com sua formatação original
                    for run_info in para_info['runs']:
                        new_run = new_p.add_run(run_info['text'])
                        new_run.font.bold = run_info['bold']
                        new_run.font.size = run_info['size']
                        new_run.font.underline = run_info['underline']
                        new_run.font.italic = run_info['italic']
            
            # Salvar documento final
            rendered_doc.save(temp_docx)
            print(f"DEBUG: Receituário individual duplicado criado - 2 páginas idênticas")
            
            # Tentar conversão para PDF
            try:
                import pythoncom
                pythoncom.CoInitializeEx(pythoncom.COINIT_APARTMENTTHREADED)
                
                try:
                    from docx2pdf import convert
                    convert(temp_docx, temp_pdf)
                    
                    if os.path.exists(temp_pdf) and os.path.getsize(temp_pdf) > 0:
                        file_path = temp_pdf
                        mimetype = 'application/pdf'
                        extension = 'pdf'
                    else:
                        file_path = temp_docx
                        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        extension = 'docx'
                finally:
                    pythoncom.CoUninitialize()
                    
            except Exception as e:
                print(f"Erro na conversão PDF: {e}")
                file_path = temp_docx
                mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                extension = 'docx'
            
            # Ler arquivo e retornar
            with open(file_path, 'rb') as f:
                file_data = f.read()
            
            # Nome do arquivo conforme solicitado
            current_date = datetime.now().strftime('%d-%m-%Y')
            patient_name = remove_acentos(paciente_dict['nome_paciente'].upper()).replace(' ', '_')
            filename = f"RECEITUARIO ({current_date}) - HIPERTENSAO - {patient_name}.{extension}"
            
            return Response(
                file_data,
                mimetype=mimetype,
                headers={
                    'Content-Disposition': f'attachment; filename="{filename}"'
                }
            )
        
    except Exception as e:
        print(f"Erro ao gerar receituário individual: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

# ================================
# ENDPOINTS PARA DIABETES (HIPERDIA-DM)
# ================================

@app.route('/painel-hiperdia-dm')
def painel_hiperdia_dm():
    """Página principal do painel de diabetes"""
    from datetime import datetime
    data_atual = datetime.now().strftime('%d/%m/%Y')
    return render_template('painel-hiperdia-dm.html', data_atual=data_atual)

@app.route('/api/pacientes_hiperdia_dm')
def api_pacientes_hiperdia_dm():
    """API para buscar pacientes diabéticos com filtros"""
    conn = None
    cur = None
    try:
        # Parâmetros da requisição
        equipe = request.args.get('equipe', 'Todas')
        microarea = request.args.get('microarea', 'Todas')
        search = request.args.get('search', '').strip()
        status = request.args.get('status', 'Todos')
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 10))
        
        offset = (page - 1) * limit
        
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Base da query para pacientes diabéticos usando view específica com agente e ação atual
        base_query = """
            SELECT DISTINCT 
                d.cod_paciente,
                d.nome_paciente,
                d.dt_nascimento,
                d.sexo,
                d.cartao_sus,
                d.nome_equipe,
                d.microarea,
                d.tipo_diabetes,
                d.situacao_problema,
                ag.nome_agente,
                -- Ação atual (mais recente)
                ultima_acao.cod_acao as acao_atual_cod,
                ultima_acao.dsc_acao as acao_atual_nome,
                ultima_acao.status_acao as acao_atual_status,
                ultima_acao.data_agendamento as acao_atual_data_agendamento,
                ultima_acao.data_realizacao as acao_atual_data_realizacao,
                -- Status inteligente baseado na lógica de acompanhamento
                CASE 
                    WHEN ultima_acao.cod_acao IS NULL THEN 'sem_avaliacao'
                    WHEN ultima_acao.status_acao = 'AGUARDANDO' THEN 'em_analise'
                    WHEN ultima_acao.status_acao = 'REALIZADA' THEN 'em_analise'
                    WHEN ultima_acao.status_acao = 'FINALIZADO' THEN 'diabetes_compensada'
                    WHEN d.situacao_problema = 1 THEN 'controlado'
                    WHEN d.situacao_problema = 0 THEN 'descompensado'
                    ELSE 'indefinido'
                END as status_dm_novo,
                -- Status antigo para compatibilidade
                CASE 
                    WHEN d.situacao_problema = 1 THEN 'controlado'
                    WHEN d.situacao_problema = 0 THEN 'descompensado'
                    ELSE 'indefinido'
                END as status_dm
            FROM sistemaaps.mv_hiperdia_diabetes d
            LEFT JOIN sistemaaps.tb_agentes ag ON d.nome_equipe = ag.nome_equipe AND d.microarea = ag.micro_area
            LEFT JOIN LATERAL (
                SELECT 
                    a.cod_acao,
                    ta.dsc_acao,
                    a.status_acao,
                    a.data_agendamento,
                    a.data_realizacao,
                    a.created_at
                FROM sistemaaps.tb_hiperdia_dm_acompanhamento a
                LEFT JOIN sistemaaps.tb_hiperdia_tipos_acao ta ON a.cod_acao = ta.cod_acao
                WHERE a.cod_cidadao = d.cod_paciente
                ORDER BY a.created_at DESC
                LIMIT 1
            ) ultima_acao ON true
            WHERE 1=1
        """

        where_clauses = []
        params = {}
        
        # Filtro por equipe
        if equipe != 'Todas':
            where_clauses.append("d.nome_equipe = %(equipe)s")
            params['equipe'] = equipe
            
        # Filtro por microárea
        if microarea != 'Todas' and microarea != 'Todas as áreas':
            try:
                microarea_int = int(microarea)
                where_clauses.append("d.microarea = %(microarea)s")
                params['microarea'] = microarea_int
            except (ValueError, TypeError):
                # Se não conseguir converter para int, ignore o filtro
                pass
            
        # Filtro por busca (nome do paciente)
        if search:
            where_clauses.append("UPPER(d.nome_paciente) LIKE UPPER(%(search)s)")
            params['search'] = f'%{search}%'
        
        # Filtro por status específico para diabetes
        if status == 'Controlados':
            where_clauses.append("d.situacao_problema = 1") # 1 = Compensado/Controlado
        elif status == 'Descompensados':
            where_clauses.append("d.situacao_problema = 0") # 0 = Ativo/Descompensado
        elif status == 'ComTratamento':
            # Verificar se tem medicamentos ativos para diabetes
            where_clauses.append("""
                EXISTS (
                    SELECT 1 FROM sistemaaps.tb_hiperdia_dm_medicamentos med 
                    WHERE med.codcidadao = d.cod_paciente 
                    AND (med.data_fim IS NULL OR med.data_fim > CURRENT_DATE)
                )
            """)
        
        # Construir query final
        if where_clauses:
            full_query = base_query + " AND " + " AND ".join(where_clauses)
        else:
            full_query = base_query
            
        # Adicionar ordenação e paginação
        full_query += " ORDER BY d.nome_paciente LIMIT %(limit)s OFFSET %(offset)s"
        params['limit'] = limit
        params['offset'] = offset
        
        cur.execute(full_query, params)
        pacientes = cur.fetchall()
        
        # Contar total de pacientes - query com agente e ação atual  
        count_query = """
            SELECT COUNT(DISTINCT d.cod_paciente) as count
            FROM sistemaaps.mv_hiperdia_diabetes d
            LEFT JOIN sistemaaps.tb_agentes ag ON d.nome_equipe = ag.nome_equipe AND d.microarea = ag.micro_area
            LEFT JOIN LATERAL (
                SELECT 
                    a.cod_acao,
                    ta.dsc_acao,
                    a.status_acao,
                    a.data_agendamento,
                    a.data_realizacao,
                    a.created_at
                FROM sistemaaps.tb_hiperdia_dm_acompanhamento a
                LEFT JOIN sistemaaps.tb_hiperdia_tipos_acao ta ON a.cod_acao = ta.cod_acao
                WHERE a.cod_cidadao = d.cod_paciente
                ORDER BY a.created_at DESC
                LIMIT 1
            ) ultima_acao ON true
            WHERE 1=1
        """
        if where_clauses:
            count_query = count_query + " AND " + " AND ".join(where_clauses)
            
        cur.execute(count_query, {k: v for k, v in params.items() if k not in ['limit', 'offset']})
        total = cur.fetchone()['count']
        
        return jsonify({
            "pacientes": [dict(p) for p in pacientes],
            "total": total,
            "page": page,
            "limit": limit,
            "total_pages": (total + limit - 1) // limit
        })
        
    except Exception as e:
        print(f"Erro ao buscar pacientes diabéticos: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/get_total_diabeticos')
def api_get_total_diabeticos():
    """API para obter totais de pacientes diabéticos"""
    conn = None
    cur = None
    try:
        equipe = request.args.get('equipe', 'Todas')
        microarea = request.args.get('microarea', 'Todas')
        status = request.args.get('status', 'Todos')
        
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Query base para contar diabéticos usando view específica
        base_query = """
            SELECT COUNT(DISTINCT d.cod_paciente) as total_pacientes
            FROM sistemaaps.mv_hiperdia_diabetes d
            WHERE 1=1
        """
        
        where_clauses = []
        params = {}
        
        if equipe != 'Todas':
            where_clauses.append("d.nome_equipe = %(equipe)s")
            params['equipe'] = equipe
            
        if microarea != 'Todas' and microarea != 'Todas as áreas':
            try:
                microarea_int = int(microarea)
                where_clauses.append("d.microarea = %(microarea)s")
                params['microarea'] = microarea_int
            except (ValueError, TypeError):
                # Se não conseguir converter para int, ignore o filtro
                pass
        
        # Aplicar filtros de status específicos para diabetes
        if status == 'Controlados':
            where_clauses.append("d.situacao_problema = 1")
        elif status == 'Descompensados':
            where_clauses.append("d.situacao_problema = 0")
        elif status == 'ComTratamento':
            where_clauses.append("""
                EXISTS (
                    SELECT 1 FROM sistemaaps.tb_hiperdia_dm_medicamentos med 
                    WHERE med.codcidadao = d.cod_paciente 
                    AND (med.data_fim IS NULL OR med.data_fim > CURRENT_DATE)
                )
            """)
        
        if where_clauses:
            full_query = base_query + " AND " + " AND ".join(where_clauses)
        else:
            full_query = base_query
            
        cur.execute(full_query, params)
        result = cur.fetchone()
        
        return jsonify({
            "total_pacientes": result[0] if result else 0
        })
        
    except Exception as e:
        print(f"Erro ao buscar total de diabéticos: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/timeline/<int:cod_paciente>')
def api_diabetes_timeline(cod_paciente):
    """API para buscar timeline de acompanhamento de um paciente diabético"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Buscar histórico de acompanhamento incluindo dados MRG e exames
        query = """
        SELECT
                a.cod_acompanhamento,
                a.cod_acao,
                ta.dsc_acao,
                ta.dsc_detalhada,
                a.status_acao,
                a.data_agendamento,
                a.data_realizacao,
                a.observacoes,
                a.responsavel_pela_acao,
                a.created_at,
                mrg.data_mrg,
                mrg.g_jejum,
                mrg.g_apos_cafe,
                mrg.g_antes_almoco,
                mrg.g_apos_almoco,
                mrg.g_antes_jantar,
                mrg.g_ao_deitar,
                mrg.analise_mrg,
                mrg.periodo_mapeamento,
                mrg.dias_mapeamento,
                ex.hemoglobina_glicada,
                ex.glicemia_media,
                ex.glicemia_jejum as exame_glicemia_jejum,
                ex.data_exame,
                ex.observacoes as observacoes_exame
            FROM sistemaaps.tb_hiperdia_dm_acompanhamento a
            JOIN sistemaaps.tb_hiperdia_tipos_acao ta ON a.cod_acao = ta.cod_acao
            LEFT JOIN sistemaaps.tb_hiperdia_mrg mrg ON a.cod_acompanhamento = mrg.cod_acompanhamento
            LEFT JOIN sistemaaps.tb_hiperdia_dm_exames ex ON a.cod_acompanhamento = ex.cod_acompanhamento
            WHERE a.cod_cidadao = %(cod_paciente)s
            ORDER BY a.created_at DESC, mrg.cod_mrg ASC
        """
        
        cur.execute(query, {'cod_paciente': cod_paciente})
        results = cur.fetchall()
        
        # Processar timeline agrupando dados por acompanhamento
        timeline = []
        acompanhamentos = {}

        for row in results:
            evento = dict(row)
            cod_acompanhamento = evento['cod_acompanhamento']

            # Se já existe o acompanhamento, agregar dados
            if cod_acompanhamento in acompanhamentos:
                acomp = acompanhamentos[cod_acompanhamento]

                # Agregar mapeamentos MRG
                if any([evento['g_jejum'] is not None, evento['g_apos_cafe'] is not None,
                       evento['g_antes_almoco'] is not None, evento['g_apos_almoco'] is not None,
                       evento['g_antes_jantar'] is not None, evento['g_ao_deitar'] is not None]):
                    if 'mrg_mappings' not in acomp:
                        acomp['mrg_mappings'] = []
                    acomp['mrg_mappings'].append({
                        'data_mrg': evento['data_mrg'].strftime('%Y-%m-%d') if evento['data_mrg'] else None,
                        'periodo_mapeamento': evento['periodo_mapeamento'] or f'Período {len(acomp["mrg_mappings"]) + 1}',
                        'dias_mapeamento': evento['dias_mapeamento'] or 7,
                        'g_jejum': evento['g_jejum'],
                        'g_apos_cafe': evento['g_apos_cafe'],
                        'g_antes_almoco': evento['g_antes_almoco'],
                        'g_apos_almoco': evento['g_apos_almoco'],
                        'g_antes_jantar': evento['g_antes_jantar'],
                        'g_ao_deitar': evento['g_ao_deitar'],
                        'analise_mrg': evento['analise_mrg']
                    })
            else:
                # Novo acompanhamento
                acompanhamentos[cod_acompanhamento] = {
                    'cod_acompanhamento': evento['cod_acompanhamento'],
                    'cod_acao': evento['cod_acao'],
                    'dsc_acao': evento['dsc_acao'],
                    'dsc_detalhada': evento['dsc_detalhada'],
                    'status_acao': evento['status_acao'],
                    'data_agendamento': evento['data_agendamento'],
                    'data_realizacao': evento['data_realizacao'],
                    'observacoes': evento['observacoes'],
                    'responsavel_pela_acao': evento['responsavel_pela_acao'],
                    'created_at': evento['created_at']
                }

                # Adicionar dados de exames laboratoriais se existirem
                if any([evento['hemoglobina_glicada'] is not None, evento['glicemia_media'] is not None,
                       evento['exame_glicemia_jejum'] is not None]):
                    acompanhamentos[cod_acompanhamento]['lab_tests'] = {
                        'hemoglobina_glicada': evento['hemoglobina_glicada'],
                        'glicemia_media': evento['glicemia_media'],
                        'glicemia_jejum': evento['exame_glicemia_jejum'],
                        'data_exame': evento['data_exame'].strftime('%Y-%m-%d') if evento['data_exame'] else None,
                        'observacoes': evento['observacoes_exame']
                    }

                # Adicionar primeiro mapeamento MRG se existir
                if any([evento['g_jejum'] is not None, evento['g_apos_cafe'] is not None,
                       evento['g_antes_almoco'] is not None, evento['g_apos_almoco'] is not None,
                       evento['g_antes_jantar'] is not None, evento['g_ao_deitar'] is not None]):
                    acompanhamentos[cod_acompanhamento]['mrg_mappings'] = [{
                        'data_mrg': evento['data_mrg'].strftime('%Y-%m-%d') if evento['data_mrg'] else None,
                        'periodo_mapeamento': evento['periodo_mapeamento'] or 'Período 1',
                        'dias_mapeamento': evento['dias_mapeamento'] or 7,
                        'g_jejum': evento['g_jejum'],
                        'g_apos_cafe': evento['g_apos_cafe'],
                        'g_antes_almoco': evento['g_antes_almoco'],
                        'g_apos_almoco': evento['g_apos_almoco'],
                        'g_antes_jantar': evento['g_antes_jantar'],
                        'g_ao_deitar': evento['g_ao_deitar'],
                        'analise_mrg': evento['analise_mrg']
                    }]

        # Converter para lista ordenada
        timeline = list(acompanhamentos.values())
        timeline.sort(key=lambda x: x['created_at'], reverse=True)
        
        return jsonify({
            "timeline": timeline
        })
        
    except Exception as e:
        print(f"Erro ao buscar timeline do diabético: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/registrar_acao', methods=['POST'])
def api_diabetes_registrar_acao():
    """API para registrar nova ação de acompanhamento para diabético"""
    conn = None
    cur = None
    try:
        data = request.get_json()
        
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Inserir ação de acompanhamento
        query = """
            INSERT INTO sistemaaps.tb_hiperdia_dm_acompanhamento 
            (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
            VALUES (%(cod_cidadao)s, %(cod_acao_atual)s, %(status_acao)s, %(data_agendamento)s, %(data_realizacao)s, %(observacoes)s, %(responsavel_pela_acao)s)
            RETURNING cod_acompanhamento
        """
        
        # Determinar status e data de realização
        status_acao = data.get('status_acao', 'AGUARDANDO')
        data_realizacao = None
        
        # Se status for REALIZADA ou FINALIZADO, definir data_realizacao
        if status_acao in ['REALIZADA', 'FINALIZADO']:
            data_realizacao = data.get('data_realizacao', data['data_acao_atual'])
        
        cur.execute(query, {
            'cod_cidadao': data['cod_cidadao'],
            'cod_acao_atual': data['cod_acao_atual'],
            'status_acao': status_acao,
            'data_agendamento': data.get('data_agendamento', data['data_acao_atual']),
            'data_realizacao': data_realizacao,
            'observacoes': data.get('observacoes'),
            'responsavel_pela_acao': data.get('responsavel_pela_acao')
        })
        
        cod_acompanhamento = cur.fetchone()[0]
        
        # Tipo 3 agora é apenas solicitação, não insere dados de MRG
        # Dados de MRG são inseridos apenas no tipo 4 (Avaliar Tratamento)
        if False:  # Desabilitado - tipo 3 agora é apenas solicitação
            mrg_query = """
                INSERT INTO sistemaaps.tb_hiperdia_mrg 
                (cod_acompanhamento, data_mrg, g_jejum, g_apos_cafe, g_antes_almoco, g_apos_almoco, g_antes_jantar, g_ao_deitar, analise_mrg)
                VALUES (%(cod_acompanhamento)s, %(data_mrg)s, %(g_jejum)s, %(g_apos_cafe)s, %(g_antes_almoco)s, %(g_apos_almoco)s, %(g_antes_jantar)s, %(g_ao_deitar)s, %(analise_mrg)s)
            """
            
            mrg_data = data['mrg_data']
            cur.execute(mrg_query, {
                'cod_acompanhamento': cod_acompanhamento,
                'data_mrg': data['data_acao_atual'],
                'g_jejum': mrg_data.get('g_jejum'),
                'g_apos_cafe': mrg_data.get('g_apos_cafe'),
                'g_antes_almoco': mrg_data.get('g_antes_almoco'),
                'g_apos_almoco': mrg_data.get('g_apos_almoco'),
                'g_antes_jantar': mrg_data.get('g_antes_jantar'),
                'g_ao_deitar': mrg_data.get('g_ao_deitar'),
                'analise_mrg': mrg_data.get('analise_mrg')
            })
        
        # Ação de Modificar tratamento (cod_acao 3) agora apenas registra a ação na timeline
        # A modificação real do tratamento será feita através do modal de tratamento dedicado
        
        conn.commit()
        
        return jsonify({
            "sucesso": True,
            "cod_acompanhamento": cod_acompanhamento,
            "mensagem": "Ação registrada com sucesso"
        })
        
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro ao registrar ação para diabético: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/timeline/<int:cod_acompanhamento>/status', methods=['PUT'])
def api_diabetes_update_timeline_status(cod_acompanhamento):
    """API para atualizar status de uma ação da timeline (REALIZADA, CANCELADA)"""
    conn = None
    cur = None
    try:
        data = request.get_json()
        
        if not data or 'status' not in data:
            return jsonify({"sucesso": False, "erro": "Status é obrigatório"}), 400
        
        status = data['status']
        data_realizacao = data.get('data_realizacao')
        
        # Validar status - aceita todos os valores permitidos pela constraint do banco
        if status not in ['REALIZADA', 'CANCELADA', 'AGUARDANDO', 'PENDENTE', 'FINALIZADO']:
            return jsonify({"sucesso": False, "erro": "Status inválido. Valores aceitos: REALIZADA, CANCELADA, AGUARDANDO, PENDENTE, FINALIZADO"}), 400
        
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Verificar se a ação existe
        cur.execute("""
            SELECT cod_acompanhamento, status_acao, cod_cidadao
            FROM sistemaaps.tb_hiperdia_dm_acompanhamento 
            WHERE cod_acompanhamento = %s
        """, (cod_acompanhamento,))
        
        action = cur.fetchone()
        if not action:
            return jsonify({"sucesso": False, "erro": "Ação não encontrada"}), 404
        
        # Atualizar status
        update_query = """
            UPDATE sistemaaps.tb_hiperdia_dm_acompanhamento 
            SET status_acao = %s, 
                data_realizacao = %s,
                updated_at = CURRENT_TIMESTAMP
            WHERE cod_acompanhamento = %s
        """
        
        cur.execute(update_query, (status, data_realizacao, cod_acompanhamento))
        conn.commit()
        
        mensagem_map = {
            'REALIZADA': 'Ação marcada como concluída com sucesso!',
            'CANCELADA': 'Ação cancelada com sucesso!',
            'AGUARDANDO': 'Status da ação atualizado para aguardando',
            'PENDENTE': 'Status da ação atualizado para pendente'
        }
        
        return jsonify({
            "sucesso": True,
            "mensagem": mensagem_map.get(status, "Status atualizado com sucesso!")
        })
        
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro ao atualizar status da timeline: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/timeline/<int:cod_acompanhamento>', methods=['DELETE'])
def api_diabetes_delete_timeline_action(cod_acompanhamento):
    """API para excluir uma ação da timeline"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Verificar se a ação existe
        cur.execute("""
            SELECT cod_acompanhamento, cod_cidadao
            FROM sistemaaps.tb_hiperdia_dm_acompanhamento 
            WHERE cod_acompanhamento = %s
        """, (cod_acompanhamento,))
        
        action = cur.fetchone()
        if not action:
            return jsonify({"sucesso": False, "erro": "Ação não encontrada"}), 404
        
        # Primeiro, excluir dados MRG relacionados (se existirem)
        cur.execute("""
            DELETE FROM sistemaaps.tb_hiperdia_mrg 
            WHERE cod_acompanhamento = %s
        """, (cod_acompanhamento,))
        
        # Excluir a ação principal
        cur.execute("""
            DELETE FROM sistemaaps.tb_hiperdia_dm_acompanhamento 
            WHERE cod_acompanhamento = %s
        """, (cod_acompanhamento,))
        
        if cur.rowcount == 0:
            return jsonify({"sucesso": False, "erro": "Ação não pôde ser excluída"}), 400
        
        conn.commit()
        
        return jsonify({
            "sucesso": True,
            "mensagem": "Ação excluída com sucesso!"
        })
        
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro ao excluir ação da timeline: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/equipes_microareas_diabetes')
def api_equipes_microareas_diabetes():
    """API para buscar equipes e microáreas específicas para pacientes diabéticos"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # Buscar equipes, suas microáreas distintas e contagem de pacientes diabéticos
        cur.execute("""
            SELECT 
                nome_equipe, 
                array_agg(DISTINCT microarea ORDER BY microarea ASC) as microareas_list,
                COUNT(DISTINCT cod_paciente) as total_pacientes_diabeticos
            FROM sistemaaps.mv_hiperdia_diabetes
            WHERE nome_equipe IS NOT NULL AND microarea IS NOT NULL
            GROUP BY nome_equipe
            ORDER BY total_pacientes_diabeticos DESC, nome_equipe ASC;
        """)
        equipes_com_microareas = cur.fetchall()

        resultado_final = []
        for equipe_row in equipes_com_microareas:
            equipe_nome = equipe_row["nome_equipe"]
            
            # Para cada equipe, buscar os agentes associados às suas microáreas de pacientes diabéticos
            cur.execute("""
                SELECT DISTINCT ta.micro_area, ta.nome_agente
                FROM sistemaaps.tb_agentes ta
                WHERE ta.nome_equipe = %s AND ta.micro_area IN (
                    SELECT DISTINCT md.microarea 
                    FROM sistemaaps.mv_hiperdia_diabetes md 
                    WHERE md.nome_equipe = %s AND md.microarea IS NOT NULL
                )
                ORDER BY ta.micro_area, ta.nome_agente;
            """, (equipe_nome, equipe_nome))
            agentes_db = cur.fetchall()
            agentes_formatados = [{"micro_area": ag["micro_area"], "nome_agente": ag["nome_agente"]} for ag in agentes_db]

            resultado_final.append({
                "nome_equipe": equipe_nome,
                "agentes": agentes_formatados,
                "num_pacientes": equipe_row["total_pacientes_diabeticos"]
            })
        
        return jsonify(resultado_final)
    except Exception as e:
        print(f"Erro ao buscar equipes e microáreas para diabetes: {e}")
        # Se a view de diabetes não existir ainda, retornar dados do endpoint de hipertensão como fallback
        try:
            return api_equipes_microareas_hiperdia()
        except:
            return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/medicamentos_atuais/<int:cod_cidadao>')
def api_diabetes_medicamentos_atuais(cod_cidadao):
    """API para buscar medicamentos ativos de um paciente diabético"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        query = """
            SELECT 
                cod_seq_medicamento,
                nome_medicamento,
                dose,
                frequencia,
                posologia,
                data_inicio,
                observacoes,
                updated_at
            FROM sistemaaps.tb_hiperdia_dm_medicamentos
            WHERE codcidadao = %(cod_cidadao)s
            AND (data_fim IS NULL OR data_fim > CURRENT_DATE)
            ORDER BY nome_medicamento
        """
        
        cur.execute(query, {'cod_cidadao': cod_cidadao})
        medicamentos = cur.fetchall()
        
        # Converter para lista de dicionários e formatar
        medicamentos_list = []
        for med in medicamentos:
            med_dict = dict(med)
            # Tratar dados nulos
            if med_dict.get('data_inicio'):
                med_dict['data_inicio'] = med_dict['data_inicio'].strftime('%Y-%m-%d')
            medicamentos_list.append(med_dict)
        
        # Ordenar por data_inicio (mais recente primeiro), tratando valores None
        medicamentos_list.sort(key=lambda x: x.get('data_inicio') or '', reverse=True)
        
        return jsonify({
            "medicamentos": medicamentos_list
        })
        
    except Exception as e:
        print(f"Erro ao buscar medicamentos do diabético: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

# --- API Routes para Análise de Adolescentes ---

@app.route('/api/adolescentes/analytics/status_snapshot')
def api_adolescentes_status_snapshot():
    equipe_req = request.args.get('equipe', 'Todas')
    microarea_req = request.args.get('microarea', 'Todas as áreas')
    inicio = request.args.get('inicio')
    fim = request.args.get('fim')
    
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        where_clauses = ["m.idade_calculada BETWEEN 14 AND 18"]
        params = {}
        
        if equipe_req != 'Todas':
            where_clauses.append("m.nome_equipe = %(equipe)s")
            params['equipe'] = equipe_req
        if microarea_req != 'Todas as áreas':
            if ' - ' in microarea_req:
                parts = microarea_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else:
                micro_area_str = microarea_req.replace('Área ', '').strip()
            if micro_area_str:
                where_clauses.append("m.microarea = %(microarea)s")
                params['microarea'] = micro_area_str
        
        where_str = " WHERE " + " AND ".join(where_clauses)
        
        # KPIs seguindo o padrão do painel geral (gestantes, sem método, atraso, em dia)
        query = f"""
            SELECT 
                SUM(CASE WHEN m.status_gravidez = 'Grávida' THEN 1 ELSE 0 END) as gestantes,
                SUM(CASE WHEN (m.metodo IS NULL OR m.metodo = '') AND (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') THEN 1 ELSE 0 END) as sem_metodo,
                SUM(CASE WHEN 
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '90 days')) OR
                                ((m.metodo ILIKE '%%implante%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '1095 days')) OR
                                ((m.metodo ILIKE '%%diu%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') < (CURRENT_DATE - INTERVAL '3650 days'))
                            ))
                        )
                    THEN 1 ELSE 0 END) as metodo_atraso,
                SUM(CASE WHEN
                        (m.metodo IS NOT NULL AND m.metodo != '') AND
                        (m.status_gravidez IS NULL OR m.status_gravidez != 'Grávida') AND
                        NOT (m.metodo ILIKE '%%laqueadura%%' OR m.metodo ILIKE '%%vasectomia%%') AND
                        (
                            (m.data_aplicacao IS NOT NULL AND m.data_aplicacao != '' AND (
                                ((m.metodo ILIKE '%%mensal%%' OR m.metodo ILIKE '%%pílula%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '30 days')) OR
                                ((m.metodo ILIKE '%%trimestral%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '90 days')) OR
                                ((m.metodo ILIKE '%%implante%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '1095 days')) OR
                                ((m.metodo ILIKE '%%diu%%') AND TO_DATE(m.data_aplicacao, 'DD/MM/YYYY') >= (CURRENT_DATE - INTERVAL '3650 days'))
                            )) OR
                            (m.metodo ILIKE '%%laqueadura%%' OR m.metodo ILIKE '%%vasectomia%%') OR
                            (m.data_aplicacao IS NULL OR m.data_aplicacao = '')
                        )
                    THEN 1 ELSE 0 END) as metodo_em_dia
            FROM sistemaaps.mv_plafam m
            LEFT JOIN sistemaaps.tb_plafam_adolescentes pa ON m.cod_paciente = pa.co_cidadao
            {where_str}
        """
        
        cur.execute(query, params)
        result = cur.fetchone()
        
        return jsonify({
            'gestantes': result['gestantes'] or 0,
            'sem_metodo': result['sem_metodo'] or 0,
            'metodo_atraso': result['metodo_atraso'] or 0,
            'metodo_em_dia': result['metodo_em_dia'] or 0
        })
        
    except Exception as e:
        print(f"Erro ao buscar status dos adolescentes: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/adolescentes/analytics/mix_metodos')
def api_adolescentes_mix_metodos():
    equipe_req = request.args.get('equipe', 'Todas')
    microarea_req = request.args.get('microarea', 'Todas as áreas')
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

        where_clauses = ["m.idade_calculada BETWEEN 14 AND 18"]
        params = {}
        
        if equipe_req != 'Todas':
            where_clauses.append("m.nome_equipe = %(equipe)s")
            params['equipe'] = equipe_req
        if microarea_req != 'Todas as áreas':
            if ' - ' in microarea_req:
                parts = microarea_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else:
                micro_area_str = microarea_req.replace('Área ', '').strip()
            if micro_area_str:
                where_clauses.append("m.microarea = %(microarea)s")
                params['microarea'] = micro_area_str
        
        where_str = " WHERE " + " AND ".join(where_clauses)

        # Mix de métodos das adolescentes de 14-18 anos (métodos atualmente em uso)
        query = f"""
        SELECT
            CASE
                WHEN m.metodo ILIKE '%%laqueadura%%' THEN 'LAQUEADURA'
                WHEN m.metodo ILIKE '%%histerectomia%%' THEN 'HISTERECTOMIA'
                WHEN m.metodo ILIKE '%%diu%%' THEN 'DIU'
                WHEN m.metodo ILIKE '%%implante%%' THEN 'IMPLANTE SUBDÉRMICO'
                WHEN m.metodo ILIKE '%%mensal%%' THEN 'MENSAL'
                WHEN m.metodo ILIKE '%%trimestral%%' THEN 'TRIMESTRAL'
                WHEN m.metodo ILIKE '%%pílula%%' OR m.metodo ILIKE '%%pílula%%' OR m.metodo ILIKE '%%pilula%%' THEN 'PÍLULA'
                WHEN m.metodo IS NULL OR m.metodo = '' THEN 'SEM MÉTODO'
                ELSE 'OUTROS'
            END AS categoria,
            COUNT(*) AS total
        FROM sistemaaps.mv_plafam m
        {where_str}
        GROUP BY categoria
        ORDER BY total DESC;
        """
        
        cur.execute(query, params)
        rows = cur.fetchall() or []
        data = [{"categoria": r['categoria'], "total": int(r['total'])} for r in rows]
        return jsonify({"mix": data})
        
    except Exception as e:
        print(f"Erro ao buscar mix de métodos de adolescentes: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/adolescentes/analytics/metodos_desejados')
def api_adolescentes_metodos_desejados():
    equipe_req = request.args.get('equipe', 'Todas')
    microarea_req = request.args.get('microarea', 'Todas as áreas')
    
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        where_clauses = ["m.idade_calculada BETWEEN 14 AND 18"]
        params = {}
        
        if equipe_req != 'Todas':
            where_clauses.append("m.nome_equipe = %(equipe)s")
            params['equipe'] = equipe_req
        if microarea_req != 'Todas as áreas':
            if ' - ' in microarea_req:
                parts = microarea_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else:
                micro_area_str = microarea_req.replace('Área ', '').strip()
            if micro_area_str:
                where_clauses.append("m.microarea = %(microarea)s")
                params['microarea'] = micro_area_str
        
        where_str = " WHERE " + " AND ".join(where_clauses)
        
        # Query métodos desejados from adolescent action registrations (numeric codes)
        query = f"""
            SELECT 
                CASE 
                    WHEN pa.metodo_desejado = 1 THEN 'Pílula anticoncepcional'
                    WHEN pa.metodo_desejado = 2 THEN 'Injetável mensal'
                    WHEN pa.metodo_desejado = 3 THEN 'Injetável trimestral'
                    WHEN pa.metodo_desejado = 4 THEN 'DIU'
                    WHEN pa.metodo_desejado = 5 THEN 'Implante subdérmico'
                    ELSE 'Outros métodos'
                END as metodo_categorizado,
                COUNT(DISTINCT pa.co_cidadao) as quantidade
            FROM sistemaaps.mv_plafam m
            LEFT JOIN sistemaaps.tb_plafam_adolescentes pa ON m.cod_paciente = pa.co_cidadao
            {where_str} 
            AND pa.metodo_desejado IS NOT NULL 
            AND pa.resultado_abordagem = 1  -- Apenas quando "Deseja iniciar um método contraceptivo"
            GROUP BY 
                CASE 
                    WHEN pa.metodo_desejado = 1 THEN 'Pílula anticoncepcional'
                    WHEN pa.metodo_desejado = 2 THEN 'Injetável mensal'
                    WHEN pa.metodo_desejado = 3 THEN 'Injetável trimestral'
                    WHEN pa.metodo_desejado = 4 THEN 'DIU'
                    WHEN pa.metodo_desejado = 5 THEN 'Implante subdérmico'
                    ELSE 'Outros métodos'
                END
            ORDER BY quantidade DESC
        """
        
        cur.execute(query, params)
        results = cur.fetchall()
        
        metodos = []
        for row in results:
            if row['metodo_categorizado']:  # Apenas métodos válidos (não nulos)
                metodos.append({
                    'nome': row['metodo_categorizado'],
                    'quantidade': row['quantidade']
                })
        
        return jsonify({'metodos': metodos})
        
    except Exception as e:
        print(f"Erro ao buscar métodos desejados: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/adolescentes/analytics/tipos_abordagem')
def api_adolescentes_tipos_abordagem():
    equipe_req = request.args.get('equipe', 'Todas')
    microarea_req = request.args.get('microarea', 'Todas as áreas')
    
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        where_clauses = ["m.idade_calculada BETWEEN 14 AND 18"]
        params = {}
        
        if equipe_req != 'Todas':
            where_clauses.append("m.nome_equipe = %(equipe)s")
            params['equipe'] = equipe_req
        if microarea_req != 'Todas as áreas':
            if ' - ' in microarea_req:
                parts = microarea_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else:
                micro_area_str = microarea_req.replace('Área ', '').strip()
            if micro_area_str:
                where_clauses.append("m.microarea = %(microarea)s")
                params['microarea'] = micro_area_str
        
        where_str = " WHERE " + " AND ".join(where_clauses)
        
        # Mapear códigos de tipo de abordagem para nomes
        tipo_abordagem_map = {
            1: 'Abordagem com pais',
            2: 'Abordagem direta com adolescente', 
            5: 'Mudou de área',
            7: 'Remover do acompanhamento'
        }
        
        # Contar total de adolescentes e abordagens por tipo
        query_total = f"""
            SELECT COUNT(DISTINCT m.cod_paciente) as total_adolescentes
            FROM sistemaaps.mv_plafam m
            {where_str}
        """
        
        cur.execute(query_total, params)
        total_adolescentes = cur.fetchone()['total_adolescentes']
        
        query_abordagens = f"""
            SELECT 
                pa.tipo_abordagem,
                COUNT(DISTINCT pa.co_cidadao) as quantidade
            FROM sistemaaps.mv_plafam m
            INNER JOIN sistemaaps.tb_plafam_adolescentes pa ON m.cod_paciente = pa.co_cidadao
            {where_str} AND pa.tipo_abordagem IN (1, 2, 5, 7)
            GROUP BY pa.tipo_abordagem
            ORDER BY quantidade DESC
        """
        
        cur.execute(query_abordagens, params)
        results = cur.fetchall()
        
        abordagens = []
        total_com_abordagem = 0
        
        for row in results:
            tipo = row['tipo_abordagem']
            quantidade = row['quantidade']
            # Apenas processar códigos válidos
            if tipo in tipo_abordagem_map:
                total_com_abordagem += quantidade
                abordagens.append({
                    'tipo': tipo_abordagem_map[tipo],
                    'codigo': tipo,
                    'quantidade': quantidade
                })
        
        # Adicionar adolescentes sem abordagem
        sem_abordagem = total_adolescentes - total_com_abordagem
        if sem_abordagem > 0:
            abordagens.insert(0, {
                'tipo': 'Sem abordagem',
                'codigo': 0,
                'quantidade': sem_abordagem
            })
        
        return jsonify({'abordagens': abordagens})
        
    except Exception as e:
        print(f"Erro ao buscar tipos de abordagem: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/adolescentes/analytics/resultados_abordagem')
def api_adolescentes_resultados_abordagem():
    equipe_req = request.args.get('equipe', 'Todas')
    microarea_req = request.args.get('microarea', 'Todas as áreas')
    
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        where_clauses = ["m.idade_calculada BETWEEN 14 AND 18"]
        params = {}
        
        if equipe_req != 'Todas':
            where_clauses.append("m.nome_equipe = %(equipe)s")
            params['equipe'] = equipe_req
        if microarea_req != 'Todas as áreas':
            if ' - ' in microarea_req:
                parts = microarea_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else:
                micro_area_str = microarea_req.replace('Área ', '').strip()
            if micro_area_str:
                where_clauses.append("m.microarea = %(microarea)s")
                params['microarea'] = micro_area_str
        
        where_str = " WHERE " + " AND ".join(where_clauses)
        
        # Mapear códigos de resultado para nomes
        resultado_map = {
            1: 'Deseja iniciar método contraceptivo',
            2: 'Recusou método contraceptivo',
            3: 'Ausente em domicílio',
            4: 'Já usa um método',
            5: 'Mudou de área',
            6: 'Mudou de cidade',
            7: 'Método particular',
            8: 'Outros motivos'
        }
        
        query = f"""
            SELECT 
                CASE 
                    WHEN pa.resultado_abordagem = 1 THEN 'Deseja iniciar método contraceptivo'
                    WHEN pa.resultado_abordagem = 2 THEN 'Recusou método contraceptivo'
                    WHEN pa.resultado_abordagem IN (4, 7) THEN 'Já usa um método'
                    WHEN pa.resultado_abordagem IN (3, 5, 6, 8) THEN 'Outros resultados'
                    ELSE 'Outros resultados'
                END as resultado_agrupado,
                COUNT(*) as quantidade
            FROM sistemaaps.mv_plafam m
            INNER JOIN sistemaaps.tb_plafam_adolescentes pa ON m.cod_paciente = pa.co_cidadao
            {where_str} AND pa.resultado_abordagem IS NOT NULL
            GROUP BY 
                CASE 
                    WHEN pa.resultado_abordagem = 1 THEN 'Deseja iniciar método contraceptivo'
                    WHEN pa.resultado_abordagem = 2 THEN 'Recusou método contraceptivo'
                    WHEN pa.resultado_abordagem IN (4, 7) THEN 'Já usa um método'
                    WHEN pa.resultado_abordagem IN (3, 5, 6, 8) THEN 'Outros resultados'
                    ELSE 'Outros resultados'
                END
            ORDER BY quantidade DESC
        """
        
        cur.execute(query, params)
        results = cur.fetchall()
        
        resultados = []
        for row in results:
            resultados.append({
                'resultado': row['resultado_agrupado'],
                'quantidade': row['quantidade']
            })
        
        return jsonify({'resultados': resultados})
        
    except Exception as e:
        print(f"Erro ao buscar resultados de abordagem: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/adolescentes/analytics/timeline_acoes')
def api_adolescentes_timeline_acoes():
    equipe_req = request.args.get('equipe', 'Todas')
    microarea_req = request.args.get('microarea', 'Todas as áreas')
    granularidade = request.args.get('granularidade', 'month')
    inicio = request.args.get('inicio')
    fim = request.args.get('fim')
    
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        where_clauses = ["m.idade_calculada BETWEEN 14 AND 18"]
        params = {}
        
        if equipe_req != 'Todas':
            where_clauses.append("m.nome_equipe = %(equipe)s")
            params['equipe'] = equipe_req
        if microarea_req != 'Todas as áreas':
            if ' - ' in microarea_req:
                parts = microarea_req.split(' - ', 1)
                micro_area_str = parts[0].replace('Área ', '').strip()
            else:
                micro_area_str = microarea_req.replace('Área ', '').strip()
            if micro_area_str:
                where_clauses.append("m.microarea = %(microarea)s")
                params['microarea'] = micro_area_str
        
        if inicio:
            where_clauses.append("pa.data_acao >= %(inicio)s")
            params['inicio'] = inicio
        if fim:
            where_clauses.append("pa.data_acao <= %(fim)s")
            params['fim'] = fim
        
        where_str = " WHERE " + " AND ".join(where_clauses)
        
        # Definir formato de data baseado na granularidade
        if granularidade == 'day':
            date_trunc = 'day'
            date_format = 'DD/MM/YYYY'
        elif granularidade == 'week':
            date_trunc = 'week'
            date_format = 'DD/MM/YYYY'
        else:  # month
            date_trunc = 'month'
            date_format = 'MM/YYYY'
        
        query = f"""
            SELECT 
                DATE_TRUNC('{date_trunc}', pa.data_acao) as periodo,
                TO_CHAR(DATE_TRUNC('{date_trunc}', pa.data_acao), '{date_format}') as periodo_str,
                COUNT(*) as total_acoes
            FROM sistemaaps.mv_plafam m
            INNER JOIN sistemaaps.tb_plafam_adolescentes pa ON m.cod_paciente = pa.co_cidadao
            {where_str} AND pa.data_acao IS NOT NULL
            GROUP BY DATE_TRUNC('{date_trunc}', pa.data_acao)
            ORDER BY periodo
        """
        
        cur.execute(query, params)
        results = cur.fetchall()
        
        timeline = []
        for row in results:
            timeline.append({
                'periodo': row['periodo_str'],
                'acoes': row['total_acoes']
            })
        
        return jsonify({'timeline': timeline})
        
    except Exception as e:
        print(f"Erro ao buscar timeline de ações: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/hiperdia/export_pdf', methods=['POST'])
def api_hiperdia_export_pdf():
    """Exporta lista de pacientes hipertensos para PDF"""
    try:
        data = request.get_json()
        pacientes = data.get('pacientes', [])
        filtros = data.get('filtros', {})
        
        if not pacientes:
            return jsonify({"erro": "Nenhum paciente fornecido para exportação"}), 400
        
        # Criar buffer para o PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4))  # Layout paisagem
        story = []
        styles = getSampleStyleSheet()
        
        # Título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        story.append(Paragraph("HIPERDIA - Lista de Pacientes Hipertensos", title_style))
        
        # Informações dos filtros
        filter_info = []
        if filtros.get('equipe') and filtros['equipe'] != 'Todas as equipes':
            filter_info.append(f"Equipe: {filtros['equipe']}")
        if filtros.get('microarea') and filtros['microarea'] != 'Todas as áreas':
            filter_info.append(f"Microárea: {filtros['microarea']}")
        if filtros.get('status') and filtros['status'] != 'Todos':
            filter_info.append(f"Status: {filtros['status']}")
        
        if filter_info:
            story.append(Paragraph(f"Filtros aplicados: {' | '.join(filter_info)}", styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Data de geração
        story.append(Paragraph(f"Data de geração: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Tabela de pacientes - apenas 4 colunas
        table_data = [
            ['Paciente', 'Tratamento Atual', 'Próxima Ação', 'Status']
        ]
        
        for paciente in pacientes:
            # Criar parágrafo para as informações do paciente na primeira coluna
            nome = paciente.get('Nome', '')
            idade = paciente.get('Idade', '')
            cns = paciente.get('CNS', '')
            area_agente = paciente.get('Area - Agente', '')
            
            # Montar texto da primeira coluna com quebras de linha
            paciente_info = f"<b>{nome}</b>"
            if idade:
                paciente_info += f", {idade} anos"
            if cns:
                paciente_info += f"<br/>CNS: {cns}"
            if area_agente:
                paciente_info += f"<br/>{area_agente}"
                
            paciente_paragraph = Paragraph(paciente_info, ParagraphStyle(
                'PatientStyle',
                fontName='Helvetica',
                fontSize=9,
                leading=11,
                leftIndent=2,
                rightIndent=2
            ))
            
            # Para a coluna tratamento atual, processar medicamentos individualmente
            tratamento_text = paciente.get('Tratamento Atual', '')
            if tratamento_text and tratamento_text != 'Nenhum tratamento cadastrado':
                # Separar medicamentos usando múltiplos padrões
                medicamentos = []
                
                # Primeiro tenta separar por quebras de linha existentes
                if '<br>' in tratamento_text:
                    medicamentos = tratamento_text.split('<br>')
                elif '\n' in tratamento_text:
                    medicamentos = tratamento_text.split('\n')
                else:
                    # Para o padrão específico: "mg (Xx/dia)" seguido de letra minúscula
                    # Exemplo: "10 mg (1x/dia)losartana 50 mg"
                    medicamentos = re.split(r'(?<=\))\s*(?=[a-z])', tratamento_text)
                    
                    # Se ainda não separou bem, usar padrão mais amplo
                    if len(medicamentos) <= 1:
                        # Separar após parênteses fechado seguido de qualquer palavra
                        medicamentos = re.split(r'(?<=\))\s*(?=\w)', tratamento_text)
                
                # Formatar cada medicamento com símbolo de check verde
                medicamentos_formatados = []
                for med in medicamentos:
                    med = med.strip()
                    if med:
                        # Usar símbolo de check verde (✓) - funciona em ReportLab
                        medicamentos_formatados.append(f"<font color='green'>✓</font> {med}")
                
                tratamento_formatado = '<br/>'.join(medicamentos_formatados)
                
                tratamento_paragraph = Paragraph(tratamento_formatado, ParagraphStyle(
                    'TreatmentStyle',
                    fontName='Helvetica',
                    fontSize=7,
                    leading=9,
                    leftIndent=2,
                    rightIndent=2,
                    spaceAfter=2
                ))
            else:
                tratamento_paragraph = Paragraph(tratamento_text, ParagraphStyle(
                    'TreatmentStyle',
                    fontName='Helvetica',
                    fontSize=8,
                    leading=9
                ))
                
            table_data.append([
                paciente_paragraph,  # Todas as informações do paciente
                tratamento_paragraph,  # Usar parágrafo para quebra de linha
                paciente.get('Próxima Ação', ''),  # Próxima ação
                paciente.get('Status', '')  # Status
            ])
        
        # Ajustar larguras das colunas para 4 colunas no layout paisagem
        # Paciente maior, Tratamento reduzido
        table = Table(table_data, colWidths=[110*mm, 70*mm, 60*mm, 25*mm])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Alinhamento vertical superior
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            # Alinhamento específico para a coluna do paciente (índice 0)
            ('ALIGN', (0, 1), (0, -1), 'LEFT'),  # Alinhar texto à esquerda na coluna paciente
            ('LEFTPADDING', (0, 1), (0, -1), 6),
            ('RIGHTPADDING', (0, 1), (0, -1), 6),
            # Alinhamento específico para a coluna de tratamento (índice 1)
            ('ALIGN', (1, 1), (1, -1), 'LEFT'),  # Alinhar texto à esquerda na coluna tratamento
            ('LEFTPADDING', (1, 1), (1, -1), 4),
            ('RIGHTPADDING', (1, 1), (1, -1), 4),
            # Padding geral
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 20))
        
        # Rodapé
        story.append(Paragraph(f"Total de pacientes: {len(pacientes)}", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename=hiperdia_pacientes_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
            }
        )
        
    except Exception as e:
        print(f"Erro ao gerar PDF de exportação: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500

@app.route('/plano_semanal_novo')
def plano_semanal_personalizado():
    """API para gerar Plano Semanal personalizado com configurações específicas"""
    conn = None
    cur = None
    try:
        print("=== NOVA FUNÇÃO PLANO SEMANAL CHAMADA ===")
        conn = get_db_connection()
        cur = conn.cursor()  # Usando cursor normal como outros endpoints
        
        # Parâmetros da requisição
        equipe = request.args.get('equipe', 'Todas')
        microarea = request.args.get('microarea', 'Todas')
        sem_metodo_qtd = int(request.args.get('sem_metodo_qtd', 0))
        metodo_vencido_qtd = int(request.args.get('metodo_vencido_qtd', 0))
        organizacao_idade = request.args.get('organizacao_idade', 'crescente')
        
        # Verificar se há pelo menos uma quantidade
        if sem_metodo_qtd == 0 and metodo_vencido_qtd == 0:
            return jsonify({"success": False, "error": "Defina pelo menos uma quantidade"})
        
        all_patients = []
        
        def get_patients_by_microarea_ordered(sem_metodo_qtd, equipe_filter, ordem_idade):
            """Buscar pacientes sem método com ordenação específica: X por microárea de cada equipe, ordenados intercaladamente"""
            patients = []
            
            # Lista de equipes válidas na ordem específica (nomes corretos)
            equipes_validas = ['PSF-01', 'PSF-02', 'PSF - 03', 'URBANA I', 'URBANA II', 'URBANA III']
            microareas_validas = [1, 2, 3, 4, 5, 6]
            
            # Se equipe específica foi selecionada, usar apenas ela
            if equipe_filter and equipe_filter != 'Todas' and equipe_filter in equipes_validas:
                equipes_validas = [equipe_filter]
            
            print(f"=== BUSCANDO PACIENTES ===")
            print(f"Quantidade sem método: {sem_metodo_qtd}")
            print(f"Equipes: {equipes_validas}")
            print(f"Microáreas: {microareas_validas}")
            
            # Buscar pacientes por equipe e microárea
            pacientes_por_equipe_micro = {}
            
            # Definir direção da ordenação baseada no parâmetro
            ordem_direction = "ASC" if ordem_idade == "crescente" else "DESC"
            print(f"Ordenação por idade: {ordem_idade} ({ordem_direction})")
            
            for equipe in equipes_validas:
                pacientes_por_equipe_micro[equipe] = {}
                for microarea in microareas_validas:
                    query = f"""
                    SELECT
                        m.cod_paciente, m.nome_paciente, m.cartao_sus, m.idade_calculada, 
                        m.microarea, m.metodo, m.nome_equipe, 
                        m.data_aplicacao, m.status_gravidez, m.data_provavel_parto,
                        pa.status_acompanhamento, pa.data_acompanhamento,
                        ag.nome_agente
                    FROM sistemaaps.mv_plafam m
                    LEFT JOIN sistemaaps.tb_plafam_acompanhamento pa ON m.cod_paciente = pa.co_cidadao
                    LEFT JOIN sistemaaps.tb_agentes ag ON m.microarea = ag.micro_area AND m.nome_equipe = ag.nome_equipe
                    WHERE m.nome_equipe = %(equipe)s
                    AND m.microarea = %(microarea)s
                    AND m.status_gravidez != 'Grávida'
                    AND m.idade_calculada >= 19 AND m.idade_calculada <= 45
                    AND (pa.status_acompanhamento IS NULL OR pa.status_acompanhamento = 0)
                    AND (m.metodo = '' OR m.metodo IS NULL)
                    ORDER BY m.idade_calculada {ordem_direction}
                    LIMIT %(limite)s
                    """
                    
                    cur.execute(query, {
                        'equipe': equipe, 
                        'microarea': microarea, 
                        'limite': sem_metodo_qtd
                    })
                    results = cur.fetchall()
                    column_names = [desc[0] for desc in cur.description]
                    
                    pacientes_microarea = []
                    for row in results:
                        patient_dict = dict(zip(column_names, row))
                        patient_dict['categoria_plano'] = 'Sem método'
                        pacientes_microarea.append(patient_dict)
                    
                    pacientes_por_equipe_micro[equipe][microarea] = pacientes_microarea
                    print(f"{equipe} - Microárea {microarea}: {len(pacientes_microarea)} pacientes")
            
            # Ordenação por equipe, depois por microárea, depois por posição
            # equipe 1, micro area 1 - 1º paciente
            # equipe 1, micro area 1 - 2º paciente  
            # equipe 1, micro area 2 - 1º paciente
            # equipe 1, micro area 2 - 2º paciente
            # ... até completar todas as microáreas da equipe 1
            # equipe 2, micro area 1 - 1º paciente
            # equipe 2, micro area 1 - 2º paciente
            # etc.
            
            for equipe in equipes_validas:  # Para cada equipe na ordem
                for microarea in microareas_validas:  # Para cada microárea
                    for pos in range(sem_metodo_qtd):  # Para cada posição (1º, 2º, etc.)
                        pacientes_micro = pacientes_por_equipe_micro[equipe][microarea]
                        if pos < len(pacientes_micro):  # Se existe paciente nesta posição
                            patients.append(pacientes_micro[pos])
            
            print(f"Total de pacientes ordenados: {len(patients)}")
            return patients
        
        # Buscar pacientes usando a nova função ordenada (por enquanto apenas sem método)
        if sem_metodo_qtd > 0:
            all_patients = get_patients_by_microarea_ordered(
                sem_metodo_qtd, 
                equipe, 
                organizacao_idade
            )
        
        # Formatar datas se necessário
        print(f"DEBUG: Iniciando formatação de datas para {len(all_patients)} pacientes...")
        for i, p in enumerate(all_patients):
            print(f"DEBUG: Formatando paciente {i}")
            try:
                if p.get('data_aplicacao') and isinstance(p['data_aplicacao'], date):
                    p['data_aplicacao'] = p['data_aplicacao'].strftime('%Y-%m-%d')
                if p.get('data_acompanhamento') and isinstance(p['data_acompanhamento'], date):
                    p['data_acompanhamento'] = p['data_acompanhamento'].strftime('%Y-%m-%d')
                if p.get('data_provavel_parto') and isinstance(p['data_provavel_parto'], date):
                    p['data_provavel_parto'] = p['data_provavel_parto'].strftime('%Y-%m-%d')
                print(f"DEBUG: Paciente {i} formatado com sucesso")
            except Exception as e:
                print(f"DEBUG: Erro formatando paciente {i}: {e}")
                raise
        
        return jsonify({
            "success": True,
            "pacientes": all_patients,
            "total": len(all_patients),
            "configuracao": {
                "sem_metodo_qtd": sem_metodo_qtd,
                "metodo_vencido_qtd": metodo_vencido_qtd,
                "organizacao_idade": organizacao_idade,
                "equipe": equipe,
                "microarea": microarea
            }
        })
        
    except Exception as e:
        print(f"Erro ao gerar plano semanal personalizado: {e}")
        return jsonify({"success": False, "error": f"Erro no servidor: {e}"})
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

# --- API Routes para Medicamentos de Diabetes ---

@app.route('/api/diabetes/medicamentos', methods=['POST'])
def api_diabetes_adicionar_medicamento():
    """API para adicionar novo medicamento para paciente diabético"""
    conn = None
    cur = None
    try:
        data = request.get_json()
        
        # Validação dos dados obrigatórios
        cod_cidadao = data.get('codcidadao')
        nome_medicamento = data.get('nome_medicamento')
        dose = data.get('dose', 1)
        frequencia = data.get('frequencia', 1)
        
        if not cod_cidadao or not nome_medicamento:
            return jsonify({"sucesso": False, "erro": "Código do cidadão e nome do medicamento são obrigatórios."}), 400
        
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Inserir medicamento
        query = """
            INSERT INTO sistemaaps.tb_hiperdia_dm_medicamentos 
            (codcidadao, nome_medicamento, dose, frequencia, posologia, data_inicio, observacoes)
            VALUES (%(codcidadao)s, %(nome_medicamento)s, %(dose)s, %(frequencia)s, %(posologia)s, %(data_inicio)s, %(observacoes)s)
            RETURNING cod_seq_medicamento
        """
        
        params = {
            'codcidadao': cod_cidadao,
            'nome_medicamento': nome_medicamento,
            'dose': dose,
            'frequencia': frequencia,
            'posologia': data.get('posologia', ''),
            'data_inicio': data.get('data_inicio'),
            'observacoes': data.get('observacoes', '')
        }
        
        cur.execute(query, params)
        cod_seq_medicamento = cur.fetchone()[0]
        conn.commit()
        
        return jsonify({
            "sucesso": True, 
            "mensagem": "Medicamento adicionado com sucesso.",
            "cod_seq_medicamento": cod_seq_medicamento
        })
        
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro ao adicionar medicamento diabético: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/medicamentos/<int:cod_seq_medicamento>', methods=['PUT'])
def api_diabetes_atualizar_medicamento(cod_seq_medicamento):
    """API para atualizar medicamento de paciente diabético"""
    conn = None
    cur = None
    try:
        data = request.get_json()
        
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Construir query de atualização dinamicamente
        campos_atualizacao = []
        params = {'cod_seq_medicamento': cod_seq_medicamento}
        
        if 'nome_medicamento' in data:
            campos_atualizacao.append("nome_medicamento = %(nome_medicamento)s")
            params['nome_medicamento'] = data['nome_medicamento']
        
        if 'dose' in data:
            campos_atualizacao.append("dose = %(dose)s")
            params['dose'] = data['dose']
        
        if 'frequencia' in data:
            campos_atualizacao.append("frequencia = %(frequencia)s")
            params['frequencia'] = data['frequencia']
        
        if 'posologia' in data:
            campos_atualizacao.append("posologia = %(posologia)s")
            params['posologia'] = data['posologia']
        
        if 'observacoes' in data:
            campos_atualizacao.append("observacoes = %(observacoes)s")
            params['observacoes'] = data['observacoes']
        
        campos_atualizacao.append("updated_at = CURRENT_TIMESTAMP")
        
        if not campos_atualizacao:
            return jsonify({"sucesso": False, "erro": "Nenhum campo para atualizar."}), 400
        
        query = f"""
            UPDATE sistemaaps.tb_hiperdia_dm_medicamentos 
            SET {', '.join(campos_atualizacao)}
            WHERE cod_seq_medicamento = %(cod_seq_medicamento)s
        """
        
        cur.execute(query, params)
        conn.commit()
        
        if cur.rowcount == 0:
            return jsonify({"sucesso": False, "erro": "Medicamento não encontrado."}), 404
        
        return jsonify({"sucesso": True, "mensagem": "Medicamento atualizado com sucesso."})
        
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro ao atualizar medicamento diabético: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/medicamentos/<int:cod_seq_medicamento>/interromper', methods=['PUT'])
def api_diabetes_interromper_medicamento(cod_seq_medicamento):
    """API para interromper medicamento de paciente diabético"""
    conn = None
    cur = None
    try:
        data = request.get_json()
        motivo_interrupcao = data.get('motivo_interrupcao', 'Interrompido pelo profissional')
        
        conn = get_db_connection()
        cur = conn.cursor()
        
        query = """
            UPDATE sistemaaps.tb_hiperdia_dm_medicamentos 
            SET data_fim = CURRENT_DATE, 
                motivo_interrupcao = %(motivo_interrupcao)s,
                updated_at = CURRENT_TIMESTAMP
            WHERE cod_seq_medicamento = %(cod_seq_medicamento)s
        """
        
        cur.execute(query, {
            'cod_seq_medicamento': cod_seq_medicamento,
            'motivo_interrupcao': motivo_interrupcao
        })
        conn.commit()
        
        if cur.rowcount == 0:
            return jsonify({"sucesso": False, "erro": "Medicamento não encontrado."}), 404
        
        return jsonify({"sucesso": True, "mensagem": "Medicamento interrompido com sucesso."})
        
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro ao interromper medicamento diabético: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/medicamentos_diabetes')
def api_diabetes_medicamentos_disponiveis():
    """API para buscar lista de medicamentos disponíveis para diabetes"""
    try:
        # Lista de medicamentos comuns para diabetes
        medicamentos_diabetes = [
            "Metformina 500mg",
            "Metformina 850mg", 
            "Glibenclamida 5mg",
            "Gliclazida 30mg",
            "Gliclazida 60mg",
            "Insulina NPH",
            "Insulina Regular",
            "Insulina Glargina",
            "Insulina Lispro",
            "Sitagliptina 50mg",
            "Sitagliptina 100mg",
            "Empagliflozina 10mg",
            "Empagliflozina 25mg",
            "Pioglitazona 15mg",
            "Pioglitazona 30mg",
            "Repaglinida 1mg",
            "Repaglinida 2mg",
            "Acarbose 50mg",
            "Acarbose 100mg"
        ]
        
        return jsonify({"medicamentos": medicamentos_diabetes})
        
    except Exception as e:
        print(f"Erro ao buscar medicamentos para diabetes: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500

# --- API Routes para Insulinas ---

@app.route('/api/diabetes/insulinas/<int:cod_cidadao>')
def api_diabetes_insulinas_atuais(cod_cidadao):
    """API para buscar insulinas ativas de um paciente diabético"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        query = """
            SELECT 
                cod_seq_insulina,
                tipo_insulina,
                frequencia_dia,
                doses_estruturadas,
                data_inicio,
                observacoes,
                created_at
            FROM sistemaaps.tb_hiperdia_dm_insulina 
            WHERE codcidadao = %s 
            AND data_fim IS NULL
            ORDER BY tipo_insulina, created_at DESC
        """
        
        cur.execute(query, (cod_cidadao,))
        insulinas = cur.fetchall()
        
        # Processar dados das insulinas para formato mais amigável
        insulinas_processadas = []
        for insulina in insulinas:
            doses_str = insulina['doses_estruturadas'] or '[]'
            try:
                doses = json.loads(doses_str)
            except (json.JSONDecodeError, TypeError):
                doses = []
            
            doses_formatadas = []
            
            for dose in doses:
                dose_str = f"{dose.get('dose', 0)} U às {dose.get('horario', 'N/A')}"
                doses_formatadas.append(dose_str)
            
            insulina_processada = dict(insulina)
            insulina_processada['doses_formatadas'] = doses_formatadas
            insulina_processada['doses_resumo'] = ' | '.join(doses_formatadas)
            insulinas_processadas.append(insulina_processada)
        
        return jsonify({"sucesso": True, "insulinas": insulinas_processadas})
        
    except Exception as e:
        print(f"Erro ao buscar insulinas para paciente {cod_cidadao}: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/insulinas', methods=['POST'])
def api_diabetes_adicionar_insulina():
    """API para adicionar nova insulina para paciente diabético"""
    conn = None
    cur = None
    try:
        data = request.get_json()
        
        # Validação dos dados obrigatórios
        cod_cidadao = data.get('codcidadao')
        tipo_insulina = data.get('tipo_insulina')
        frequencia_dia = data.get('frequencia_dia')
        doses_estruturadas = data.get('doses_estruturadas')
        
        if not all([cod_cidadao, tipo_insulina, frequencia_dia, doses_estruturadas]):
            return jsonify({"sucesso": False, "erro": "Campos obrigatórios não preenchidos"}), 400
        
        # Validações específicas
        if frequencia_dia not in [1, 2, 3, 4]:
            return jsonify({"sucesso": False, "erro": "Frequência deve ser entre 1 e 4 vezes ao dia"}), 400
        
        if tipo_insulina not in ['Insulina NPH', 'Insulina Regular', 'Insulina Glargina', 'Insulina Lispro']:
            return jsonify({"sucesso": False, "erro": "Tipo de insulina inválido"}), 400
        
        if len(doses_estruturadas) != frequencia_dia:
            return jsonify({"sucesso": False, "erro": "Número de doses deve corresponder à frequência"}), 400
        
        # Validar formato das doses
        for dose in doses_estruturadas:
            if not isinstance(dose, dict) or 'dose' not in dose or 'horario' not in dose:
                return jsonify({"sucesso": False, "erro": "Formato de doses inválido"}), 400
            
            dose_valor = dose.get('dose')
            if not isinstance(dose_valor, int) or dose_valor < 1 or dose_valor > 100:
                return jsonify({"sucesso": False, "erro": "Dose deve ser entre 1 e 100 unidades"}), 400
        
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Verificar se já existe insulina do mesmo tipo ativa para o paciente
        check_query = """
            SELECT COUNT(*) FROM sistemaaps.tb_hiperdia_dm_insulina 
            WHERE codcidadao = %s AND tipo_insulina = %s AND data_fim IS NULL
        """
        cur.execute(check_query, (cod_cidadao, tipo_insulina))
        existing_count = cur.fetchone()[0]
        
        if existing_count > 0:
            return jsonify({"sucesso": False, "erro": f"Paciente já possui {tipo_insulina} ativa. Interrompa a insulina anterior primeiro."}), 400
        
        # Inserir nova insulina
        insert_query = """
            INSERT INTO sistemaaps.tb_hiperdia_dm_insulina 
            (codcidadao, tipo_insulina, frequencia_dia, doses_estruturadas, data_inicio, observacoes)
            VALUES (%s, %s, %s, %s, %s, %s)
            RETURNING cod_seq_insulina
        """
        
        data_inicio = data.get('data_inicio') or None
        observacoes = data.get('observacoes', '')
        
        cur.execute(insert_query, (
            cod_cidadao,
            tipo_insulina,
            frequencia_dia,
            json.dumps(doses_estruturadas),
            data_inicio,
            observacoes
        ))
        
        cod_seq_insulina = cur.fetchone()[0]
        conn.commit()
        
        print(f"Nova insulina adicionada: {tipo_insulina} para paciente {cod_cidadao}")
        return jsonify({
            "sucesso": True, 
            "mensagem": f"{tipo_insulina} adicionada com sucesso",
            "cod_seq_insulina": cod_seq_insulina
        })
        
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro ao adicionar insulina: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/insulinas/<int:cod_seq_insulina>/detalhes', methods=['GET'])
def api_diabetes_obter_detalhes_insulina(cod_seq_insulina):
    """API para obter detalhes de uma insulina específica para edição"""
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Obter detalhes da insulina
        query = """
            SELECT cod_seq_insulina, codcidadao, tipo_insulina, frequencia_dia, 
                   doses_estruturadas, data_inicio, data_fim, observacoes,
                   status, motivo_interrupcao, created_at, updated_at
            FROM sistemaaps.tb_hiperdia_dm_insulina 
            WHERE cod_seq_insulina = %s AND data_fim IS NULL
        """
        cur.execute(query, (cod_seq_insulina,))
        result = cur.fetchone()
        
        if not result:
            return jsonify({"sucesso": False, "erro": "Insulina não encontrada ou já foi interrompida"}), 404
        
        (cod_seq, codcidadao, tipo_insulina, frequencia_dia, doses_estruturadas, 
         data_inicio, data_fim, observacoes, status, motivo_interrupcao, 
         created_at, updated_at) = result
        
        # Parse das doses estruturadas
        doses = []
        if doses_estruturadas:
            try:
                doses = json.loads(doses_estruturadas)
            except json.JSONDecodeError:
                doses = []
        
        insulina_detalhes = {
            "cod_seq_insulina": cod_seq,
            "codcidadao": codcidadao,
            "tipo_insulina": tipo_insulina,
            "frequencia_dia": frequencia_dia,
            "doses_estruturadas": doses,
            "data_inicio": data_inicio.strftime('%Y-%m-%d') if data_inicio else None,
            "observacoes": observacoes or '',
            "status": status or 'ATIVO'
        }
        
        return jsonify({"sucesso": True, "insulina": insulina_detalhes})
        
    except Exception as e:
        print(f"Erro ao obter detalhes da insulina {cod_seq_insulina}: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/insulinas/<int:cod_seq_insulina>', methods=['PUT'])
def api_diabetes_atualizar_insulina(cod_seq_insulina):
    """API para modificar insulina de paciente diabético (cria novo registro mantendo histórico)"""
    conn = None
    cur = None
    try:
        data = request.get_json()
        motivo_modificacao = data.get('motivo_modificacao', 'Modificação de dosagem/frequência')
        
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Obter dados da insulina atual
        select_query = """
            SELECT codcidadao, tipo_insulina, frequencia_dia, doses_estruturadas, 
                   data_inicio, observacoes 
            FROM sistemaaps.tb_hiperdia_dm_insulina 
            WHERE cod_seq_insulina = %s AND data_fim IS NULL
        """
        cur.execute(select_query, (cod_seq_insulina,))
        insulina_atual = cur.fetchone()
        
        if not insulina_atual:
            return jsonify({"sucesso": False, "erro": "Insulina não encontrada ou já foi interrompida"}), 404
        
        codcidadao, tipo_atual, freq_atual, doses_atuais, data_inicio_atual, obs_atuais = insulina_atual
        
        # Validar dados de entrada
        nova_frequencia = data.get('frequencia_dia', freq_atual)
        if nova_frequencia not in [1, 2, 3, 4]:
            return jsonify({"sucesso": False, "erro": "Frequência deve ser entre 1 e 4 vezes ao dia"}), 400
        
        novo_tipo = data.get('tipo_insulina', tipo_atual)
        tipos_validos = ['Insulina NPH', 'Insulina Regular', 'Insulina Glargina', 'Insulina Lispro']
        if novo_tipo not in tipos_validos:
            return jsonify({"sucesso": False, "erro": f"Tipo de insulina deve ser um dos: {', '.join(tipos_validos)}"}), 400
        
        novas_doses = data.get('doses_estruturadas', json.loads(doses_atuais) if doses_atuais else [])
        # Validar doses
        if novas_doses:
            for dose in novas_doses:
                if not isinstance(dose, dict) or 'dose' not in dose or 'horario' not in dose:
                    return jsonify({"sucesso": False, "erro": "Formato de doses inválido"}), 400
                
                dose_valor = dose.get('dose')
                if not isinstance(dose_valor, int) or dose_valor < 1 or dose_valor > 100:
                    return jsonify({"sucesso": False, "erro": "Dose deve ser entre 1 e 100 unidades"}), 400
        
        nova_data_inicio = data.get('data_inicio')
        if nova_data_inicio:
            try:
                datetime.strptime(nova_data_inicio, '%Y-%m-%d')
            except ValueError:
                return jsonify({"sucesso": False, "erro": "Data de início inválida"}), 400
        
        novas_observacoes = data.get('observacoes', obs_atuais or '')
        
        # Verificar se houve mudança
        doses_atuais_parsed = json.loads(doses_atuais) if doses_atuais else []
        if (nova_frequencia == freq_atual and 
            novo_tipo == tipo_atual and 
            novas_doses == doses_atuais_parsed and
            novas_observacoes == (obs_atuais or '') and
            nova_data_inicio == (data_inicio_atual.strftime('%Y-%m-%d') if data_inicio_atual else None)):
            return jsonify({"sucesso": False, "erro": "Nenhuma modificação foi detectada"}), 400
        
        # Marcar insulina atual como substituída
        cur.execute("""
            UPDATE sistemaaps.tb_hiperdia_dm_insulina 
            SET status = 'SUBSTITUIDO', 
                data_fim = CURRENT_DATE,
                motivo_interrupcao = %s,
                updated_at = CURRENT_TIMESTAMP
            WHERE cod_seq_insulina = %s
        """, (motivo_modificacao, cod_seq_insulina))
        
        # Criar novo registro de insulina
        insert_query = """
            INSERT INTO sistemaaps.tb_hiperdia_dm_insulina 
            (codcidadao, tipo_insulina, frequencia_dia, doses_estruturadas, 
             data_inicio, observacoes, status, cod_insulina_anterior, motivo_modificacao,
             created_at, updated_at)
            VALUES (%s, %s, %s, %s, %s, %s, 'ATIVO', %s, %s, 
                    CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
            RETURNING cod_seq_insulina
        """
        
        cur.execute(insert_query, (
            codcidadao, novo_tipo, nova_frequencia, json.dumps(novas_doses),
            nova_data_inicio or data_inicio_atual, novas_observacoes,
            cod_seq_insulina, motivo_modificacao
        ))
        
        novo_cod_seq = cur.fetchone()[0]
        conn.commit()
        
        print(f"Insulina {cod_seq_insulina} modificada para paciente {codcidadao}. Nova insulina: {novo_cod_seq}")
        return jsonify({
            "sucesso": True, 
            "mensagem": "Insulina modificada com sucesso",
            "novo_cod_seq": novo_cod_seq
        })
        
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro ao modificar insulina {cod_seq_insulina}: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/insulinas/<int:cod_seq_insulina>/interromper', methods=['PUT'])
def api_diabetes_interromper_insulina(cod_seq_insulina):
    """API para interromper insulina de paciente diabético"""
    conn = None
    cur = None
    try:
        data = request.get_json()
        motivo_interrupcao = data.get('motivo_interrupcao', 'Insulina interrompida pelo profissional')
        
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Verificar se a insulina existe e está ativa
        check_query = """
            SELECT codcidadao, tipo_insulina FROM sistemaaps.tb_hiperdia_dm_insulina 
            WHERE cod_seq_insulina = %s AND data_fim IS NULL
        """
        cur.execute(check_query, (cod_seq_insulina,))
        insulina_atual = cur.fetchone()
        
        if not insulina_atual:
            return jsonify({"sucesso": False, "erro": "Insulina não encontrada ou já foi interrompida"}), 404
        
        # Interromper insulina
        update_query = """
            UPDATE sistemaaps.tb_hiperdia_dm_insulina 
            SET data_fim = CURRENT_DATE,
                motivo_interrupcao = %s,
                updated_at = CURRENT_TIMESTAMP
            WHERE cod_seq_insulina = %s
        """
        
        cur.execute(update_query, (motivo_interrupcao, cod_seq_insulina))
        conn.commit()
        
        print(f"Insulina {insulina_atual[1]} interrompida para paciente {insulina_atual[0]}")
        return jsonify({"sucesso": True, "mensagem": "Insulina interrompida com sucesso"})
        
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro ao interromper insulina {cod_seq_insulina}: {e}")
        return jsonify({"sucesso": False, "erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/generate_prescriptions_pdf', methods=['POST'])
def api_diabetes_generate_prescriptions_pdf():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        data = request.json
        patients = data.get('patients', [])
        
        if not patients:
            return jsonify({"erro": "Nenhum paciente selecionado"}), 400
        
        # Diretório temporário para armazenar os PDFs individuais
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_files = []
            
            for patient in patients:
                try:
                    # Buscar dados do paciente
                    sql_paciente = """
                        SELECT cod_paciente, nome_paciente, dt_nascimento, sexo, cartao_sus
                        FROM sistemaaps.mv_hiperdia_diabetes
                        WHERE cod_paciente = %(cod_paciente)s
                    """
                    
                    cur.execute(sql_paciente, {'cod_paciente': patient['cod_paciente']})
                    paciente_dados = cur.fetchone()
                    
                    if not paciente_dados:
                        continue
                        
                    paciente_dict = dict(paciente_dados)
                    
                    # Buscar medicamentos orais ativos
                    sql_medicamentos = """
                        SELECT 
                            nome_medicamento,
                            dose,
                            frequencia,
                            posologia,
                            data_inicio,
                            observacoes,
                            created_at
                        FROM sistemaaps.tb_hiperdia_dm_medicamentos
                        WHERE codcidadao = %(cod_paciente)s
                        AND (data_fim IS NULL OR data_fim > CURRENT_DATE)
                        ORDER BY created_at DESC, nome_medicamento
                    """
                    
                    cur.execute(sql_medicamentos, {'cod_paciente': patient['cod_paciente']})
                    medicamentos = cur.fetchall()
                    
                    # Buscar insulinas ativas
                    sql_insulinas = """
                        SELECT 
                            tipo_insulina,
                            doses_estruturadas,
                            data_inicio,
                            observacoes,
                            created_at
                        FROM sistemaaps.tb_hiperdia_dm_insulina
                        WHERE codcidadao = %(cod_paciente)s
                        AND (data_fim IS NULL OR data_fim > CURRENT_DATE)
                        ORDER BY created_at DESC, tipo_insulina
                    """
                    
                    cur.execute(sql_insulinas, {'cod_paciente': patient['cod_paciente']})
                    insulinas = cur.fetchall()
                    
                    if not medicamentos and not insulinas:
                        continue
                    
                    # Combinar medicamentos e insulinas em uma lista usando as novas funções de formatação
                    medicamentos_lista = []
                    
                    # Adicionar medicamentos orais usando função de formatação
                    for med in medicamentos:
                        resultado = formatar_medicamento_oral_diabetes(
                            med['nome_medicamento'],
                            med['dose'],
                            med['frequencia'],
                            med['posologia'],
                            med['observacoes']
                        )
                        
                        medicamentos_lista.append({
                            'nome': resultado['nome'],
                            'quantidade': resultado['quantidade'],
                            'instrucoes': resultado['instrucoes'],
                            'created_at': med['created_at']
                        })
                    
                    # Adicionar insulinas usando função de formatação
                    for ins in insulinas:
                        try:
                            resultado = formatar_insulina_diabetes(
                                ins['tipo_insulina'],
                                ins['doses_estruturadas'],
                                ins['observacoes']
                            )
                            
                            medicamentos_lista.append({
                                'nome': resultado['nome'],
                                'quantidade': resultado['quantidade'],
                                'instrucoes': resultado['instrucoes'],
                                'created_at': ins['created_at']
                            })
                            
                        except Exception as e:
                            print(f"Erro ao processar doses de insulina para paciente {patient['cod_paciente']}: {e}")
                            continue
                    
                    if not medicamentos_lista:
                        continue
                        
                    # Aplicar a mesma lógica de fonte do sistema de hipertensão
                    num_medicamentos = len(medicamentos_lista)
                    if num_medicamentos <= 2:
                        font_size = 14
                    elif num_medicamentos == 3:
                        font_size = 14
                    elif num_medicamentos == 4:
                        font_size = 12
                    elif num_medicamentos == 5:
                        font_size = 12
                    elif num_medicamentos == 6:
                        font_size = 11
                    elif num_medicamentos == 7:
                        font_size = 10
                    else:
                        font_size = max(6, 16 - num_medicamentos)
                    
                    # Definir quantidade de traços baseada no tamanho da fonte
                    if num_medicamentos <= 3:
                        tracos = "--------------------------------"  # 32 traços
                    else:
                        tracos = "---------------------------------------------"  # 45 traços
                    
                    # Gerar texto completo de medicamentos dinamicamente
                    medicamentos_texto = ""
                    for idx, med in enumerate(medicamentos_lista, 1):
                        medicamentos_texto += f"{idx}) {med['nome']} {tracos} {med['quantidade']} \n"
                        
                        for instrucao in med['instrucoes']:
                            medicamentos_texto += f"{instrucao}\n"
                        
                        if idx < len(medicamentos_lista):
                            medicamentos_texto += "\n"
                    
                    # Encontrar data mais recente (created_at mais recente)
                    datas_created = []
                    if medicamentos:
                        datas_created.extend([med['created_at'] for med in medicamentos if med['created_at']])
                    if insulinas:
                        datas_created.extend([ins['created_at'] for ins in insulinas if ins['created_at']])
                    
                    data_mais_recente = max(datas_created) if datas_created else None
                    
                    print(f"DEBUG: {num_medicamentos} medicamentos/insulinas, fonte {font_size}pt")
                    
                    # Preparar contexto para o template
                    context = {
                        'nome_paciente': remove_acentos(paciente_dict['nome_paciente'].upper()),
                        'data_nascimento': paciente_dict['dt_nascimento'].strftime('%d/%m/%Y') if paciente_dict['dt_nascimento'] else "xx/xx/xxxx",
                        'sexo': paciente_dict.get('sexo', 'Não informado'),
                        'cns': paciente_dict['cartao_sus'] if paciente_dict['cartao_sus'] else "CNS não registrado no PEC",
                        'ultima_atualizacao': data_mais_recente.strftime('%d/%m/%Y') if data_mais_recente else "Não disponível",
                        'medicamentos_texto': medicamentos_texto,
                        'font_size': font_size
                    }
                    
                    # Processar template
                    template_path = os.path.join(os.path.dirname(__file__), 'modelos', 'template_receituario_diabetes.docx')
                    temp_docx = os.path.join(temp_dir, f"receituario_diabetes_{patient['cod_paciente']}.docx")
                    temp_pdf = os.path.join(temp_dir, f"receituario_diabetes_{patient['cod_paciente']}.pdf")
                    
                    # Verificar se o template específico de diabetes existe, senão usar o padrão
                    if not os.path.exists(template_path):
                        template_path = os.path.join(os.path.dirname(__file__), 'modelos', 'template_receituario_dynamic.docx')
                    
                    # Processar template
                    doc_template = DocxTemplate(template_path)
                    doc_template.render(context)
                    doc_template.save(temp_docx)
                    
                    # Aplicar formatação específica
                    print(f"DEBUG: Aplicando formatação de fonte {font_size}pt e negrito...")
                    from docx import Document
                    from docx.shared import Pt
                    doc = Document(temp_docx)
                    for paragraph in doc.paragraphs:
                        # Substituir "Hipertensão" por "Diabetes" em todo o documento
                        for run in paragraph.runs:
                            if 'Hipertensão' in run.text:
                                run.text = run.text.replace('Hipertensão', 'Diabetes')
                        
                        if 'medicamentos_texto' in paragraph.text or any('medicamentos_texto' in run.text for run in paragraph.runs):
                            for run in paragraph.runs:
                                if 'medicamentos_texto' in run.text:
                                    run.text = run.text.replace('{{ medicamentos_texto }}', medicamentos_texto)
                                    run.font.size = Pt(font_size)
                                    run.font.name = 'Arial'
                                elif any(char.isdigit() and ')' in run.text for char in run.text[:3]):
                                    run.font.size = Pt(font_size)
                                    run.font.bold = True
                                else:
                                    run.font.size = Pt(font_size)
                    
                    doc.save(temp_docx)
                    
                    # Converter para PDF
                    import pythoncom
                    import win32com.client
                    pythoncom.CoInitialize()
                    word = None
                    try:
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        doc_word = word.Documents.Open(temp_docx)
                        doc_word.SaveAs(temp_pdf, 17)  # 17 = PDF format
                        doc_word.Close()
                        word.Quit()
                    except Exception as e:
                        print(f"Erro na conversão Word->PDF: {e}")
                        if word:
                            try:
                                word.Quit()
                            except:
                                pass
                    finally:
                        pythoncom.CoUninitialize()
                    
                    # Aguardar e verificar se o arquivo está acessível
                    max_retries = 5
                    retry_delay = 0.3
                    
                    for attempt in range(max_retries):
                        time.sleep(retry_delay)
                        
                        # Verificar se o arquivo existe e está acessível
                        if os.path.exists(temp_pdf):
                            try:
                                # Tentar abrir o arquivo para verificar se está liberado
                                with open(temp_pdf, 'rb') as test_file:
                                    test_file.read(1024)  # Ler um pouco do arquivo
                                print(f"PDF do paciente {patient['cod_paciente']} liberado após {attempt + 1} tentativas")
                                pdf_files.append(temp_pdf)
                                break
                            except (PermissionError, OSError) as e:
                                print(f"Paciente {patient['cod_paciente']} - Tentativa {attempt + 1}: Arquivo ainda bloqueado - {e}")
                                if attempt == max_retries - 1:
                                    print(f"Pulando paciente {patient['cod_paciente']} - arquivo não liberado")
                        else:
                            print(f"Paciente {patient['cod_paciente']} - Tentativa {attempt + 1}: PDF ainda não existe")
                            if attempt == max_retries - 1:
                                print(f"Pulando paciente {patient['cod_paciente']} - PDF não foi criado")
                    
                except Exception as e:
                    print(f"Erro ao processar paciente {patient.get('cod_paciente', 'N/A')}: {e}")
                    continue
            
            if not pdf_files:
                return jsonify({"erro": "Nenhum receituário foi gerado com sucesso"}), 404
            
            # Combinar todos os PDFs em um único arquivo
            combined_pdf_path = os.path.join(temp_dir, "receituarios_diabetes_combinados.pdf")
            
            with open(combined_pdf_path, 'wb') as combined_file:
                writer = PdfWriter()
                
                for pdf_path in pdf_files:
                    with open(pdf_path, 'rb') as pdf_file:
                        reader = PdfReader(pdf_file)
                        for page in reader.pages:
                            writer.add_page(page)
                
                writer.write(combined_file)
            
            # Retornar PDF combinado
            timestamp = datetime.now().strftime('%d%m%Y_%H%M%S')
            return send_file(
                combined_pdf_path,
                as_attachment=True,
                download_name=f"receituarios_diabetes_{timestamp}.pdf",
                mimetype='application/pdf'
            )
            
    except Exception as e:
        print(f"Erro na API /api/diabetes/generate_prescriptions_pdf: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/generate_prescription_pdf_individual', methods=['POST'])
def api_diabetes_generate_prescription_pdf_individual():
    conn = None
    cur = None
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        data = request.json
        patient = data.get('patient')
        
        if not patient or 'cod_paciente' not in patient:
            return jsonify({"erro": "Dados do paciente inválidos"}), 400
            
        # Buscar dados do paciente
        sql_paciente = """
            SELECT cod_paciente, nome_paciente, dt_nascimento, sexo, cartao_sus
            FROM sistemaaps.mv_hiperdia_diabetes
            WHERE cod_paciente = %(cod_paciente)s
        """
        
        cur.execute(sql_paciente, {'cod_paciente': patient['cod_paciente']})
        paciente_dados = cur.fetchone()
        
        if not paciente_dados:
            return jsonify({"erro": "Paciente não encontrado"}), 404
            
        paciente_dict = dict(paciente_dados)
        
        # Buscar medicamentos orais ativos
        sql_medicamentos = """
            SELECT 
                nome_medicamento,
                dose,
                frequencia,
                posologia,
                data_inicio,
                observacoes,
                created_at
            FROM sistemaaps.tb_hiperdia_dm_medicamentos
            WHERE codcidadao = %(cod_paciente)s
            AND (data_fim IS NULL OR data_fim > CURRENT_DATE)
            ORDER BY created_at DESC, nome_medicamento
        """
        
        cur.execute(sql_medicamentos, {'cod_paciente': patient['cod_paciente']})
        medicamentos = cur.fetchall()
        
        # Buscar insulinas ativas
        sql_insulinas = """
            SELECT 
                tipo_insulina,
                doses_estruturadas,
                data_inicio,
                observacoes,
                created_at
            FROM sistemaaps.tb_hiperdia_dm_insulina
            WHERE codcidadao = %(cod_paciente)s
            AND (data_fim IS NULL OR data_fim > CURRENT_DATE)
            ORDER BY created_at DESC, tipo_insulina
        """
        
        cur.execute(sql_insulinas, {'cod_paciente': patient['cod_paciente']})
        insulinas = cur.fetchall()
        
        if not medicamentos and not insulinas:
            return jsonify({"erro": "Nenhum medicamento ou insulina ativo encontrado para este paciente"}), 404
        
        # Combinar medicamentos e insulinas em uma lista usando as novas funções de formatação
        medicamentos_lista = []
        
        # Adicionar medicamentos orais com formatação correta
        for med in medicamentos:
            med_formatado = formatar_medicamento_oral_diabetes(
                nome=med['nome_medicamento'],
                dose=med['dose'] or 1,
                frequencia=med['frequencia'] or 1,
                posologia=med['posologia'],
                observacoes=med['observacoes']
            )
            
            # Adicionar created_at para cálculo da data
            med_formatado['created_at'] = med['created_at']
            medicamentos_lista.append(med_formatado)
        
        # Adicionar insulinas com formatação correta
        for ins in insulinas:
            ins_formatada = formatar_insulina_diabetes(
                tipo_insulina=ins['tipo_insulina'],
                doses_estruturadas_json=ins['doses_estruturadas'],
                observacoes=ins['observacoes']
            )
            
            if ins_formatada:
                # Adicionar created_at para cálculo da data
                ins_formatada['created_at'] = ins['created_at']
                medicamentos_lista.append(ins_formatada)
        
        if not medicamentos_lista:
            return jsonify({"erro": "Erro ao processar medicamentos e insulinas"}), 500
        
        # Aplicar a mesma lógica de fonte do sistema de hipertensão
        num_medicamentos = len(medicamentos_lista)
        if num_medicamentos <= 2:
            font_size = 14
        elif num_medicamentos == 3:
            font_size = 14
        elif num_medicamentos == 4:
            font_size = 12
        elif num_medicamentos == 5:
            font_size = 12
        elif num_medicamentos == 6:
            font_size = 11
        elif num_medicamentos == 7:
            font_size = 10
        else:
            font_size = max(6, 16 - num_medicamentos)
        
        # Definir quantidade de traços baseada no tamanho da fonte
        if num_medicamentos <= 3:
            tracos = "--------------------------------"  # 32 traços
        else:
            tracos = "---------------------------------------------"  # 45 traços
        
        # Gerar texto completo de medicamentos dinamicamente
        medicamentos_texto = ""
        for idx, med in enumerate(medicamentos_lista, 1):
            medicamentos_texto += f"{idx}) {med['nome']} {tracos} {med['quantidade']} \n"
            
            for instrucao in med['instrucoes']:
                medicamentos_texto += f"{instrucao}\n"
            
            if idx < len(medicamentos_lista):
                medicamentos_texto += "\n"
        
        # Encontrar data mais recente (created_at mais recente)
        datas_created = []
        if medicamentos:
            datas_created.extend([med['created_at'] for med in medicamentos if med['created_at']])
        if insulinas:
            datas_created.extend([ins['created_at'] for ins in insulinas if ins['created_at']])
        
        data_mais_recente = max(datas_created) if datas_created else None
        
        # Preparar contexto para o template
        context = {
            'nome_paciente': remove_acentos(paciente_dict['nome_paciente'].upper()),
            'data_nascimento': paciente_dict['dt_nascimento'].strftime('%d/%m/%Y') if paciente_dict['dt_nascimento'] else "xx/xx/xxxx",
            'sexo': paciente_dict.get('sexo', 'Não informado'),
            'cns': paciente_dict['cartao_sus'] if paciente_dict['cartao_sus'] else "CNS não registrado no PEC",
            'ultima_atualizacao': data_mais_recente.strftime('%d/%m/%Y') if data_mais_recente else "Não disponível",
            'medicamentos_texto': medicamentos_texto,
            'font_size': font_size
        }
        
        # Criar arquivo temporário
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = os.path.join(os.path.dirname(__file__), 'modelos', 'template_receituario_diabetes.docx')
            temp_docx = os.path.join(temp_dir, f"receituario_diabetes_{patient['cod_paciente']}.docx")
            temp_pdf = os.path.join(temp_dir, f"receituario_diabetes_{patient['cod_paciente']}.pdf")
            
            # Verificar se o template específico de diabetes existe, senão usar o padrão
            if not os.path.exists(template_path):
                template_path = os.path.join(os.path.dirname(__file__), 'modelos', 'template_receituario_dynamic.docx')
            
            # Processar template
            doc_template = DocxTemplate(template_path)
            doc_template.render(context)
            doc_template.save(temp_docx)
            
            # Aplicar formatação específica
            from docx import Document
            from docx.shared import Pt
            doc = Document(temp_docx)
            for paragraph in doc.paragraphs:
                # Substituir "Hipertensão" por "Diabetes" em todo o documento
                for run in paragraph.runs:
                    if 'Hipertensão' in run.text:
                        run.text = run.text.replace('Hipertensão', 'Diabetes')
                
                if 'medicamentos_texto' in paragraph.text or any('medicamentos_texto' in run.text for run in paragraph.runs):
                    for run in paragraph.runs:
                        if 'medicamentos_texto' in run.text:
                            run.text = run.text.replace('{{ medicamentos_texto }}', medicamentos_texto)
                            run.font.size = Pt(font_size)
                            run.font.name = 'Arial'
                        elif any(char.isdigit() and ')' in run.text for char in run.text[:3]):
                            run.font.size = Pt(font_size)
                            run.font.bold = True
                        else:
                            run.font.size = Pt(font_size)
            
            doc.save(temp_docx)
            
            # Converter para PDF
            import pythoncom
            import win32com.client
            pythoncom.CoInitialize()
            word = None
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc_word = word.Documents.Open(temp_docx)
                doc_word.SaveAs(temp_pdf, 17)  # 17 = PDF format
                doc_word.Close()
                word.Quit()
            except Exception as e:
                print(f"Erro na conversão Word->PDF: {e}")
                if word:
                    try:
                        word.Quit()
                    except:
                        pass
                raise e  # Re-raise para ser capturado pelo try/except externo
            finally:
                pythoncom.CoUninitialize()
            
            # Aguardar e verificar se o arquivo está acessível
            max_retries = 10
            retry_delay = 0.5
            
            for attempt in range(max_retries):
                time.sleep(retry_delay)
                
                # Verificar se o arquivo existe e está acessível
                if os.path.exists(temp_pdf):
                    try:
                        # Tentar abrir o arquivo para verificar se está liberado
                        with open(temp_pdf, 'rb') as test_file:
                            test_file.read(1024)  # Ler um pouco do arquivo
                        print(f"Arquivo PDF liberado após {attempt + 1} tentativas")
                        break
                    except (PermissionError, OSError) as e:
                        print(f"Tentativa {attempt + 1}: Arquivo ainda bloqueado - {e}")
                        if attempt == max_retries - 1:
                            raise Exception(f"Arquivo PDF não foi liberado após {max_retries} tentativas: {e}")
                else:
                    print(f"Tentativa {attempt + 1}: Arquivo PDF ainda não existe")
                    if attempt == max_retries - 1:
                        raise Exception("Arquivo PDF não foi criado")
            
            # Ler o arquivo PDF em memória para evitar problemas de acesso
            final_pdf_name = f"receituario_diabetes_{remove_acentos(paciente_dict['nome_paciente'])}_{datetime.now().strftime('%d%m%Y')}.pdf"
            
            with open(temp_pdf, 'rb') as pdf_file:
                pdf_data = pdf_file.read()
            
            print(f"Arquivo PDF lido em memória ({len(pdf_data)} bytes)")
            
            # Criar BytesIO para enviar
            pdf_buffer = io.BytesIO(pdf_data)
            pdf_buffer.seek(0)
            
            # Retornar PDF
            return send_file(
                pdf_buffer,
                as_attachment=True,
                download_name=final_pdf_name,
                mimetype='application/pdf'
            )
            
    except Exception as e:
        print(f"Erro na API /api/diabetes/generate_prescription_pdf_individual: {e}")
        return jsonify({"erro": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/api/diabetes/avaliar_tratamento', methods=['POST'])
def api_diabetes_avaliar_tratamento():
    """API para avaliar tratamento de diabético com exames e múltiplos mapeamentos"""
    conn = None
    cur = None
    try:
        data = request.get_json()

        conn = get_db_connection()
        cur = conn.cursor()

        # Inserir ação de acompanhamento principal
        acao_query = """
            INSERT INTO sistemaaps.tb_hiperdia_dm_acompanhamento
            (cod_cidadao, cod_acao, status_acao, data_agendamento, data_realizacao, observacoes, responsavel_pela_acao)
            VALUES (%(cod_cidadao)s, %(cod_acao)s, %(status_acao)s, %(data_agendamento)s, %(data_realizacao)s, %(observacoes)s, %(responsavel_pela_acao)s)
            RETURNING cod_acompanhamento
        """

        status_acao = 'REALIZADA'  # Avaliação já foi realizada
        data_realizacao = data.get('data_agendamento')

        cur.execute(acao_query, {
            'cod_cidadao': data['cod_cidadao'],
            'cod_acao': data['cod_acao'],
            'status_acao': status_acao,
            'data_agendamento': data.get('data_agendamento'),
            'data_realizacao': data_realizacao,
            'observacoes': data.get('observacoes'),
            'responsavel_pela_acao': data.get('responsavel_pela_acao')
        })

        cod_acompanhamento = cur.fetchone()[0]

        # Inserir dados de exames laboratoriais se fornecidos
        if 'exames' in data and data['exames']:
            exames = data['exames']
            if any([exames.get('hemoglobina_glicada'), exames.get('glicemia_media'), exames.get('glicemia_jejum')]):
                exames_query = """
                    INSERT INTO sistemaaps.tb_hiperdia_dm_exames
                    (cod_acompanhamento, hemoglobina_glicada, glicemia_media, glicemia_jejum, data_exame)
                    VALUES (%(cod_acompanhamento)s, %(hemoglobina_glicada)s, %(glicemia_media)s, %(glicemia_jejum)s, %(data_exame)s)
                """

                cur.execute(exames_query, {
                    'cod_acompanhamento': cod_acompanhamento,
                    'hemoglobina_glicada': exames.get('hemoglobina_glicada'),
                    'glicemia_media': exames.get('glicemia_media'),
                    'glicemia_jejum': exames.get('glicemia_jejum'),
                    'data_exame': exames.get('data_exame')
                })

        # Inserir múltiplos mapeamentos residenciais se fornecidos
        if 'mapeamentos' in data and data['mapeamentos']:
            mrg_query = """
                INSERT INTO sistemaaps.tb_hiperdia_mrg
                (cod_acompanhamento, data_mrg, g_jejum, g_apos_cafe, g_antes_almoco, g_apos_almoco, g_antes_jantar, g_ao_deitar, periodo_mapeamento, dias_mapeamento)
                VALUES (%(cod_acompanhamento)s, %(data_mrg)s, %(g_jejum)s, %(g_apos_cafe)s, %(g_antes_almoco)s, %(g_apos_almoco)s, %(g_antes_jantar)s, %(g_ao_deitar)s, %(periodo_mapeamento)s, %(dias_mapeamento)s)
            """

            for mapeamento in data['mapeamentos']:
                cur.execute(mrg_query, {
                    'cod_acompanhamento': cod_acompanhamento,
                    'data_mrg': mapeamento.get('data_mrg'),
                    'g_jejum': mapeamento.get('g_jejum'),
                    'g_apos_cafe': mapeamento.get('g_apos_cafe'),
                    'g_antes_almoco': mapeamento.get('g_antes_almoco'),
                    'g_apos_almoco': mapeamento.get('g_apos_almoco'),
                    'g_antes_jantar': mapeamento.get('g_antes_jantar'),
                    'g_ao_deitar': mapeamento.get('g_ao_deitar'),
                    'periodo_mapeamento': mapeamento.get('periodo', 'Período 1'),
                    'dias_mapeamento': mapeamento.get('dias_mapeamento', 7)
                })

        # Atualizar status do controle do paciente se fornecido
        if 'status_controle' in data and data['status_controle']:
            # Aqui você pode adicionar lógica para atualizar o status geral do paciente
            # Por exemplo, atualizar uma tabela de status ou adicionar em observações
            pass

        conn.commit()

        return jsonify({
            "success": True,
            "cod_acompanhamento": cod_acompanhamento,
            "message": "Avaliação de tratamento salva com sucesso"
        })

    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Erro ao salvar avaliação de tratamento: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": f"Erro no servidor: {e}"}), 500
    finally:
        if cur: cur.close()
        if conn: conn.close()

if __name__ == '__main__':
    app.run(debug=True, port=3030, host='0.0.0.0')